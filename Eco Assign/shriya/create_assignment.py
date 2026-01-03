#!/usr/bin/env python3
"""
Create comprehensive Microeconomics Product Assignment on Hair Straightening Brush
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_page_margins(doc):
    """Set A4 page with specified margins"""
    for section in doc.sections:
        section.page_width = Inches(8.27)  # A4 width
        section.page_height = Inches(11.69)  # A4 height
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1)

def add_page_numbers(doc):
    """Add page numbers to footer"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

def add_heading(doc, text, level=1):
    """Add heading with proper formatting"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.5
    return para

def add_subheading(doc, text):
    """Add subheading with proper formatting"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.5
    return para

def add_body_text(doc, text):
    """Add body paragraph with proper formatting"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(8)
    return para

def add_table(doc, headers, rows):
    """Add formatted table"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for para in header_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data)
            for para in row_cells[col_idx].paragraphs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

    doc.add_paragraph()  # Space after table
    return table

def create_cover_page(doc):
    """Create cover page"""
    # Add spacing at top
    for _ in range(3):
        doc.add_paragraph()

    # Institute name
    para = doc.add_paragraph()
    run = para.add_run("SYDENHAM INSTITUTE OF MANAGEMENT STUDIES,")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run("RESEARCH AND ENTREPRENEURSHIP EDUCATION")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run("(SIMSREE)")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # Assignment title
    para = doc.add_paragraph()
    run = para.add_run("MICROECONOMICS PRODUCT ASSIGNMENT")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run("A Comprehensive Study of")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run("HAIR STRAIGHTENING BRUSH")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(20)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Student details
    details = [
        ("Name:", "Shriya Parab"),
        ("Roll No:", "HBM25120"),
        ("Class:", "MMS'27"),
        ("Semester:", "01"),
        ("Academic Year:", "2025-26"),
        ("Subject:", "Managerial Economics"),
    ]

    for label, value in details:
        para = doc.add_paragraph()
        run1 = para.add_run(f"{label} ")
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(12)
        run2 = para.add_run(value)
        run2.bold = True
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Page break
    doc.add_page_break()

def main():
    doc = Document()

    # Set page formatting
    set_page_margins(doc)
    add_page_numbers(doc)

    # Create cover page
    create_cover_page(doc)

    # ============================================
    # SECTION 1: INTRODUCTION AND EVOLUTION
    # ============================================
    add_heading(doc, "1. INTRODUCTION AND EVOLUTION OF HAIR STRAIGHTENING BRUSH")

    add_subheading(doc, "1.1 Definition and Economic Relevance")

    add_body_text(doc, """The hair straightening brush represents a convergence of thermal engineering and personal grooming technology, functioning as a hybrid appliance that combines heated bristles with traditional brush design to straighten hair through thermal reconditioning of keratin bonds. Unlike conventional flat irons that apply direct plate-to-hair compression, straightening brushes distribute heat across multiple contact points, reducing styling time while maintaining ergonomic handling. From a microeconomic perspective, this product occupies a distinct market position within the broader USD 30 billion global hair styling tools industry, addressing consumer demand for time-efficient, damage-minimizing hair care solutions.""")

    add_body_text(doc, """The economic significance of hair straightening brushes extends beyond their functional utility. They represent a classic example of product innovation responding to evolving consumer preferences and technological capabilities. The global market for hair straightening brushes reached USD 620-750 million in 2024, with projections indicating growth to USD 1.53 billion by 2033 at a compound annual growth rate (CAGR) of 8.2 percent. In India specifically, the market stood at USD 22.17 million in 2024, expected to reach USD 33.32 million by 2033 with a CAGR of 4.28 percent, reflecting the product's penetration into emerging markets with expanding middle-class populations.""")

    add_body_text(doc, """In the Indian market, hair straightening brushes are generally available in the price range of Rs. 1,500 to Rs. 3,500, depending on factors such as brand reputation, quality of materials used, technological features, and safety mechanisms. Premium brands may charge higher prices due to better build quality, advanced temperature control, and stronger brand perception. The product has gained popularity mainly in urban and semi-urban regions among students, working professionals, and young consumers who seek efficient grooming solutions within limited time.""")

    add_subheading(doc, "1.2 Historical Evolution of Hair Styling Technology")

    add_body_text(doc, """The trajectory of hair straightening technology exemplifies the economic principle of creative destruction, where successive innovations displace existing methods through superior cost-benefit propositions. Archaeological evidence indicates that hair manipulation tools date to ancient civilizations, with combs crafted from wood, bone, and ivory serving primarily hygiene rather than styling functions. The commercialization of thermal hair straightening began in the 1870s when French hairstylist Marcel Grateau developed heated iron rods paired with chemical lotions. Despite achieving straightening effects, the combination frequently caused scalp burns and hair damage, illustrating an early market failure due to inadequate safety features.""")

    add_body_text(doc, """The breakthrough came in 1909 when Isaac K. Shero patented metal plates capable of gliding through hair, incorporating adjustable heat settings that would become standard in modern devices. This innovation reduced transaction costs associated with professional salon visits by enabling home usage. The twentieth century witnessed exponential technological advancement driven by electrification and materials science. The marcel iron's introduction of heat regulation mechanisms represented a significant improvement in user control, reducing information asymmetry between manufacturers and consumers regarding safe operating temperatures.""")

    add_body_text(doc, """The adoption of plastic materials through injection moulding in the mid-1900s democratized access by dramatically lowering production costs, a textbook example of how manufacturing process innovations can expand market size through reduced average costs. Contemporary hair straightening brushes emerged in the early 2000s as manufacturers recognized an underserved market segment: consumers seeking straightening results with minimal technical skill and time investment. By integrating heating elements into brush formats familiar from daily grooming routines, designers reduced the learning curve and switching costs from traditional methods.""")

    add_subheading(doc, "1.3 Technological Innovations in Hair Straightening Brushes")

    add_body_text(doc, """Modern hair straightening brushes incorporate several technological innovations that create differentiated value propositions and enable premium pricing strategies. Ceramic heating technology, formed by pressing clay mixtures and firing at high temperatures, provides even heat distribution that minimizes hot spots capable of causing thermal damage. This addresses a principal-agent problem where consumers cannot perfectly observe product quality ex-ante, making materials certification a credible quality signal. Ceramic's superior heat conductivity relative to aluminium reduces energy requirements by 15-20 percent, translating to lower operating costs over the product lifecycle.""")

    add_body_text(doc, """Ionic technology represents another significant advancement. Devices equipped with ionic generators emit negatively charged ions that break down water molecules on hair surfaces, reducing drying time and static electricity. From an economic standpoint, ionic technology represents a form of process innovation that increases productivity (styled hair per unit time) without proportionally increasing inputs. Premium models leveraging this technology command price premiums of 40-60 percent over basic alternatives.""")

    add_body_text(doc, """PTC (Positive Temperature Coefficient) heating elements utilize semiconductor properties that cause electrical resistance to increase sharply at predetermined temperatures, typically 120 degrees Celsius, creating self-regulating heating without external thermostats. This innovation addresses safety concerns and reduces liability risks for manufacturers, effectively lowering expected costs associated with product recalls or litigation. Tourmaline and titanium coatings further enhance product performance. Tourmaline, a semi-precious mineral, when crushed and infused into ceramic plates, generates negative ions and far-infrared heat that penetrate hair shafts more effectively than conventional heating. Titanium plates offer superior heat conductivity and durability, justifying price premiums of 50-80 percent in professional-grade segments.""")

    add_subheading(doc, "1.4 Comparison with Substitutes")

    add_body_text(doc, """Hair straightening brushes compete with several substitute products in the personal grooming market. Traditional flat irons represent the closest substitute, offering precise straightening but requiring careful handling and higher skill levels. Blow dryers with straightening attachments provide an alternative for consumers seeking multi-functional devices, though they may not achieve the same level of straightening as dedicated tools. Professional salon services remain a substitute, particularly for consumers seeking expert styling or chemical straightening treatments, though these options involve higher costs and time investments.""")

    # Table: Comparison of Hair Styling Methods
    add_body_text(doc, "Table 1.1: Comparison of Hair Straightening Methods")
    comparison_headers = ["Attribute", "Straightening Brush", "Flat Iron", "Blow Dryer", "Salon Service"]
    comparison_rows = [
        ["Average Price (Rs.)", "1,500-3,500", "800-2,500", "1,000-3,000", "500-2,000/visit"],
        ["Ease of Use", "High", "Medium", "Medium", "Professional"],
        ["Time Required", "10-15 minutes", "15-25 minutes", "20-30 minutes", "60-120 minutes"],
        ["Heat Damage Risk", "Low", "Medium-High", "Low", "Variable"],
        ["Skill Requirement", "Low", "Medium", "Medium", "None (customer)"],
        ["Portability", "High", "High", "Medium", "Not applicable"],
    ]
    add_table(doc, comparison_headers, comparison_rows)

    add_subheading(doc, "1.5 Product Variants and Classifications")

    add_body_text(doc, """Hair straightening brushes are available in multiple variants designed to address different consumer needs and price sensitivities. Corded variants represent the most common type, offering consistent power supply and typically lower prices ranging from Rs. 1,200 to Rs. 2,500. Cordless variants provide enhanced portability and convenience but command price premiums of 40-60 percent due to battery technology costs, typically ranging from Rs. 2,500 to Rs. 5,000. Material-based differentiation includes ceramic-coated models emphasizing gentle heat distribution, ionic technology variants reducing frizz, and tourmaline-infused options for professional-grade results.""")

    add_body_text(doc, """Temperature control features further segment the market. Basic models offer fixed temperature settings suitable for general use, while advanced models provide adjustable temperature ranges from 150 to 230 degrees Celsius, allowing customization based on hair type and styling requirements. Smart temperature control models incorporate sensors that automatically adjust heat based on hair moisture levels, representing the premium segment targeting technology-conscious consumers willing to pay 80-100 percent premiums over basic models.""")

    add_subheading(doc, "1.6 Product Lifecycle Classification")

    add_body_text(doc, """Applying the standard product lifecycle framework comprising introduction, growth, maturity, and decline phases, hair straightening brushes currently occupy the growth-to-early-maturity transition phase. Several indicators support this classification. Market penetration data shows that 45 percent of women in developed markets owned hair straighteners by 2024, up from 35 percent in 2007, with straightening brushes representing approximately 30 percent of this category. The product has moved beyond early adopters into the mainstream majority, though significant untapped potential remains in emerging markets where ownership rates remain below 15 percent.""")

    add_body_text(doc, """The sustained CAGR of 5.2-8.2 percent across major markets indicates continued expansion rather than the sales plateauing characteristic of mature products. The introduction of cordless variants and smart-enabled models with app connectivity demonstrates ongoing innovation typical of growth-stage products. Market fragmentation with numerous competitors across price tiers indicates that dominant designs have not fully crystallized, consistent with late-growth phase competitive intensity. Premium brands continue investing heavily in research and development, behavior more typical of growth-phase firms seeking to establish sustainable competitive advantages.""")

    add_body_text(doc, """Unlike mature commodity markets characterized by price compression, hair straightening brushes maintain significant price dispersion from budget models at Rs. 300-1,200 to premium options at Rs. 15,000-45,000, suggesting that product differentiation remains viable and market segmentation persists. This lifecycle positioning has crucial implications for firm strategy and market structure, informing the subsequent analysis of production functions, cost structures, and competitive dynamics.""")

    # ============================================
    # SECTION 2: BACKEND ANALYSIS
    # ============================================
    add_heading(doc, "2. BACKEND ANALYSIS: PRODUCTION, COSTS, AND SUPPLY CHAIN")

    add_subheading(doc, "2.1 Production Modality and Technology")

    add_body_text(doc, """The production of hair straightening brushes exemplifies modern manufacturing's integration of materials science, electrical engineering, and industrial design. The manufacturing value chain can be decomposed into discrete stages, each with distinct cost structures and efficiency determinants. Hair straightening brushes are produced using electro-mechanical manufacturing processes involving the assembly of ceramic heating elements, electrical wiring, temperature control circuits, and plastic casings manufactured through injection moulding.""")

    add_body_text(doc, """Product development begins with computer-aided design (CAD) modelling to optimize heat distribution patterns, bristle configurations, and ergonomic handling characteristics. This stage incurs high fixed costs including research and development personnel and testing equipment, but near-zero marginal costs for design iterations. A typical product development cycle spans 12-18 months with expenditures ranging from USD 50,000 for budget models to over USD 2 million for premium innovations incorporating novel heating technologies. Engineers conduct finite element analysis to model thermal conductivity and identify optimal bristle spacing that maximizes hair contact while preventing heat concentration.""")

    add_body_text(doc, """Small and medium-scale manufacturers usually operate semi-automatic assembly lines producing approximately 500 to 2,000 units per day. Large manufacturers employ automated production systems capable of producing 5,000 units or more per day. Over time, production has shifted from labour-intensive methods to capital-intensive methods due to technological advancement and the need for cost efficiency. Automation increases output per hour, reduces dependence on manual labour, and ensures uniformity in product quality. In microeconomic terms, technological improvement increases productivity and leads to a downward shift in the average cost curve, enabling firms to produce larger quantities at lower per-unit costs in the long run.""")

    add_body_text(doc, "Table 2.1: Production Stages and Requirements")
    prod_headers = ["Production Stage", "Key Activities", "Capital Intensity", "Labour Requirement"]
    prod_rows = [
        ["Design & Prototyping", "CAD modelling, testing", "High", "Skilled engineers"],
        ["Component Manufacturing", "Heating elements, housings", "High", "Semi-skilled"],
        ["Electronic Assembly", "PCB assembly, wiring", "Medium-High", "Skilled technicians"],
        ["Final Assembly", "Integration, testing", "Medium", "Semi-skilled workers"],
        ["Quality Control", "Safety testing, calibration", "Medium", "Skilled inspectors"],
        ["Packaging", "Boxing, documentation", "Low", "Unskilled labour"],
    ]
    add_table(doc, prod_headers, prod_rows)

    add_subheading(doc, "2.2 Raw Materials and Factor Market Analysis")

    add_body_text(doc, """The production of hair straightening brushes requires a diverse combination of raw materials, labour, and capital inputs. Hair straightener production consumes diverse material inputs, each with distinct supply chain characteristics and cost determinants. Understanding the factor market is essential for analyzing production costs and supply chain efficiency.""")

    add_body_text(doc, """Engineered plastics form the primary housing material. ABS (Acrylonitrile Butadiene Styrene) comprises 50-60 percent of budget and mid-range housings with heat deflection temperature of 98-105 degrees Celsius, impact resistance, and ease of injection moulding. Global ABS pricing averaged USD 1.40-1.80 per kilogram in 2024, with Asian spot prices typically 8-12 percent below Western markets. A typical straightener housing requires 120-180 grams of ABS, translating to material costs of USD 0.17-0.32 per unit. Polycarbonate is used in premium models due to superior heat resistance of 140-150 degrees Celsius and optical clarity for temperature displays, with pricing ranging USD 2.20-3.00 per kilogram.""")

    add_body_text(doc, """Metal inputs serve structural and thermal conduction functions. Aluminium alloys are used for heating plates in budget models and heat sinks, with London Metal Exchange prices averaging USD 2,200-2,600 per metric ton in 2024. A typical straightener contains 40-80 grams of aluminium, contributing USD 0.09-0.21 to material costs. Titanium is used in premium straighteners for superior heat conductivity and corrosion resistance, with commercial pricing ranging USD 15,000-25,000 per metric ton, approximately 8-10 times aluminium prices, restricting its use to products retailing above USD 100.""")

    add_body_text(doc, """Ceramic heating elements are the core functional component, either procured from domestic manufacturers or imported from countries such as China. Advanced PTC ceramic heaters, which utilize semiconductor properties for self-regulating heating, require specialized manufacturing with barium titanate ceramics doped with rare-earth elements. Manufacturing tolerances for PTC elements are exceptionally tight with plus or minus 3 degrees Celsius temperature accuracy, requiring sophisticated quality control systems that increase per-unit costs by 15-25 percent relative to standard resistance heating.""")

    add_body_text(doc, "Table 2.2: Raw Materials and Procurement Sources")
    materials_headers = ["Input Material", "Source Region", "Cost per Unit", "Nature of Cost"]
    materials_rows = [
        ["Ceramic heating elements", "China / Domestic", "Rs. 150-300", "Variable"],
        ["ABS plastic housing", "Domestic manufacturers", "Rs. 80-120", "Variable"],
        ["Electrical wiring", "India / East Asia", "Rs. 50-80", "Variable"],
        ["Temperature control circuits", "China / Taiwan", "Rs. 120-200", "Variable"],
        ["Bristles (nylon/ceramic)", "China / Domestic", "Rs. 40-80", "Variable"],
        ["Power cord and plug", "Domestic", "Rs. 60-100", "Variable"],
        ["Packaging materials", "Domestic", "Rs. 30-50", "Variable"],
    ]
    add_table(doc, materials_headers, materials_rows)

    add_body_text(doc, """Labour is hired locally with wage rates depending on skill levels and labour market conditions. The factor market for hair straightener production demonstrates characteristics of competitive input markets with multiple suppliers. Firms attempt to procure inputs at the lowest possible cost while maintaining quality standards, reflecting the microeconomic principle of cost minimization where producers choose the most efficient combination of factors of production.""")

    add_subheading(doc, "2.3 Cost Structure: Detailed Microeconomic Analysis")

    add_body_text(doc, """The cost structure of producing hair straightening brushes consists of operating expenditure (Op-Ex) and capital expenditure (Cap-Ex). Understanding this cost structure is fundamental to analyzing firm behaviour, pricing decisions, and market competition. Capital expenditure includes investments in manufacturing plants, machinery, moulds, research and development facilities, and regulatory certifications. These costs are incurred upfront and represent fixed costs that do not vary with output in the short run.""")

    add_body_text(doc, """Operating expenditure comprises variable costs that change with production volume, including raw materials, labour wages, electricity consumption, quality control testing, packaging materials, transportation, and logistics. The relationship between fixed and variable costs determines the firm's cost behaviour and optimal production decisions.""")

    add_body_text(doc, "Table 2.3: Detailed Cost Structure per Unit")
    cost_headers = ["Cost Component", "Cost Range (Rs.)", "Percentage of Total", "Cost Type"]
    cost_rows = [
        ["Raw materials", "900-1,200", "45-50%", "Variable"],
        ["Labour and wages", "200-300", "10-12%", "Variable"],
        ["Electricity and utilities", "80-120", "4-5%", "Variable"],
        ["Quality control testing", "50-80", "2-3%", "Variable"],
        ["Packaging materials", "100-150", "5-6%", "Variable"],
        ["Transportation and logistics", "150-250", "8-10%", "Variable"],
        ["Depreciation and overheads", "200-300", "10-12%", "Fixed"],
        ["Administrative costs", "100-150", "5-6%", "Fixed"],
        ["Total Production Cost", "1,800-2,550", "100%", "-"],
    ]
    add_table(doc, cost_headers, cost_rows)

    add_body_text(doc, """Raw materials constitute the largest share of total cost at approximately 45-50 percent, making production highly sensitive to fluctuations in input prices. Fixed costs including machinery depreciation and administrative overheads remain constant in the short run, while variable costs increase proportionally with output. The average total cost of producing one unit ranges between Rs. 1,800 and Rs. 2,550, depending on scale of production and input procurement efficiency.""")

    add_subheading(doc, "2.3.1 Short-Run Cost Analysis")

    add_body_text(doc, """In the short run, at least one factor of production remains fixed, typically capital equipment and factory space. The short-run cost structure exhibits characteristic patterns that inform production decisions. Total Fixed Cost (TFC) remains constant regardless of output level, including rent, machinery depreciation, and administrative salaries. Total Variable Cost (TVC) increases with output, encompassing raw materials, direct labour, and utilities. Total Cost (TC) equals the sum of TFC and TVC.""")

    add_body_text(doc, """Average Fixed Cost (AFC) declines continuously as output increases because the same fixed cost is spread over more units. This phenomenon explains why firms seek to increase production volume to achieve lower per-unit fixed costs. Average Variable Cost (AVC) initially decreases due to specialization and efficiency gains, reaches a minimum at optimal capacity utilization, then increases due to diminishing marginal returns as variable inputs are added to fixed factors.""")

    add_body_text(doc, "Table 2.4: Short-Run Cost Schedule (Illustrative)")
    sr_cost_headers = ["Output (Units)", "TFC (Rs.)", "TVC (Rs.)", "TC (Rs.)", "AFC (Rs.)", "AVC (Rs.)", "ATC (Rs.)", "MC (Rs.)"]
    sr_cost_rows = [
        ["0", "50,000", "0", "50,000", "-", "-", "-", "-"],
        ["100", "50,000", "180,000", "230,000", "500", "1,800", "2,300", "1,800"],
        ["200", "50,000", "340,000", "390,000", "250", "1,700", "1,950", "1,600"],
        ["300", "50,000", "480,000", "530,000", "167", "1,600", "1,767", "1,400"],
        ["400", "50,000", "600,000", "650,000", "125", "1,500", "1,625", "1,200"],
        ["500", "50,000", "750,000", "800,000", "100", "1,500", "1,600", "1,500"],
        ["600", "50,000", "960,000", "1,010,000", "83", "1,600", "1,683", "2,100"],
        ["700", "50,000", "1,260,000", "1,310,000", "71", "1,800", "1,871", "3,000"],
    ]
    add_table(doc, sr_cost_headers, sr_cost_rows)

    add_body_text(doc, """The table illustrates typical short-run cost behaviour. Average Total Cost (ATC) initially declines due to the spreading of fixed costs and efficiency gains, reaches minimum at approximately 500 units representing optimal short-run output, then rises as diminishing returns set in. Marginal Cost (MC) represents the additional cost of producing one more unit and intersects ATC at its minimum point, indicating the most efficient scale of production in the short run.""")

    add_subheading(doc, "2.3.2 Long-Run Cost Analysis and Economies of Scale")

    add_body_text(doc, """In the long run, all factors of production become variable, allowing firms to adjust plant size and capital stock. The Long-Run Average Cost (LRAC) curve represents the lowest average cost achievable at each output level when all inputs can be varied. The LRAC curve is typically U-shaped, reflecting economies and diseconomies of scale.""")

    add_body_text(doc, """Economies of scale reduce per-unit costs as output expands. Technical economies arise from using larger, more efficient machinery and production lines. Automated assembly lines achieving breakeven at production volumes of approximately 150,000-200,000 units annually can reduce labour costs from USD 2.50 per unit to USD 0.40 per unit. Managerial economies result from specialization of management functions and more efficient organizational structures. Marketing economies enable bulk advertising and stronger bargaining power with distributors. Financial economies provide access to lower interest rates and better credit terms for larger firms. Purchasing economies allow bulk procurement at discounted rates.""")

    add_body_text(doc, """Chinese mega-factories producing millions of units annually have automated extensively, achieving per-unit costs 30-40 percent below smaller facilities through spreading fixed costs and exploiting specialization. Individual factories producing 2-5 million units annually achieve significant cost advantages over smaller competitors. This cost structure creates substantial barriers to entry for new firms lacking scale.""")

    add_body_text(doc, "Table 2.5: Economies of Scale in Production")
    scale_headers = ["Production Scale", "Annual Units", "Average Cost (Rs.)", "Cost Advantage"]
    scale_rows = [
        ["Small-scale", "5,000-20,000", "2,400-2,800", "Baseline"],
        ["Medium-scale", "50,000-200,000", "1,900-2,200", "15-20% lower"],
        ["Large-scale", "500,000-2,000,000", "1,500-1,800", "30-35% lower"],
        ["Mass production", ">2,000,000", "1,200-1,500", "45-50% lower"],
    ]
    add_table(doc, scale_headers, scale_rows)

    add_body_text(doc, """The learning curve effect further reduces costs as cumulative production experience improves efficiency. Workers become more skilled at assembly tasks, quality control processes become more refined, and production scheduling becomes more efficient. Studies in electronics manufacturing suggest learning rates of 10-15 percent, meaning costs decline by this percentage each time cumulative output doubles.""")

    add_subheading(doc, "2.4 Supply Chain Analysis")

    add_body_text(doc, """The supply chain for hair straightening brushes encompasses both backend operations connecting raw material suppliers to manufacturers, and frontend operations linking manufacturers to end consumers. Effective supply chain management is crucial for cost minimization and competitive positioning.""")

    add_body_text(doc, """The backend supply chain begins with raw material extraction and processing. Plastic resin producers supply ABS and polycarbonate to injection moulding firms. Metal refiners supply aluminium and copper to component manufacturers. Ceramic manufacturers produce heating elements using specialized firing processes. Electronic component producers supply integrated circuits, thermistors, and other control components. These tier-2 suppliers feed into tier-1 suppliers who manufacture subassemblies including heating element assemblies, electronic control boards, and housing components.""")

    add_body_text(doc, """Final assembly occurs at original equipment manufacturer facilities, integrating tier-1 components into finished products. Approximately 60 percent of hair straightener manufacturers utilize in-house final assembly while outsourcing component production. This configuration balances specialization benefits from sourcing components from efficient specialists with control over quality-critical final assembly. Chinese manufacturers in Shenzhen leverage supplier proximity with average distances of 30-50 kilometres to minimize inventory holding through just-in-time systems, reducing working capital requirements by 25-35 percent versus Western manufacturers.""")

    add_body_text(doc, """The frontend supply chain distributes finished products through multiple channels. The traditional wholesale-retail channel involves manufacturers selling to distributors at ex-factory prices, distributors adding 15-25 percent markup before selling to retailers, and retailers adding 40-60 percent markup before selling to consumers. A product with Rs. 1,800 factory cost thus reaches consumers at Rs. 2,900-3,600 through traditional channels.""")

    add_body_text(doc, "Table 2.6: Supply Chain Value Distribution")
    sc_headers = ["Supply Chain Stage", "Value Added", "Margin Range", "Cumulative Price"]
    sc_rows = [
        ["Manufacturer (ex-factory)", "Production cost + margin", "15-20%", "Rs. 1,800-2,100"],
        ["Distributor/Wholesaler", "Logistics, warehousing", "15-25%", "Rs. 2,100-2,600"],
        ["Retailer", "Display, sales, service", "30-40%", "Rs. 2,700-3,600"],
        ["E-commerce platform", "Digital infrastructure", "15-25%", "Rs. 2,100-2,600"],
    ]
    add_table(doc, sc_headers, sc_rows)

    add_subheading(doc, "2.5 Wholesale Markets in Mumbai")

    add_body_text(doc, """Mumbai serves as a major wholesale hub for consumer electrical appliances in Western India. Hair straightening brushes are commonly available at wholesale rates in several key markets. Crawford Market, located in South Mumbai, functions as a traditional wholesale centre for diverse consumer goods including electrical appliances. Lohar Chawl specializes in electrical and hardware items, offering competitive wholesale prices due to high trader concentration. Lamington Road is known for electronics and electrical goods, with numerous shops offering hair styling appliances at wholesale rates. Bhiwandi, located on Mumbai's outskirts, serves as a major warehousing and distribution hub with lower operating costs enabling competitive wholesale pricing.""")

    add_body_text(doc, """Wholesale prices at these markets are typically 15-25 percent lower than retail prices, allowing retailers to earn margins. Bulk purchasing further reduces per-unit costs, with orders exceeding 100 units typically receiving additional discounts of 5-10 percent. These markets facilitate transactions between manufacturers, wholesalers, and retailers, contributing to efficient price discovery and distribution in the product market.""")

    add_subheading(doc, "2.6 Government Policies and Regulations")

    add_body_text(doc, """Government policies significantly influence the production, pricing, and marketing of hair straightening brushes in India. Understanding these regulatory frameworks is essential for comprehensive economic analysis of the product market.""")

    add_body_text(doc, """The Goods and Services Tax (GST) structure applies an 18 percent rate to hair straightening brushes, classified under electrical appliances. This tax is levied at each stage of the value chain with input tax credit available, reducing cascading effects compared to the previous tax regime. Import duties on finished products and electronic components affect cost structures for manufacturers relying on imported inputs. Basic customs duty of 20 percent applies to imported hair styling appliances, while electronic components attract duties ranging from 5-15 percent depending on classification.""")

    add_body_text(doc, """Bureau of Indian Standards (BIS) certification is mandatory for electrical appliances sold in India, ensuring products meet safety and performance standards. The certification process involves product testing, factory audits, and ongoing compliance monitoring. Compliance costs include initial certification fees of Rs. 50,000-100,000, annual licence fees, and testing costs for each product variant. While increasing production costs, BIS certification serves as a quality signal reducing consumer uncertainty and enabling market differentiation.""")

    add_body_text(doc, """Consumer protection laws including the Consumer Protection Act 2019 mandate clear labelling, warranty provisions, and redressal mechanisms. Manufacturers must provide minimum warranty periods and establish service networks for after-sales support. Environmental regulations including e-waste management rules require manufacturers to establish collection and recycling mechanisms for end-of-life products, adding compliance costs but promoting sustainable production practices.""")

    add_body_text(doc, "Table 2.7: Regulatory Framework Summary")
    reg_headers = ["Regulation", "Applicable Rate/Requirement", "Impact on Costs"]
    reg_rows = [
        ["GST", "18%", "Adds to final price"],
        ["Import duty (finished goods)", "20%", "Increases import costs"],
        ["Import duty (components)", "5-15%", "Affects raw material costs"],
        ["BIS certification", "Mandatory", "Rs. 50,000-100,000 initial"],
        ["Consumer protection", "Warranty required", "After-sales service costs"],
        ["E-waste compliance", "EPR registration", "Collection/recycling costs"],
    ]
    add_table(doc, reg_headers, reg_rows)

    # ============================================
    # SECTION 3: FRONTEND ANALYSIS
    # ============================================
    add_heading(doc, "3. FRONTEND ANALYSIS: DEMAND, REVENUE, AND MARKET STRUCTURE")

    add_subheading(doc, "3.1 Nature and Determinants of Demand")

    add_body_text(doc, """Demand for hair straightening brushes reflects derived demand from consumers' desires for specific hairstyles and appearance attributes rather than intrinsic product demand. The product is primarily demanded by urban and semi-urban consumers, particularly students, working professionals, and young adults who value convenience and time-saving grooming solutions. Since hair straightening brushes are not essential commodities, their demand depends significantly on income levels, tastes, and preferences.""")

    add_body_text(doc, """Consumers derive utility from hair straightening brushes through multiple attributes including styling outcome, time savings, ease of use, safety features, and price. This multi-attribute utility framework explains product differentiation success, as manufacturers segment markets by emphasizing different attributes. Premium brands focus on styling quality and durability while budget brands emphasize affordability and basic functionality.""")

    add_body_text(doc, """Income effects represent a primary demand determinant. Hair straighteners exhibit characteristics of normal goods with positive income elasticity. As consumer incomes rise, demand increases more than proportionally in emerging markets where ownership rates remain low. In India, middle-class expansion with 8-10 percent annual income growth during 2020-2024 correlates with 7.1 percent annual demand growth for the product category. This pattern suggests luxury good characteristics in low-income contexts transitioning toward necessity-good status at higher income levels.""")

    add_body_text(doc, """Demographic factors significantly influence demand patterns. The 18-35 age cohort accounts for approximately 68 percent of purchases, demonstrating higher frequency usage of 4-5 times weekly compared to 2-3 times for consumers above 35 years. While historically female-dominated with 92 percent of users in 2010, male adoption increased to approximately 20 percent by 2024, reflecting evolving grooming norms. Urban residents demonstrate 3-4 times higher ownership rates than rural populations due to income differentials, social media exposure, and reliable electricity access.""")

    add_body_text(doc, "Table 3.1: Demand Determinants and Effects")
    demand_headers = ["Demand Determinant", "Effect on Demand", "Economic Interpretation"]
    demand_rows = [
        ["Increase in income", "Demand increases", "Positive income elasticity"],
        ["Availability of substitutes", "Demand sensitive to price", "High cross elasticity"],
        ["Fashion and grooming trends", "Demand increases", "Taste and preference shift"],
        ["Urban lifestyle adoption", "Higher demand", "Convenience-driven demand"],
        ["Seasonal occasions", "Temporary demand spike", "Rightward demand shift"],
        ["Social media influence", "Demand increases", "Information effect"],
        ["Working women population", "Sustained demand growth", "Time opportunity cost"],
    ]
    add_table(doc, demand_headers, demand_rows)

    add_subheading(doc, "3.2 Elasticity of Demand Analysis")

    add_body_text(doc, """Elasticity of demand measures the responsiveness of quantity demanded to changes in various factors. Understanding elasticity is crucial for pricing decisions, revenue optimization, and market strategy formulation. This analysis examines price elasticity, income elasticity, and cross elasticity for hair straightening brushes.""")

    add_body_text(doc, """Price Elasticity of Demand (PED) measures the percentage change in quantity demanded resulting from a one percent change in price. For hair straightening brushes, price elasticity varies significantly across market segments and price ranges. The budget segment below Rs. 2,500 exhibits highly elastic demand with PED estimated at -1.8 to -2.5. A 10 percent price increase would reduce quantity demanded by 18-25 percent. This high elasticity reflects multiple substitutes within the budget tier, income constraints among target consumers, and low brand loyalty enabling easy switching.""")

    add_body_text(doc, """The mid-range segment from Rs. 2,500 to Rs. 8,000 demonstrates moderately elastic demand with PED of -1.0 to -1.4. Brand differentiation by established players like Philips, Vega, and Havells reduces perfect substitutability, while quality perception creates minimum threshold effects limiting downward substitution. The premium segment above Rs. 8,000 shows inelastic demand with PED of -0.4 to -0.7. A 10 percent price increase reduces quantity demanded by only 4-7 percent. Premium consumers prioritize quality and brand over price, limited substitutes exist at equivalent quality levels, and high incomes reduce price sensitivity.""")

    add_body_text(doc, "Table 3.2: Price Elasticity by Market Segment")
    ped_headers = ["Market Segment", "Price Range (Rs.)", "PED Estimate", "Demand Nature", "Revenue Implication"]
    ped_rows = [
        ["Budget", "<2,500", "-1.8 to -2.5", "Highly elastic", "Price cuts increase TR"],
        ["Mid-range", "2,500-8,000", "-1.0 to -1.4", "Unit elastic", "TR stable with price"],
        ["Premium", ">8,000", "-0.4 to -0.7", "Inelastic", "Price hikes increase TR"],
    ]
    add_table(doc, ped_headers, ped_rows)

    add_body_text(doc, """Income Elasticity of Demand (YED) measures demand responsiveness to income changes. Hair straightening brushes exhibit positive income elasticity, classifying them as normal goods. In emerging markets like India with lower average incomes, estimated YED ranges from 1.5 to 2.2, indicating luxury good characteristics where demand grows faster than income. In developed markets with higher saturation, YED falls to 0.8-1.2 as the product approaches necessity status. This transition explains faster market growth in developing economies where income growth translates into proportionally larger demand increases.""")

    add_body_text(doc, """Cross Elasticity of Demand (XED) measures demand responsiveness to price changes of related goods. Cross elasticity between straightening brushes and flat irons is estimated at +0.4 to +0.6, indicating substitute relationships. A 10 percent flat iron price increase would boost straightening brush demand by 4-6 percent. Cross elasticity with salon straightening services shows higher values of +0.6 to +0.8, as salon visits represent a more direct substitute for home styling. Cross elasticity with complementary products like heat protectant sprays and styling products shows negative values of -0.2 to -0.3, where price increases in complements reduce straightening brush demand.""")

    add_body_text(doc, "Table 3.3: Elasticity Summary")
    elasticity_headers = ["Elasticity Type", "Estimate", "Interpretation", "Strategic Implication"]
    elasticity_rows = [
        ["Price elasticity (budget)", "-2.0", "Highly elastic", "Volume-based strategy"],
        ["Price elasticity (premium)", "-0.5", "Inelastic", "Premium pricing viable"],
        ["Income elasticity (India)", "+1.8", "Luxury good", "Target growing middle class"],
        ["Cross elasticity (flat irons)", "+0.5", "Substitutes", "Competitive positioning"],
        ["Cross elasticity (salons)", "+0.7", "Substitutes", "Home convenience emphasis"],
    ]
    add_table(doc, elasticity_headers, elasticity_rows)

    add_subheading(doc, "3.3 Market Structure Analysis")

    add_body_text(doc, """The hair straightening brush market operates under monopolistic competition, characterized by many sellers offering differentiated products with relatively free entry and exit. Each firm attempts to distinguish its product through design, technology, safety features, branding, and advertising. Due to the presence of many sellers and close substitutes, individual firms have limited control over market prices.""")

    add_body_text(doc, """Product differentiation occurs along multiple dimensions. Functional differentiation includes heating technology (ceramic, ionic, tourmaline), temperature range and control precision, heating speed, and additional features like auto shut-off. Design differentiation encompasses ergonomic handle shapes, aesthetic styling, size variations for portability, and cord length options. Brand differentiation creates perceived quality differences through reputation, celebrity endorsements, and marketing communications.""")

    add_body_text(doc, """Market fragmentation is evident with numerous competitors including global brands like Philips, Dyson, and Remington, regional players like Vega and Havells in India, and numerous unbranded manufacturers producing budget alternatives. The Herfindahl-Hirschman Index (HHI) for the Indian market is estimated below 1,500, indicating moderate concentration consistent with monopolistic competition. No single firm controls more than 15-20 percent market share.""")

    add_body_text(doc, """Entry barriers are moderate. Technical barriers include product development costs, manufacturing equipment investment, and quality certification requirements. Brand barriers require substantial marketing investment to establish recognition in a crowded market. Distribution barriers necessitate relationships with retail chains and e-commerce platforms. However, contract manufacturing options and e-commerce platforms have reduced entry barriers, enabling new entrants to compete without massive capital investments.""")

    add_body_text(doc, """In the short run, firms may earn supernormal profits through successful product differentiation or cost advantages. In the long run, the absence of significant entry barriers attracts new competitors, eroding excess profits until firms earn normal profits where price equals average cost. This competitive pressure encourages continuous innovation and efficiency improvements to maintain market position.""")

    add_subheading(doc, "3.4 Pricing Mechanism")

    add_body_text(doc, """Pricing in the hair straightening brush market reflects various strategies depending on market segment, competitive positioning, and distribution channel. Understanding pricing mechanisms is essential for analyzing revenue generation and market dynamics.""")

    add_body_text(doc, """Cost-plus pricing forms the foundation for manufacturer pricing decisions. The cost of production forms the base and a fixed margin is added, ensuring cost recovery and reasonable profit. For a product with Rs. 2,000 production cost, applying a 20 percent margin yields a manufacturer selling price of Rs. 2,400. This approach provides pricing stability and ensures profitability but may not optimize revenue in all market conditions.""")

    add_body_text(doc, """Value-based pricing is employed by premium brands where price reflects perceived consumer value rather than production cost. Dyson's hair styling products command significant premiums (Rs. 30,000-45,000) based on technological innovation, design excellence, and brand prestige rather than proportionally higher production costs. This strategy is viable when brand differentiation creates inelastic demand segments.""")

    add_body_text(doc, """Wholesale versus retail pricing reflects channel economics. Manufacturers sell to distributors at ex-factory prices plus margin. Distributors add 15-25 percent for logistics and warehousing services. Retailers add 30-40 percent for display, sales support, and service, yielding final retail prices 60-100 percent above manufacturer prices. E-commerce channels compress these margins, enabling 10-20 percent lower consumer prices while maintaining seller profitability.""")

    add_body_text(doc, """Dynamic pricing is increasingly common on e-commerce platforms. Algorithms adjust prices based on demand patterns, competitor pricing, inventory levels, and consumer browsing behaviour. Prices may vary 15-25 percent across platforms and time periods. Promotional pricing during festival seasons and sale events offers 20-40 percent discounts, stimulating demand during peak shopping periods.""")

    add_subheading(doc, "3.5 Revenue Analysis")

    add_body_text(doc, """Revenue analysis examines the earnings generated at various stages of the product value chain. Understanding revenue relationships helps explain firm behaviour and market outcomes.""")

    add_body_text(doc, """Total Revenue (TR) equals price multiplied by quantity sold. For a firm selling 10,000 units at Rs. 2,500 average price, TR equals Rs. 2.5 crore. Average Revenue (AR) equals TR divided by quantity, representing the price received per unit. In competitive markets, AR equals the market price. Marginal Revenue (MR) represents additional revenue from selling one more unit. Under monopolistic competition, MR lies below AR due to the downward-sloping demand curve.""")

    add_body_text(doc, """Revenue distribution across the value chain shows manufacturers earning 25-35 percent of final retail price, distributors earning 8-12 percent, and retailers earning 15-25 percent. E-commerce platforms capture 12-18 percent through commissions, fulfillment fees, and advertising charges. These margins reflect the value added at each stage and the bargaining power of participants.""")

    add_body_text(doc, "Table 3.4: Revenue Distribution Example (Rs. 3,000 Retail Price)")
    revenue_headers = ["Value Chain Stage", "Revenue Share", "Amount (Rs.)", "Value Added"]
    revenue_rows = [
        ["Manufacturer", "55-60%", "1,650-1,800", "Production, R&D, quality"],
        ["Distributor", "12-15%", "360-450", "Logistics, warehousing"],
        ["Retailer", "25-30%", "750-900", "Display, sales, service"],
        ["Total Retail Price", "100%", "3,000", "-"],
    ]
    add_table(doc, revenue_headers, revenue_rows)

    add_body_text(doc, """The relationship between elasticity and revenue is crucial for pricing decisions. When demand is elastic (PED > 1), price reductions increase total revenue as quantity gains exceed percentage price drops. When demand is inelastic (PED < 1), price increases raise total revenue as quantity losses are proportionally smaller. This relationship explains why budget brands focus on volume strategies while premium brands can maintain higher prices.""")

    add_subheading(doc, "3.6 Selling and Marketing Costs")

    add_body_text(doc, """Selling costs play an important role in influencing demand for hair straightening brushes. Unlike production costs, selling costs do not directly increase output but aim to shift the demand curve rightward by influencing consumer preferences. In monopolistically competitive markets, selling costs are essential for product differentiation and brand building.""")

    add_body_text(doc, """Advertising expenditure encompasses television commercials, digital marketing campaigns, print advertisements, and outdoor media. Large brands allocate 8-15 percent of revenue to advertising, with digital channels increasingly dominant. Social media advertising and influencer partnerships provide targeted reach to the 18-35 demographic at lower cost per impression than traditional media.""")

    add_body_text(doc, """Influencer marketing has become particularly significant for personal grooming products. Beauty influencers demonstrating products to followers create authentic engagement and purchase intent. Micro-influencers with 10,000-100,000 followers often provide better engagement rates than celebrity endorsements at fraction of the cost. Brands allocate 3-8 percent of marketing budgets to influencer collaborations.""")

    add_body_text(doc, """Packaging and design investments serve dual purposes of product protection and brand communication. Premium packaging with quality materials and attractive design signals product quality and justifies price premiums. Packaging costs range from Rs. 30 for basic boxes to Rs. 150 for premium presentation, representing 2-5 percent of retail price.""")

    add_body_text(doc, """After-sales service costs include warranty fulfilment, customer support infrastructure, and service centre operations. Manufacturers typically allocate 2-4 percent of revenue to after-sales service. Strong service reputation builds brand loyalty and enables premium positioning, particularly important for durable goods with multi-year usage periods.""")

    add_body_text(doc, "Table 3.5: Selling Costs Breakdown")
    selling_headers = ["Selling Cost Type", "Percentage of Revenue", "Purpose", "Demand Effect"]
    selling_rows = [
        ["Advertising", "8-15%", "Brand awareness", "Shifts demand right"],
        ["Influencer marketing", "3-8%", "Product credibility", "Increases consideration"],
        ["Discounts and promotions", "5-10%", "Volume stimulation", "Movement along curve"],
        ["Packaging", "2-5%", "Brand perception", "Premium positioning"],
        ["After-sales service", "2-4%", "Customer retention", "Repeat purchase"],
    ]
    add_table(doc, selling_headers, selling_rows)

    # ============================================
    # SECTION 4: SUGGESTIONS FOR IMPROVEMENT
    # ============================================
    add_heading(doc, "4. SUGGESTIONS TO IMPROVE PRODUCTION AND DELIVERY EFFICIENCY")

    add_body_text(doc, """Based on the comprehensive analysis of production, supply chain, and market dynamics, several recommendations emerge for improving the efficiency of hair straightening brush production and delivery. These suggestions aim to reduce costs, enhance quality, and improve market responsiveness.""")

    add_body_text(doc, """First, manufacturers should increase automation in the production process to reduce dependence on manual labour, improve consistency in product quality, and lower average cost through economies of scale. Investment in automated assembly lines becomes economically viable at production volumes exceeding 150,000 units annually. Automated processes reduce per-unit labour costs by 70-80 percent while improving quality consistency through precise component placement and standardized processes.""")

    add_body_text(doc, """Second, bulk procurement of raw materials and establishment of long-term contracts with suppliers should be adopted to reduce input price fluctuations and achieve lower per-unit raw material costs. Negotiating annual supply agreements with key component suppliers can secure 5-10 percent price reductions compared to spot market purchasing. Building strategic partnerships with reliable suppliers also ensures consistent quality and timely delivery.""")

    add_body_text(doc, """Third, the supply chain should be streamlined by reducing unnecessary intermediaries, which would lower transportation, storage, and inventory holding costs, leading to more competitive pricing. Direct-to-retailer distribution models and e-commerce partnerships can eliminate distributor margins of 15-25 percent. Investment in warehouse management systems and logistics optimization can reduce inventory holding costs by 20-30 percent.""")

    add_body_text(doc, """Fourth, improved demand forecasting based on past sales trends, seasonal demand patterns, and market intelligence should be implemented to prevent overproduction or underproduction, thereby reducing wastage and excess inventory costs. Advanced analytics and machine learning algorithms can improve forecast accuracy by 15-25 percent compared to traditional methods.""")

    add_body_text(doc, """Fifth, investment in efficient logistics and faster distribution networks can reduce delivery time, lower distribution costs, and improve product availability. Partnerships with third-party logistics providers and strategic warehouse placement can reduce average delivery time from 5-7 days to 2-3 days while lowering per-unit logistics costs.""")

    add_body_text(doc, """Sixth, adoption of sustainable materials and energy-efficient production processes addresses growing consumer preference for environmentally responsible products while potentially reducing material costs. Recycled plastics can reduce material costs by 10-15 percent while enhancing brand image. Energy-efficient machinery reduces utility costs and supports environmental compliance.""")

    add_body_text(doc, """Seventh, investment in quality management systems beyond minimum compliance requirements can reduce warranty claims and returns, improving customer satisfaction and brand reputation. Six Sigma quality processes can reduce defect rates from 3-5 percent to below 1 percent, substantially reducing after-sales costs and improving customer retention.""")

    # ============================================
    # SECTION 5: LESSONS LEARNED
    # ============================================
    add_heading(doc, "5. LESSONS LEARNED")

    add_body_text(doc, """Through the comprehensive study of hair straightening brushes as a microeconomic product, I have gained valuable insights into the practical application of economic theory to real-world markets. The following points summarize my key learnings from this assignment.""")

    add_body_text(doc, """First, I learned how production decisions are fundamentally influenced by the cost structure of a firm, particularly the distinction between fixed costs and variable costs. Understanding that fixed costs like machinery and factory rent remain constant regardless of output while variable costs like raw materials increase with production volume helped me appreciate why firms seek economies of scale. The analysis revealed how average cost initially declines as fixed costs are spread over larger output but eventually rises due to diminishing marginal returns. This understanding of cost behaviour provides crucial insight into firm decision-making regarding optimal production levels and capacity utilization.""")

    add_body_text(doc, """Second, the study highlighted the critical importance of price elasticity of demand in determining pricing strategies. I learned that when demand is elastic, as observed in the budget segment of hair straightening brushes, firms must focus on increasing sales volume rather than raising prices because price cuts increase total revenue. Conversely, premium segments with inelastic demand can sustain higher prices without significant volume loss. This differentiated elasticity across market segments explains why firms employ segment-specific pricing strategies and why premium brands can maintain substantial price premiums over budget alternatives.""")

    add_body_text(doc, """Third, the analysis demonstrated how efficiency in the supply chain directly affects final market prices and profitability. I understood the economic rationale for intermediaries including wholesalers and retailers who add value through inventory holding, assortment creation, and geographic distribution despite adding to final prices. The emergence of e-commerce as a more efficient distribution channel illustrated how technological innovation can reduce transaction costs and create consumer welfare gains. The shift to online channels capturing 58 percent of Indian market sales demonstrates market responses to cost efficiency opportunities.""")

    add_body_text(doc, """Fourth, the assignment provided comprehensive insight into the functioning of monopolistic competition, where product differentiation and selling costs play significant roles in influencing consumer demand. I learned how firms under monopolistic competition compete not only on price but also on product features, branding, and marketing. The presence of many sellers with differentiated products means no single firm controls market prices, yet differentiation creates limited market power. Understanding this market structure explains the prevalence of advertising and brand building in consumer goods markets.""")

    add_body_text(doc, """Fifth, I gained appreciation for how government policies including taxation, safety standards, and import regulations impact production costs, market prices, and overall consumer welfare. The 18 percent GST, import duties on components, and mandatory BIS certification all contribute to the final product price while serving legitimate policy objectives including revenue generation, domestic manufacturing protection, and consumer safety. This regulatory analysis demonstrates the interconnection between microeconomic outcomes and macroeconomic policy frameworks.""")

    # Save document
    output_path = '/mnt/e/AI and Projects/MMS-Prep/Eco Assign/shriya/Microeconomics_Hair_Straightening_Brush_Assignment.docx'
    doc.save(output_path)
    print(f"Document saved to: {output_path}")

    # Count approximate words
    word_count = 0
    for para in doc.paragraphs:
        word_count += len(para.text.split())
    print(f"Approximate word count: {word_count}")

if __name__ == "__main__":
    main()
