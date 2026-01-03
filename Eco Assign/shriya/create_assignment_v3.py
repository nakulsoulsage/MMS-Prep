#!/usr/bin/env python3
"""
Create Microeconomics Product Assignment on Hair Straightening Brush
Final Version: 16 pages, 3-4 tables, economic diagrams
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_page_margins(doc):
    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1)

def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

def add_heading(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.line_spacing = 1.5

def add_subheading(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.5

def add_body(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(8)

def add_figure_caption(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(10)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for para in header_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data)
            for para in row_cells[col_idx].paragraphs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
    doc.add_paragraph()

def add_diagram(doc, diagram_text, width=6):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.width = Inches(width)
    cell_para = cell.paragraphs[0]
    run = cell_para.add_run(diagram_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    cell_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()

def create_cover_page(doc):
    for _ in range(2):
        doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run("MICROECONOMICS PRODUCT ASSIGNMENT")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run("HAIR STRAIGHTENING BRUSH")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(20)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run("A Microeconomic Analysis of Production, Cost, Demand, and Market Dynamics")
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(4):
        doc.add_paragraph()
    details = [
        ("Name:", "Shriya Parab"),
        ("Roll No:", "HBM25120"),
        ("Class:", "MMS Batch 2025-27"),
        ("Semester:", "01st"),
        ("Academic Year:", "2025-2026"),
        ("Institute:", "Sydenham Institute of Management Studies,"),
        ("", "Research and Entrepreneurship Education, Mumbai"),
    ]
    for label, value in details:
        para = doc.add_paragraph()
        if label:
            run1 = para.add_run(f"{label}  ")
            run1.bold = True
            run1.font.name = 'Times New Roman'
            run1.font.size = Pt(12)
        run2 = para.add_run(value)
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

def main():
    doc = Document()
    set_page_margins(doc)
    add_page_numbers(doc)
    create_cover_page(doc)

    # ========== SECTION 1: INTRODUCTION ==========
    add_heading(doc, "1. INTRODUCTION AND EVOLUTION")

    add_subheading(doc, "1.1 Definition of Hair Straightening Brush as an Economic Product")

    add_body(doc, """A hair straightening brush is an electrical personal grooming appliance that combines the functionality of a conventional hairbrush with the heating mechanism of an electric hair straightener. The product is designed to straighten hair in a convenient, time-saving, and relatively safe manner, making it suitable for regular home use. Unlike traditional flat irons, which require careful handling and a certain level of skill, hair straightening brushes are comparatively easier to use and reduce the risk of burns or excessive hair damage. As a result, they are widely preferred by non-professional consumers seeking efficient grooming solutions.""")

    add_body(doc, """From a microeconomic perspective, this product represents a convergence of thermal engineering and consumer technology, occupying a distinct market position within the broader USD 30 billion global hair styling tools industry. The hair straightening brush addresses consumer demand for time-efficient, damage-minimizing hair care solutions, reflecting the economic principle of product innovation responding to evolving consumer preferences and technological capabilities. The global market for hair straightening brushes reached USD 620-750 million in 2024, with projections indicating growth to USD 1.53 billion by 2033 at a compound annual growth rate (CAGR) of 8.2 percent.""")

    add_body(doc, """In India specifically, the market stood at USD 22.17 million in 2024, expected to reach USD 33.32 million by 2033 with a CAGR of 4.28 percent, reflecting the product's penetration into emerging markets with expanding middle-class populations. Hair straightening brushes are generally available in the Indian market in the price range of Rs. 1,500 to Rs. 3,500, depending on factors such as brand reputation, quality of materials used, technological features, and safety mechanisms. Premium brands may charge higher prices due to better build quality, advanced temperature control, and stronger brand perception. The product has gained popularity mainly in urban and semi-urban regions among students, working professionals, and young consumers who seek efficient grooming solutions within limited time.""")

    add_subheading(doc, "1.2 Historical Evolution of Hair Styling Technology")

    add_body(doc, """The trajectory of hair straightening technology exemplifies the economic principle of creative destruction, where successive innovations displace existing methods through superior cost-benefit propositions. Historically, hair straightening was carried out using manual methods such as heated metal rods or chemical straightening treatments. These methods were mostly salon-based, expensive, and time-consuming. Chemical straightening, in particular, involved long treatment hours and often caused long-term damage to hair, illustrating an early market failure due to inadequate safety features and high transaction costs.""")

    add_body(doc, """The commercialization of thermal hair straightening began in the 1870s when French hairstylist Marcel Grateau developed heated iron rods paired with chemical lotions. Despite achieving straightening effects, the combination frequently caused scalp burns and hair damage. The breakthrough came in 1909 when Isaac K. Shero patented metal plates capable of gliding through hair, incorporating adjustable heat settings that would become standard in modern devices. This innovation reduced transaction costs associated with professional salon visits by enabling home usage and introduced the concept of consumer control over styling processes.""")

    add_body(doc, """The twentieth century witnessed exponential technological advancement driven by electrification and materials science. The marcel iron's introduction of heat regulation mechanisms represented a significant improvement in user control, reducing information asymmetry between manufacturers and consumers regarding safe operating temperatures. The adoption of plastic materials through injection moulding in the mid-1900s democratized access by dramatically lowering production costs, a textbook example of how manufacturing process innovations can expand market size through reduced average costs.""")

    add_body(doc, """Contemporary hair straightening brushes emerged in the early 2000s as manufacturers recognized an underserved market segment: consumers seeking straightening results with minimal technical skill and time investment. By integrating heating elements into brush formats familiar from daily grooming routines, designers reduced the learning curve and switching costs from traditional methods, thereby expanding the potential consumer base beyond skilled users to the mass market.""")

    add_subheading(doc, "1.3 Technological Innovations in Hair Straightening Brushes")

    add_body(doc, """Modern hair straightening brushes incorporate several technological innovations that create differentiated value propositions and enable premium pricing strategies. Ceramic heating technology, formed by pressing clay mixtures and firing at high temperatures, provides even heat distribution that minimizes hot spots capable of causing thermal damage. This addresses a principal-agent problem where consumers cannot perfectly observe product quality ex-ante, making materials certification a credible quality signal. Ceramic's superior heat conductivity relative to aluminium reduces energy requirements by 15-20 percent, translating to lower operating costs over the product lifecycle.""")

    add_body(doc, """Ionic technology represents another significant advancement in the product category. Devices equipped with ionic generators emit negatively charged ions that break down water molecules on hair surfaces, reducing drying time and static electricity. From an economic standpoint, ionic technology represents a form of process innovation that increases productivity (styled hair per unit time) without proportionally increasing inputs. Premium models leveraging this technology command price premiums of 40-60 percent over basic alternatives, demonstrating how technological differentiation enables market segmentation.""")

    add_body(doc, """PTC (Positive Temperature Coefficient) heating elements utilize semiconductor properties that cause electrical resistance to increase sharply at predetermined temperatures, typically 120 degrees Celsius, creating self-regulating heating without external thermostats. This innovation addresses safety concerns and reduces liability risks for manufacturers, effectively lowering expected costs associated with product recalls or litigation. Tourmaline and titanium coatings further enhance product performance. Tourmaline, a semi-precious mineral, when crushed and infused into ceramic plates, generates negative ions and far-infrared heat that penetrate hair shafts more effectively. Titanium plates offer superior heat conductivity and durability, justifying price premiums of 50-80 percent in professional-grade segments.""")

    add_subheading(doc, "1.4 Product Variants and Comparison with Substitutes")

    add_body(doc, """Hair straightening brushes are available in multiple variants designed to address different consumer needs and price sensitivities. Corded variants represent the most common type, offering consistent power supply and typically lower prices ranging from Rs. 1,200 to Rs. 2,500. Cordless variants provide enhanced portability and convenience but command price premiums of 40-60 percent due to battery technology costs. Material-based differentiation includes ceramic-coated models emphasizing gentle heat distribution, ionic technology variants reducing frizz, and tourmaline-infused options for professional-grade results.""")

    add_body(doc, """Hair straightening brushes compete with several substitute products in the personal grooming market. Traditional flat irons represent the closest substitute, offering precise straightening but requiring careful handling and higher skill levels. Blow dryers with straightening attachments provide an alternative for consumers seeking multi-functional devices. Professional salon services remain a substitute, particularly for consumers seeking expert styling or chemical straightening treatments, though these options involve higher costs and time investments. The straightening brush offers advantages in ease of use, reduced heat damage risk, and time efficiency, making it attractive for convenience-seeking consumers.""")

    add_subheading(doc, "1.5 Product Lifecycle Classification: Sunrise Product")

    add_body(doc, """Applying the standard product lifecycle framework comprising introduction, growth, maturity, and decline phases, hair straightening brushes are classified as a Sunrise product currently in the growth-to-early-maturity transition phase. Several indicators support this classification. Market penetration data shows that 45 percent of women in developed markets owned hair straighteners by 2024, up from 35 percent in 2007, with straightening brushes representing approximately 30 percent of this category. The product has moved beyond early adopters into the mainstream majority, though significant untapped potential remains in emerging markets where ownership rates remain below 15 percent.""")

    add_body(doc, """The sustained CAGR of 5.2-8.2 percent across major markets indicates continued expansion rather than the sales plateauing characteristic of mature products. The introduction of cordless variants and smart-enabled models with app connectivity demonstrates ongoing innovation typical of growth-stage products. Market fragmentation with numerous competitors across price tiers indicates that dominant designs have not fully crystallized, consistent with late-growth phase competitive intensity. Unlike mature commodity markets characterized by price compression, hair straightening brushes maintain significant price dispersion from budget models at Rs. 300-1,200 to premium options at Rs. 15,000-45,000, suggesting that product differentiation remains viable and market segmentation persists.""")

    # ========== SECTION 2: BACKEND ANALYSIS ==========
    add_heading(doc, "2. BACKEND ANALYSIS")

    add_subheading(doc, "2.1 Production Modality and Technology")

    add_body(doc, """Hair straightening brushes are produced using electro-mechanical manufacturing processes. The production process involves assembling ceramic heating elements, electrical wiring, temperature control circuits, and plastic casings manufactured through injection moulding. Once assembled, each unit undergoes electrical safety testing, heat resistance testing, and quality inspection before being packaged for sale. Product development begins with computer-aided design (CAD) modelling to optimize heat distribution patterns, bristle configurations, and ergonomic handling characteristics. A typical product development cycle spans 12-18 months with expenditures ranging from USD 50,000 for budget models to over USD 2 million for premium innovations incorporating novel heating technologies.""")

    add_body(doc, """Small and medium-scale manufacturers usually operate semi-automatic assembly lines producing approximately 500 to 2,000 units per day. Large manufacturers, on the other hand, employ automated production systems capable of producing 5,000 units or more per day. Over time, production has shifted from labour-intensive methods to capital-intensive methods due to technological advancement and the need for cost efficiency. Automation increases output per hour, reduces dependence on manual labour, and ensures uniformity in product quality. In microeconomic terms, technological improvement increases productivity and leads to a downward shift in the average cost curve. This enables firms to produce larger quantities at lower per-unit costs in the long run, helping them achieve economies of scale and remain competitive in the market.""")

    add_subheading(doc, "2.2 Raw Materials and Factor Market Analysis")

    add_body(doc, """The production of hair straightening brushes requires a diverse combination of raw materials, labour, and capital inputs. Ceramic heating elements are either procured from domestic manufacturers or imported from countries such as China. Plastic casings using ABS (Acrylonitrile Butadiene Styrene) with heat deflection temperature of 98-105 degrees Celsius are sourced from domestic plastic manufacturing units. Electrical components such as wiring, thermostats, and control circuits are procured from India and East Asian countries. Labour is hired locally, while machinery and moulds are sourced from domestic and international suppliers.""")

    add_body(doc, """Global ABS pricing averaged USD 1.40-1.80 per kilogram in 2024, with Asian spot prices typically 8-12 percent below Western markets. A typical straightener housing requires 120-180 grams of ABS, translating to material costs of Rs. 15-25 per unit. Aluminium for heating plates and heat sinks costs USD 2,200-2,600 per metric ton, with a typical straightener containing 40-80 grams, contributing Rs. 8-15 to material costs. PTC ceramic heating elements with tight manufacturing tolerances of plus or minus 3 degrees Celsius temperature accuracy increase per-unit costs by 15-25 percent relative to standard resistance heating but provide superior safety and energy efficiency.""")

    add_body(doc, """The factor market for hair straightener production demonstrates characteristics of competitive input markets with multiple suppliers. Firms attempt to procure inputs at the lowest possible cost while maintaining quality standards, reflecting the microeconomic principle of cost minimization where producers choose the most efficient combination of factors of production. Electronic components are often imported due to cost advantages and availability of advanced technology, while plastic casings and labour are sourced domestically to reduce transportation and logistics costs.""")

    add_subheading(doc, "2.3 Cost Structure Analysis")

    add_body(doc, """The cost structure of producing hair straightening brushes consists of operating expenditure (Op-Ex) and capital expenditure (Cap-Ex). Capital expenditure includes investments in manufacturing plants, machinery, moulds, research and development facilities, and regulatory certifications. These costs are incurred upfront and represent fixed costs that do not vary with output in the short run. Operating expenditure comprises variable costs that change with production volume, including raw materials, labour wages, electricity consumption, quality control testing, packaging materials, transportation, and logistics.""")

    # TABLE 1: Cost Structure
    add_body(doc, "Table 2.1: Cost Structure per Unit")
    cost_headers = ["Cost Component", "Cost Range (Rs.)", "Percentage", "Cost Type"]
    cost_rows = [
        ["Raw materials (ceramics, plastics, electronics)", "900-1,200", "45-50%", "Variable"],
        ["Labour and wages", "200-300", "10-12%", "Variable"],
        ["Electricity and utilities", "80-120", "4-5%", "Variable"],
        ["Quality control and testing", "50-80", "3-4%", "Variable"],
        ["Packaging materials", "100-150", "5-6%", "Variable"],
        ["Transportation and logistics", "150-250", "8-10%", "Variable"],
        ["Depreciation and factory overheads", "200-300", "10-12%", "Fixed"],
        ["Total Production Cost per Unit", "1,800-2,400", "100%", "-"],
    ]
    add_table(doc, cost_headers, cost_rows)

    add_body(doc, """Raw materials constitute the largest share of total cost at approximately 45-50 percent, making production highly sensitive to fluctuations in input prices. Fixed costs including machinery depreciation, factory rent, and administrative overheads remain constant in the short run, while variable costs increase proportionally with output. The average total cost of producing one unit ranges between Rs. 1,800 and Rs. 2,400, depending on scale of production, input procurement efficiency, and capacity utilization levels.""")

    add_subheading(doc, "2.4 Short-Run and Long-Run Cost Behaviour")

    add_body(doc, """In the short run, at least one factor of production remains fixed, typically capital equipment and factory space. Total Fixed Cost (TFC) remains constant regardless of output level, including rent, machinery depreciation, insurance, and administrative salaries. Total Variable Cost (TVC) increases with output, encompassing raw materials, direct labour, and utilities. Total Cost (TC) equals the sum of TFC and TVC. Average Fixed Cost (AFC) declines continuously as output increases because the same fixed cost is spread over more units. This phenomenon explains why firms seek to increase production volume to achieve lower per-unit fixed costs.""")

    add_body(doc, """Average Variable Cost (AVC) initially decreases due to specialization and efficiency gains, reaches a minimum at optimal capacity utilization, then increases due to diminishing marginal returns as variable inputs are added to fixed factors. Marginal Cost (MC) represents the additional cost of producing one more unit and intersects Average Total Cost (ATC) at its minimum point, indicating the most efficient scale of production in the short run.""")

    add_figure_caption(doc, "Figure 2.1: Short-Run Cost Curves in Hair Straightening Brush Manufacturing")

    cost_diagram = """
                Cost (Rs.)
                    |
               3000 +                                    MC
                    |                                  /
               2500 +                               /
                    |                            /      ATC
               2000 +         *--------------*---------/
                    |       /                  \\     /
               1500 +     /                      \\ /     AVC
                    |   /           *-------------*-----
               1000 + /           /
                    |           /
                500 +--------/------------------------------------- AFC
                    |
                  0 +--+----+----+----+----+----+----+----+----+-->
                       0   100  200  300  400  500  600  700
                                   Output (Units/day)

    Key Observations:
    - AFC continuously declines (spreading fixed costs)
    - AVC is U-shaped due to diminishing marginal returns
    - MC intersects ATC at minimum ATC (optimal scale: ~450 units)
    - Beyond optimal scale, diseconomies set in
    """
    add_diagram(doc, cost_diagram, 6)

    add_body(doc, """In the long run, all factors of production become variable, allowing firms to adjust plant size and capital stock. The Long-Run Average Cost (LRAC) curve represents the lowest average cost achievable at each output level when all inputs can be varied. The LRAC curve is typically U-shaped, reflecting economies and diseconomies of scale. Economies of scale reduce per-unit costs as output expands through technical economies from larger machinery, managerial economies from specialized functions, marketing economies from bulk advertising, and purchasing economies from bulk procurement. Chinese mega-factories producing millions of units annually achieve per-unit costs 30-40 percent below smaller facilities through spreading fixed costs and exploiting specialization.""")

    add_subheading(doc, "2.5 Supply Chain Analysis")

    add_body(doc, """The supply chain for hair straightening brushes encompasses both backend operations connecting raw material suppliers to manufacturers, and frontend operations linking manufacturers to end consumers. The backend supply chain begins with raw material extraction and processing. Plastic resin producers supply ABS and polycarbonate to injection moulding firms. Metal refiners supply aluminium and copper to component manufacturers. Ceramic manufacturers produce heating elements using specialized firing processes. Electronic component producers supply integrated circuits, thermistors, and control components.""")

    add_figure_caption(doc, "Figure 2.2: Supply Chain Architecture")

    supply_diagram = """
    +------------------------------------------------------------------+
    |                    BACKEND SUPPLY CHAIN                          |
    +------------------------------------------------------------------+
    |                                                                  |
    |  Raw Material Suppliers          Component Manufacturers         |
    |  +------------------+            +----------------------+        |
    |  | Plastic Resins   |----------->| Housing Components   |        |
    |  | (ABS, PC)        |            | (Injection Moulding) |        |
    |  +------------------+            +----------------------+        |
    |                                           |                      |
    |  +------------------+            +----------------------+        |
    |  | Ceramic/Metals   |----------->| Heating Elements     |-----+  |
    |  | (Al, Ti, Clay)   |            | (PTC, Ceramic)       |     |  |
    |  +------------------+            +----------------------+     |  |
    |                                                               |  |
    |  +------------------+            +----------------------+     |  |
    |  | Electronics      |----------->| PCB Assembly         |--+  |  |
    |  | (ICs, Sensors)   |            | (Control Circuits)   |  |  |  |
    |  +------------------+            +----------------------+  |  |  |
    |                                                            v  v  |
    |                                  +----------------------+        |
    |                                  | FINAL ASSEMBLY &     |        |
    |                                  | QUALITY CONTROL      |        |
    |                                  +----------+-----------+        |
    +------------------------------------------------------------------+
                                                 |
    +------------------------------------------------------------------+
    |                    FRONTEND SUPPLY CHAIN                         |
    +------------------------------------------------------------------+
    |                                |                                 |
    |            +-------------------+-------------------+             |
    |            v                   v                   v             |
    |  +------------------+  +----------------+  +----------------+    |
    |  | Distributors/    |  | E-commerce    |  | Direct Brand   |    |
    |  | Wholesalers      |  | Platforms     |  | Stores         |    |
    |  | (15-25% margin)  |  | (12-18% fee)  |  |                |    |
    |  +--------+---------+  +-------+-------+  +-------+--------+    |
    |           |                    |                  |             |
    |           v                    v                  v             |
    |  +------------------+         +-------------------------+       |
    |  | Retail Stores    |         |    END CONSUMERS        |       |
    |  | (30-40% margin)  |-------->| Urban, Semi-urban       |       |
    |  +------------------+         | Students, Professionals |       |
    |                               +-------------------------+       |
    +------------------------------------------------------------------+
    """
    add_diagram(doc, supply_diagram, 6.2)

    add_body(doc, """The frontend supply chain distributes finished products through multiple channels. The traditional wholesale-retail channel involves manufacturers selling to distributors at ex-factory prices, distributors adding 15-25 percent markup before selling to retailers, and retailers adding 30-40 percent markup before selling to consumers. A product with Rs. 1,800-2,100 factory cost thus reaches consumers at Rs. 2,700-3,600 through traditional channels. E-commerce channels have captured 58 percent of Indian market sales, compressing margins and enabling 10-20 percent lower consumer prices while maintaining seller profitability.""")

    add_subheading(doc, "2.6 Wholesale Markets in Mumbai")

    add_body(doc, """Mumbai serves as a major wholesale hub for consumer electrical appliances in Western India. Hair straightening brushes are commonly available at wholesale rates in several key markets. Crawford Market, located in South Mumbai, functions as a traditional wholesale centre for diverse consumer goods including electrical appliances. Lohar Chawl specializes in electrical and hardware items, offering competitive wholesale prices due to high trader concentration. Lamington Road is known for electronics and electrical goods, with numerous shops offering hair styling appliances at wholesale rates. Bhiwandi, located on Mumbai's outskirts, serves as a major warehousing and distribution hub with lower operating costs enabling competitive wholesale pricing. Wholesale prices are typically 15-25 percent lower than retail prices, with bulk orders exceeding 100 units receiving additional 5-10 percent discounts.""")

    add_subheading(doc, "2.7 Government Policies and Regulations")

    add_body(doc, """Government policies significantly influence the production, pricing, and marketing of hair straightening brushes in India. The Goods and Services Tax (GST) applies an 18 percent rate to hair straightening brushes, classified under electrical appliances. Import duties on finished products are set at 20 percent basic customs duty, while electronic components attract duties ranging from 5-15 percent. Bureau of Indian Standards (BIS) certification is mandatory for electrical appliances sold in India, with initial certification fees of Rs. 50,000-100,000, annual licence fees, and testing costs. Consumer Protection Act 2019 mandates clear labelling, warranty provisions, and redressal mechanisms. E-waste management rules require manufacturers to establish collection and recycling mechanisms, adding compliance costs but promoting sustainable practices.""")

    # ========== SECTION 3: FRONTEND ANALYSIS ==========
    add_heading(doc, "3. FRONTEND ANALYSIS")

    add_subheading(doc, "3.1 Nature and Determinants of Demand")

    add_body(doc, """Demand for hair straightening brushes reflects derived demand from consumers' desires for specific hairstyles and appearance attributes rather than intrinsic product demand. The product is primarily demanded by urban and semi-urban consumers, particularly students, working professionals, and young adults who value convenience and time-saving grooming solutions. Since hair straightening brushes are not essential commodities, their demand depends significantly on income levels, tastes, and preferences. As disposable income rises, consumers are more willing to spend on personal grooming appliances, indicating positive income elasticity of demand.""")

    add_body(doc, """Key demand determinants include income effects where middle-class expansion with 8-10 percent annual income growth correlates with 7.1 percent annual demand growth in India. Demographic factors show the 18-35 age cohort accounts for approximately 68 percent of purchases, demonstrating higher frequency usage of 4-5 times weekly compared to 2-3 times for consumers above 35 years. Urban residents demonstrate 3-4 times higher ownership rates than rural populations due to income differentials, social media exposure, and reliable electricity access. Social media influence and influencer marketing increase purchase propensity by 35-45 percent. Seasonal factors show demand increasing during festivals and weddings when grooming needs rise.""")

    add_subheading(doc, "3.2 Elasticity of Demand Analysis")

    add_body(doc, """Price Elasticity of Demand (PED) measures the percentage change in quantity demanded resulting from a one percent change in price. For hair straightening brushes, price elasticity varies significantly across market segments. The budget segment below Rs. 2,500 exhibits highly elastic demand with PED estimated at -1.8 to -2.5, meaning a 10 percent price increase reduces quantity demanded by 18-25 percent. This high elasticity reflects multiple substitutes within the budget tier, income constraints among target consumers, and low brand loyalty enabling easy switching.""")

    add_body(doc, """The mid-range segment from Rs. 2,500 to Rs. 8,000 demonstrates moderately elastic demand with PED of -1.0 to -1.4. Brand differentiation by established players reduces perfect substitutability, while quality perception creates minimum threshold effects limiting downward substitution. The premium segment above Rs. 8,000 shows inelastic demand with PED of -0.4 to -0.7, where a 10 percent price increase reduces quantity demanded by only 4-7 percent. Premium consumers prioritize quality and brand over price, limited substitutes exist at equivalent quality levels, and higher incomes reduce price sensitivity.""")

    add_figure_caption(doc, "Figure 3.1: Demand Curve Showing Price Elasticity Across Segments")

    demand_diagram = """
                Price (Rs.)
                    |
               6000 + *
                    |  \\
               5000 +   \\      PREMIUM SEGMENT
                    |    \\     (Inelastic: PED = -0.5)
               4000 +     \\
                    |      \\
               3000 +       *-----------------
                    |        \\
               2500 +         \\    MID-RANGE SEGMENT
                    |          \\   (Unit Elastic: PED = -1.2)
               2000 +           \\
                    |            *----------------
               1500 +             \\
                    |              \\   BUDGET SEGMENT
               1000 +               \\  (Elastic: PED = -2.0)
                    |                \\
                500 +                 *
                    |                  \\
                  0 +---+----+----+----+----+----+----+----+----+-> Q
                       0   2k   4k   6k   8k  10k  12k  14k  16k
                            Quantity Demanded (Units/month)

    Revenue Implications:
    * Elastic demand: Price CUTS increase Total Revenue (TR)
    * Inelastic demand: Price HIKES increase Total Revenue (TR)
    * Unit elastic: TR remains constant with price changes
    """
    add_diagram(doc, demand_diagram, 6)

    add_body(doc, """Income Elasticity of Demand (YED) measures demand responsiveness to income changes. Hair straightening brushes exhibit positive income elasticity, classifying them as normal goods. In emerging markets like India with lower average incomes, estimated YED ranges from +1.5 to +2.2, indicating luxury good characteristics where demand grows faster than income. This transition explains faster market growth in developing economies where income growth translates into proportionally larger demand increases.""")

    add_body(doc, """Cross Elasticity of Demand (XED) measures demand responsiveness to price changes of related goods. Cross elasticity between straightening brushes and flat irons is estimated at +0.4 to +0.6, indicating substitute relationships. A 10 percent flat iron price increase would boost straightening brush demand by 4-6 percent. Cross elasticity with salon straightening services shows higher values of +0.6 to +0.8, as salon visits represent a more direct substitute for home styling.""")

    # TABLE 2: Elasticity Summary
    add_body(doc, "Table 3.1: Elasticity Measures Summary")
    elasticity_headers = ["Elasticity Type", "Estimate", "Classification", "Strategic Implication"]
    elasticity_rows = [
        ["Price Elasticity (Budget Segment)", "-1.8 to -2.5", "Elastic", "Volume-based pricing strategy"],
        ["Price Elasticity (Premium Segment)", "-0.4 to -0.7", "Inelastic", "Premium pricing viable"],
        ["Income Elasticity (India)", "+1.5 to +2.2", "Luxury good", "Target rising middle class"],
        ["Cross Elasticity (Flat Irons)", "+0.4 to +0.6", "Substitutes", "Monitor competitor pricing"],
        ["Cross Elasticity (Salon Services)", "+0.6 to +0.8", "Substitutes", "Emphasize convenience"],
    ]
    add_table(doc, elasticity_headers, elasticity_rows)

    add_subheading(doc, "3.3 Market Structure: Monopolistic Competition")

    add_body(doc, """The hair straightening brush market operates under monopolistic competition, characterized by many sellers offering differentiated products with relatively free entry and exit. Each firm attempts to distinguish its product through design, technology, safety features, branding, and advertising. Due to the presence of many sellers and close substitutes, individual firms have limited control over market prices. Product differentiation occurs along multiple dimensions including functional differentiation through heating technology and features, design differentiation through ergonomic and aesthetic elements, and brand differentiation through reputation and marketing.""")

    add_body(doc, """Market fragmentation is evident with numerous competitors including global brands like Philips, Dyson, and Remington, regional players like Vega and Havells in India, and numerous unbranded manufacturers. No single firm controls more than 15-20 percent market share. Entry barriers are moderate, including product development costs, manufacturing equipment investment, and quality certification requirements. In the short run, firms may earn supernormal profits through successful differentiation. In the long run, new entrants erode excess profits until firms earn normal profits where price equals average cost.""")

    add_subheading(doc, "3.4 Pricing Mechanism")

    add_body(doc, """Pricing in the hair straightening brush market reflects various strategies. Cost-plus pricing forms the foundation where the cost of production forms the base and a fixed margin is added, ensuring cost recovery and reasonable profit. For a product with Rs. 2,000 production cost, applying a 20 percent margin yields a manufacturer selling price of Rs. 2,400. Value-based pricing is employed by premium brands where price reflects perceived consumer value rather than production cost. Dyson products command Rs. 30,000-45,000 based on innovation and prestige. Dynamic pricing on e-commerce platforms adjusts prices based on demand patterns, competitor pricing, inventory levels, and consumer behaviour, with prices varying 15-25 percent across platforms and time periods.""")

    add_subheading(doc, "3.5 Revenue Analysis")

    add_body(doc, """Total Revenue (TR) equals price multiplied by quantity sold. For a firm selling 10,000 units at Rs. 2,500 average price, TR equals Rs. 2.5 crore. Average Revenue (AR) equals TR divided by quantity, representing the price received per unit. Under monopolistic competition, Marginal Revenue (MR) lies below AR due to the downward-sloping demand curve faced by each firm. Revenue distribution across the value chain shows manufacturers earning 55-60 percent of final retail price, distributors earning 12-15 percent, and retailers earning 25-30 percent.""")

    add_figure_caption(doc, "Figure 3.2: Revenue Curves under Monopolistic Competition")

    revenue_diagram = """
            Revenue/Cost (Rs.)
                    |
               3500 +
                    |*
               3000 + \\
                    |  \\        AR (Average Revenue = Demand Curve)
               2500 +   \\
                    |    \\
               2000 +     *--------------------------
                    |      \\                  MC (Marginal Cost)
               1500 +       \\              /
                    |        \\           /
               1000 +         \\        /
                    |          \\     /       MR (Marginal Revenue)
                500 +           \\  /
                    |            */
                  0 +----+----+--*-+----+----+----+----+----+----> Q
                        0   1k  2k  3k  4k  5k  6k  7k  8k
                            Quantity (Units)

    Profit Maximization Condition: MC = MR
    * At Q = 2,500 units: MC intersects MR
    * Price charged = AR at that quantity
    * Supernormal profit if AR > ATC at profit-maximizing output
    """
    add_diagram(doc, revenue_diagram, 6)

    add_subheading(doc, "3.6 Selling and Marketing Costs")

    add_body(doc, """Selling costs play an important role in influencing demand for hair straightening brushes. Unlike production costs, selling costs do not directly increase output but aim to shift the demand curve rightward by influencing consumer preferences. In monopolistically competitive markets, selling costs are essential for product differentiation and brand building. Advertising expenditure encompasses digital marketing, social media campaigns, and influencer partnerships, with large brands allocating 8-15 percent of revenue to advertising. Influencer marketing has become particularly significant for personal grooming products, with brands allocating 3-8 percent of marketing budgets to collaborations. Packaging investments serve dual purposes of product protection and brand communication, costing Rs. 30-150 per unit. After-sales service including warranty fulfilment and customer support accounts for 2-4 percent of revenue.""")

    # ========== SECTION 4: SUGGESTIONS ==========
    add_heading(doc, "4. SUGGESTIONS TO IMPROVE PRODUCTION AND DELIVERY EFFICIENCY")

    add_body(doc, """Based on the comprehensive microeconomic analysis conducted, the following recommendations are proposed to enhance production efficiency, reduce costs, and improve supply chain effectiveness in the hair straightening brush industry.""")

    add_body(doc, """First, manufacturers should increase automation in the production process to reduce dependence on manual labour, improve consistency in product quality, and lower average cost through economies of scale. Investment in automated assembly lines becomes economically viable at production volumes exceeding 150,000 units annually. Automated processes reduce per-unit labour costs by 70-80 percent while improving quality consistency through precise component placement and standardized processes.""")

    add_body(doc, """Second, bulk procurement of raw materials and establishment of long-term contracts with suppliers should be adopted to reduce input price fluctuations and achieve lower per-unit raw material costs. Negotiating annual supply agreements with key component suppliers can secure 5-10 percent price reductions compared to spot market purchasing. Building strategic partnerships with reliable suppliers also ensures consistent quality and timely delivery.""")

    add_body(doc, """Third, the supply chain should be streamlined by reducing unnecessary intermediaries, which would lower transportation, storage, and inventory holding costs, leading to more competitive pricing. Direct-to-retailer distribution models and e-commerce partnerships can eliminate distributor margins of 15-25 percent. Investment in warehouse management systems and logistics optimization can reduce inventory holding costs by 20-30 percent.""")

    add_body(doc, """Fourth, improved demand forecasting based on past sales trends, seasonal demand patterns, and market intelligence should be implemented to prevent overproduction or underproduction, thereby reducing wastage and excess inventory costs. Advanced analytics and machine learning algorithms can improve forecast accuracy by 15-25 percent compared to traditional methods.""")

    add_body(doc, """Fifth, adoption of sustainable materials and energy-efficient production processes addresses growing consumer preference for environmentally responsible products while potentially reducing costs. Recycled plastics can reduce material costs by 10-15 percent while enhancing brand image. Energy-efficient machinery reduces utility costs and supports environmental compliance requirements.""")

    add_body(doc, """Sixth, investment in quality management systems beyond minimum compliance requirements can reduce warranty claims and returns from 3-5 percent to below 1 percent, substantially reducing after-sales costs and improving customer retention through enhanced brand reputation.""")

    # ========== SECTION 5: LESSONS LEARNED ==========
    add_heading(doc, "5. LESSONS LEARNED")

    add_body(doc, """Through the comprehensive study of hair straightening brushes as a microeconomic product, I have gained valuable insights into the practical application of economic theory to real-world markets. The following points summarize my key learnings from this assignment.""")

    add_body(doc, """First, I learned how production decisions are fundamentally influenced by the cost structure of a firm, particularly the distinction between fixed costs and variable costs. Understanding that fixed costs like machinery and factory rent remain constant regardless of output while variable costs like raw materials increase with production volume helped me appreciate why firms seek economies of scale. The analysis revealed how average cost initially declines as fixed costs are spread over larger output but eventually rises due to diminishing marginal returns. This understanding of cost behaviour provides crucial insight into firm decision-making regarding optimal production levels and capacity utilization.""")

    add_body(doc, """Second, the study highlighted the critical importance of price elasticity of demand in determining pricing strategies. I learned that when demand is elastic, as observed in the budget segment of hair straightening brushes, firms must focus on increasing sales volume rather than raising prices because price cuts increase total revenue. Conversely, premium segments with inelastic demand can sustain higher prices without significant volume loss. This differentiated elasticity across market segments explains why firms employ segment-specific pricing strategies and why premium brands can maintain substantial price premiums over budget alternatives.""")

    add_body(doc, """Third, the analysis demonstrated how efficiency in the supply chain directly affects final market prices and profitability. I understood the economic rationale for intermediaries including wholesalers and retailers who add value through inventory holding, assortment creation, and geographic distribution despite adding to final prices. The emergence of e-commerce as a more efficient distribution channel illustrated how technological innovation can reduce transaction costs and create consumer welfare gains. The shift to online channels capturing 58 percent of Indian market sales demonstrates market responses to cost efficiency opportunities.""")

    add_body(doc, """Fourth, the assignment provided comprehensive insight into the functioning of monopolistic competition, where product differentiation and selling costs play significant roles in influencing consumer demand. I learned how firms under monopolistic competition compete not only on price but also on product features, branding, and marketing. The presence of many sellers with differentiated products means no single firm controls market prices, yet differentiation creates limited market power. Understanding this market structure explains the prevalence of advertising and brand building in consumer goods markets.""")

    add_body(doc, """Fifth, I gained appreciation for how government policies including taxation, safety standards, and import regulations impact production costs, market prices, and overall consumer welfare. The 18 percent GST, import duties on components, and mandatory BIS certification all contribute to the final product price while serving legitimate policy objectives including revenue generation, domestic manufacturing protection, and consumer safety. This regulatory analysis demonstrates the interconnection between microeconomic outcomes and macroeconomic policy frameworks.""")

    # ========== REFERENCES ==========
    add_heading(doc, "REFERENCES")

    refs = [
        "[1] Grand View Research. (2024). Hair Styling Tools Market Size & Share Report.",
        "[2] Market Research Future. (2024). India Hair Care Appliances Market Analysis.",
        "[3] Verified Market Reports. (2024). Hair Straightening Brush Market Report.",
        "[4] IBEF. (2024). Indian Consumer Electronics Industry Analysis.",
        "[5] ClearTax. (2024). GST Rates for Electrical Appliances - HSN Code.",
        "[6] Bureau of Indian Standards. (2024). Certification Requirements.",
        "[7] Ministry of Commerce. (2024). Import Duty Structure for Electronics.",
        "[8] Investopedia. (2024). Price Elasticity of Demand - Concepts.",
        "[9] Salvatore, D. Microeconomics: Theory and Applications.",
        "[10] Dwivedi, D.N. Managerial Economics.",
    ]

    for ref in refs:
        para = doc.add_paragraph()
        run = para.add_run(ref)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        para.paragraph_format.line_spacing = 1.5

    # Save
    output_path = '/mnt/e/AI and Projects/MMS-Prep/Eco Assign/shriya/Hair_Straightening_Brush_Assignment_Final.docx'
    doc.save(output_path)
    print(f"Document saved: {output_path}")

    word_count = sum(len(p.text.split()) for p in doc.paragraphs)
    print(f"Word count: {word_count}")
    print(f"Estimated pages: {word_count/280:.1f}")

if __name__ == "__main__":
    main()
