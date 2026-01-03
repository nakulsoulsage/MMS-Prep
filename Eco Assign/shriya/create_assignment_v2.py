#!/usr/bin/env python3
"""
Create Microeconomics Product Assignment on Hair Straightening Brush
Revised: 16 pages, 3-4 tables, economic diagrams
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_page_margins(doc):
    """Set A4 page with specified margins"""
    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1)

def add_page_numbers(doc):
    """Add page numbers to footer"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

def add_heading(doc, text):
    """Add main heading"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.5

def add_subheading(doc, text):
    """Add subheading"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.5

def add_body(doc, text):
    """Add body paragraph"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(8)

def add_figure_caption(doc, text):
    """Add figure caption"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(12)

def add_table(doc, headers, rows):
    """Add formatted table"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for para in header_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data)
            for para in row_cells[col_idx].paragraphs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)

    doc.add_paragraph()

def add_diagram_placeholder(doc, diagram_text, width_inches=5.5):
    """Add a text-based diagram/graph representation"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Create a bordered text box effect using a single-cell table
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell = table.rows[0].cells[0]
    cell.width = Inches(width_inches)

    # Add the diagram content
    cell_para = cell.paragraphs[0]
    run = cell_para.add_run(diagram_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

def create_cover_page(doc):
    """Create cover page"""
    for _ in range(2):
        doc.add_paragraph()

    # Title
    para = doc.add_paragraph()
    run = para.add_run("MICROECONOMICS PRODUCT ASSIGNMENT")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run("HAIR STRAIGHTENING BRUSH")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(20)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run("A Microeconomic Analysis of Production, Cost, Demand, and Market Dynamics")
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _ in range(4):
        doc.add_paragraph()

    # Student details
    details = [
        ("Name:", "Shriya Parab"),
        ("Roll No:", "HBM25120"),
        ("Class:", "MMS Batch 2025-27"),
        ("Semester:", "01st"),
        ("Academic Year:", "2025-2026"),
        ("Institute:", "Sydenham Institute of Management Studies,"),
        ("", "Research and Entrepreneurship Education, Mumbai"),
    ]

    for label, value in details:
        para = doc.add_paragraph()
        if label:
            run1 = para.add_run(f"{label}  ")
            run1.bold = True
            run1.font.name = 'Times New Roman'
            run1.font.size = Pt(12)
        run2 = para.add_run(value)
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

def main():
    doc = Document()
    set_page_margins(doc)
    add_page_numbers(doc)
    create_cover_page(doc)

    # ========== SECTION 1: INTRODUCTION ==========
    add_heading(doc, "1. INTRODUCTION AND EVOLUTION")

    add_subheading(doc, "1.1 Definition of Hair Straightening Brush as an Economic Product")

    add_body(doc, """A hair straightening brush is an electrical personal grooming appliance that combines the functionality of a conventional hairbrush with the heating mechanism of an electric hair straightener. From a microeconomic perspective, this product represents a convergence of thermal engineering and consumer technology, addressing the derived demand for personal grooming solutions. The product occupies a distinct market position within the USD 30 billion global hair styling tools industry.""")

    add_body(doc, """The global market for hair straightening brushes reached USD 620-750 million in 2024, with projections indicating growth to USD 1.53 billion by 2033 at a compound annual growth rate (CAGR) of 8.2 percent. In India, the market stood at USD 22.17 million in 2024 and is expected to reach USD 33.32 million by 2033 with a CAGR of 4.28 percent. Hair straightening brushes are generally available in the Indian market in the price range of Rs. 1,500 to Rs. 3,500, depending on brand reputation, material quality, technological features, and safety mechanisms.""")

    add_subheading(doc, "1.2 Historical Evolution of Hair Styling Technology")

    add_body(doc, """The evolution of hair straightening technology exemplifies the economic principle of creative destruction, where successive innovations displace existing methods through superior cost-benefit propositions. Historically, hair straightening was carried out using manual methods such as heated metal rods or chemical straightening treatments. These methods were mostly salon-based, expensive, and time-consuming. The breakthrough came in 1909 when Isaac K. Shero patented metal plates capable of gliding through hair with adjustable heat settings.""")

    add_body(doc, """The twentieth century witnessed exponential technological advancement driven by electrification and materials science. The adoption of plastic materials through injection moulding in the mid-1900s democratized access by dramatically lowering production costs. Contemporary hair straightening brushes emerged in the early 2000s as manufacturers recognized an underserved market segment: consumers seeking straightening results with minimal technical skill and time investment. By integrating heating elements into brush formats, designers reduced the learning curve and switching costs from traditional methods.""")

    add_subheading(doc, "1.3 Technological Innovations")

    add_body(doc, """Modern hair straightening brushes incorporate several technological innovations. Ceramic heating technology provides even heat distribution that minimizes hot spots, reducing energy requirements by 15-20 percent compared to aluminium. Ionic technology emits negatively charged ions that break down water molecules on hair surfaces, reducing drying time and static electricity. Premium models with ionic technology command price premiums of 40-60 percent over basic alternatives.""")

    add_body(doc, """PTC (Positive Temperature Coefficient) heating elements utilize semiconductor properties that cause electrical resistance to increase at predetermined temperatures, creating self-regulating heating without external thermostats. Tourmaline and titanium coatings further enhance performance, with titanium plates offering superior heat conductivity and durability, justifying price premiums of 50-80 percent in professional-grade segments.""")

    add_subheading(doc, "1.4 Comparison with Substitutes")

    add_body(doc, """Hair straightening brushes compete with several substitute products including traditional flat irons requiring higher skill levels, blow dryers with straightening attachments, and professional salon services involving higher costs and time investments. The straightening brush offers advantages in ease of use, reduced heat damage risk, and time efficiency, making it attractive for the convenience-seeking urban consumer segment.""")

    add_subheading(doc, "1.5 Product Lifecycle Classification: Sunrise Product")

    add_body(doc, """Hair straightening brushes are classified as a Sunrise product within the product lifecycle framework. The sustained CAGR of 5.2-8.2 percent across major markets indicates continued expansion characteristic of growth-stage products. Market penetration shows 45 percent of women in developed markets owned hair straighteners by 2024, with straightening brushes representing approximately 30 percent of this category. The introduction of cordless variants and smart-enabled models demonstrates ongoing innovation typical of sunrise products. Unlike mature commodity markets characterized by price compression, hair straightening brushes maintain significant price dispersion from Rs. 300 to Rs. 45,000, suggesting viable product differentiation and market segmentation.""")

    # ========== SECTION 2: BACKEND ANALYSIS ==========
    add_heading(doc, "2. BACKEND ANALYSIS")

    add_subheading(doc, "2.1 Production Modality and Technology")

    add_body(doc, """Hair straightening brushes are produced using electro-mechanical manufacturing processes. The production process involves assembling ceramic heating elements, electrical wiring, temperature control circuits, and plastic casings manufactured through injection moulding. Product development begins with computer-aided design (CAD) modelling to optimize heat distribution patterns, bristle configurations, and ergonomic handling. A typical product development cycle spans 12-18 months with expenditures ranging from USD 50,000 for budget models to over USD 2 million for premium innovations.""")

    add_body(doc, """Small and medium-scale manufacturers operate semi-automatic assembly lines producing approximately 500 to 2,000 units per day. Large manufacturers employ automated production systems capable of producing 5,000 units or more per day. Over time, production has shifted from labour-intensive to capital-intensive methods. In microeconomic terms, technological improvement increases productivity and leads to a downward shift in the average cost curve, enabling firms to achieve economies of scale.""")

    add_subheading(doc, "2.2 Raw Materials and Factor Market")

    add_body(doc, """The production requires a combination of raw materials, labour, and capital inputs. Ceramic heating elements are either procured domestically or imported from China. ABS plastic housings with heat deflection temperature of 98-105 degrees Celsius are sourced from domestic manufacturers. Electrical components including wiring, thermostats, and control circuits are procured from India and East Asian countries. Labour is hired locally with wage rates depending on skill levels, while machinery and moulds are sourced from domestic and international suppliers.""")

    add_body(doc, """Global ABS pricing averaged USD 1.40-1.80 per kilogram in 2024, with a typical straightener housing requiring 120-180 grams of ABS, translating to material costs of Rs. 15-25 per unit. Aluminium for heating plates costs USD 2,200-2,600 per metric ton, contributing Rs. 8-15 to material costs. PTC ceramic heating elements with tight manufacturing tolerances increase per-unit costs by 15-25 percent relative to standard resistance heating but provide superior safety and energy efficiency.""")

    add_subheading(doc, "2.3 Cost Structure Analysis")

    add_body(doc, """The cost structure consists of operating expenditure (Op-Ex) and capital expenditure (Cap-Ex). Capital expenditure includes machinery, moulds, factory infrastructure, research and development, and regulatory certifications. Operating costs include raw materials, labour wages, electricity, packaging, transportation, and logistics.""")

    # TABLE 1: Cost Structure
    add_body(doc, "Table 2.1: Cost Structure per Unit")
    cost_headers = ["Cost Component", "Cost (Rs.)", "Percentage", "Type"]
    cost_rows = [
        ["Raw materials", "900-1,200", "45-50%", "Variable"],
        ["Labour and wages", "200-300", "10-12%", "Variable"],
        ["Electricity and utilities", "80-120", "4-5%", "Variable"],
        ["Quality control testing", "50-80", "3-4%", "Variable"],
        ["Packaging materials", "100-150", "5-6%", "Variable"],
        ["Transportation and logistics", "150-250", "8-10%", "Variable"],
        ["Depreciation and overheads", "200-300", "10-12%", "Fixed"],
        ["Total Production Cost", "1,800-2,400", "100%", "-"],
    ]
    add_table(doc, cost_headers, cost_rows)

    add_body(doc, """Raw materials constitute the largest share of total cost at approximately 45-50 percent, making production highly sensitive to input price fluctuations. Fixed costs remain constant in the short run, while variable costs increase with output.""")

    add_subheading(doc, "2.4 Short-Run and Long-Run Cost Behaviour")

    add_body(doc, """In the short run, at least one factor of production remains fixed. Total Fixed Cost (TFC) remains constant regardless of output level. Total Variable Cost (TVC) increases with output. Average Fixed Cost (AFC) declines continuously as output increases because the same fixed cost is spread over more units. Average Variable Cost (AVC) initially decreases due to specialization and efficiency gains, reaches a minimum at optimal capacity utilization, then increases due to diminishing marginal returns.""")

    # DIAGRAM: Cost Curves
    add_figure_caption(doc, "Figure 2.1: Short-Run Cost Curves in Hair Straightening Brush Manufacturing")

    cost_diagram = """
                    Cost (Rs.)
                        |
                   3000 |                              MC
                        |                            /
                   2500 |                          /
                        |                        /    ATC
                   2000 |        *-------------*--------
                        |      /                 \\
                   1500 |    /                     \\    AVC
                        |  /          *-------------*----
                   1000 |/          /
                        |         /
                    500 |-------/-------------------------------- AFC
                        |
                      0 +----------------------------------------
                        0    100   200   300   400   500   600
                                    Output (Units)

    Where: MC = Marginal Cost, ATC = Average Total Cost
           AVC = Average Variable Cost, AFC = Average Fixed Cost

    * MC intersects ATC at its minimum point (optimal scale)
    * AFC continuously declines as output increases
    * AVC is U-shaped due to diminishing marginal returns
    """
    add_diagram_placeholder(doc, cost_diagram, 6)

    add_body(doc, """In the long run, all factors become variable, allowing firms to adjust plant size. The Long-Run Average Cost (LRAC) curve represents the lowest average cost achievable at each output level. Economies of scale reduce per-unit costs as output expands through technical economies from larger machinery, managerial economies from specialized functions, marketing economies from bulk advertising, and purchasing economies from bulk procurement. Chinese mega-factories producing millions of units annually achieve per-unit costs 30-40 percent below smaller facilities.""")

    add_subheading(doc, "2.5 Supply Chain Analysis")

    add_body(doc, """The backend supply chain begins with raw material suppliers including plastic resin producers, ceramic manufacturers, and electronic component suppliers. These tier-2 suppliers feed into tier-1 suppliers who manufacture subassemblies. Final assembly integrates components into finished products. Approximately 60 percent of manufacturers utilize in-house final assembly while outsourcing component production.""")

    # DIAGRAM: Supply Chain
    add_figure_caption(doc, "Figure 2.2: Supply Chain Architecture")

    supply_chain_diagram = """
    BACKEND SUPPLY CHAIN                    FRONTEND SUPPLY CHAIN

    +------------------+                    +------------------+
    | Raw Material     |                    | Wholesalers/     |
    | Suppliers        |                    | Distributors     |
    | - Plastics       |                    | (Mumbai Markets) |
    | - Ceramics       |                    +--------+---------+
    | - Electronics    |                             |
    +--------+---------+                             v
             |                              +------------------+
             v                              | Retailers        |
    +------------------+                    | - Electronics    |
    | Component        |                    |   Stores         |
    | Manufacturers    |                    | - E-commerce     |
    | - Heating units  |                    +--------+---------+
    | - Housings       |                             |
    | - PCB Assembly   |                             v
    +--------+---------+                    +------------------+
             |                              | End Consumers    |
             v                              | - Urban women    |
    +------------------+                    | - Students       |
    | Final Assembly   |-------------------->| - Professionals  |
    | & Quality Control|                    +------------------+
    +------------------+

    Value Flow: Manufacturer Price (Rs.1,800-2,100) --> Wholesale (Rs.2,100-2,600)
                --> Retail (Rs.2,700-3,600)
    """
    add_diagram_placeholder(doc, supply_chain_diagram, 6)

    add_body(doc, """The frontend supply chain distributes finished products through multiple channels. The traditional wholesale-retail channel involves manufacturers selling to distributors at ex-factory prices, distributors adding 15-25 percent markup, and retailers adding 30-40 percent markup. E-commerce channels, capturing 58 percent of Indian market sales, compress these margins, enabling 10-20 percent lower consumer prices.""")

    add_subheading(doc, "2.6 Wholesale Markets in Mumbai")

    add_body(doc, """Mumbai serves as a major wholesale hub for consumer electrical appliances. Hair straightening brushes are available at wholesale rates in Crawford Market, Lohar Chawl specializing in electrical items, Lamington Road known for electronics, and Bhiwandi serving as a warehousing hub. Wholesale prices are typically 15-25 percent lower than retail prices, with bulk orders exceeding 100 units receiving additional 5-10 percent discounts.""")

    add_subheading(doc, "2.7 Government Policies and Regulations")

    add_body(doc, """Government policies significantly influence production and pricing. The Goods and Services Tax (GST) applies an 18 percent rate to hair straightening brushes. Import duties of 20 percent apply to finished products, while electronic components attract duties of 5-15 percent. Bureau of Indian Standards (BIS) certification is mandatory, with initial certification fees of Rs. 50,000-100,000. Consumer Protection Act 2019 mandates clear labelling and warranty provisions. E-waste management rules require manufacturers to establish collection and recycling mechanisms.""")

    # ========== SECTION 3: FRONTEND ANALYSIS ==========
    add_heading(doc, "3. FRONTEND ANALYSIS")

    add_subheading(doc, "3.1 Nature and Determinants of Demand")

    add_body(doc, """Demand for hair straightening brushes reflects derived demand from consumers' desires for specific hairstyles and appearance attributes. The product is primarily demanded by urban and semi-urban consumers, particularly students, working professionals, and young adults aged 18-35 who account for approximately 68 percent of purchases. Since hair straightening brushes are not essential commodities, their demand depends significantly on income levels, tastes, and preferences.""")

    add_body(doc, """Key demand determinants include income levels with positive income elasticity, availability of substitutes affecting price sensitivity, fashion and grooming trends shifting preferences, urban lifestyle driving convenience-based demand, and seasonal factors with demand increasing during festivals and weddings. Social media influence and influencer marketing increase purchase propensity by 35-45 percent through reduced uncertainty about product performance.""")

    add_subheading(doc, "3.2 Elasticity of Demand")

    add_body(doc, """Price Elasticity of Demand (PED) varies across market segments. The budget segment below Rs. 2,500 exhibits highly elastic demand with PED of -1.8 to -2.5, meaning a 10 percent price increase reduces quantity demanded by 18-25 percent. This reflects multiple substitutes, income constraints, and low brand loyalty. The mid-range segment (Rs. 2,500-8,000) shows moderately elastic demand with PED of -1.0 to -1.4. The premium segment above Rs. 8,000 demonstrates inelastic demand with PED of -0.4 to -0.7, where consumers prioritize quality over price.""")

    # DIAGRAM: Demand Curve and Elasticity
    add_figure_caption(doc, "Figure 3.1: Demand Curve and Price Elasticity")

    demand_diagram = """
                    Price (Rs.)
                        |
                   5000 |*
                        | \\
                   4000 |  \\     Premium Segment (Inelastic)
                        |   \\    PED = -0.5
                   3000 |    *-------------------
                        |     \\
                   2500 |      \\  Mid-Range (Unit Elastic)
                        |       \\ PED = -1.2
                   2000 |        *------------------
                        |         \\
                   1500 |          \\  Budget Segment (Elastic)
                        |           \\ PED = -2.0
                   1000 |            *
                        |             \\
                    500 |              *----------------> D
                        |
                      0 +----------------------------------------
                        0    2000  4000  6000  8000  10000
                                Quantity Demanded (Units)

    Revenue Implications:
    - Elastic demand (Budget): Price cuts increase Total Revenue
    - Inelastic demand (Premium): Price hikes increase Total Revenue
    """
    add_diagram_placeholder(doc, demand_diagram, 6)

    add_body(doc, """Income Elasticity of Demand (YED) is estimated at +1.5 to +2.2 in emerging markets like India, indicating luxury good characteristics where demand grows faster than income. Cross Elasticity of Demand (XED) with flat irons is +0.4 to +0.6 indicating substitute relationships, while cross elasticity with salon services shows higher values of +0.6 to +0.8.""")

    add_subheading(doc, "3.3 Market Structure: Monopolistic Competition")

    add_body(doc, """The hair straightening brush market operates under monopolistic competition, characterized by many sellers offering differentiated products with relatively free entry and exit. Each firm attempts to distinguish its product through design, technology, safety features, branding, and advertising. Market fragmentation is evident with global brands like Philips, Dyson, and Remington, regional players like Vega and Havells, and numerous unbranded manufacturers. No single firm controls more than 15-20 percent market share.""")

    add_body(doc, """Entry barriers are moderate including product development costs, manufacturing equipment investment, and quality certification requirements. In the short run, firms may earn supernormal profits through successful differentiation. In the long run, new entrants erode excess profits until firms earn normal profits where price equals average cost.""")

    add_subheading(doc, "3.4 Pricing Mechanism")

    add_body(doc, """Pricing reflects various strategies. Cost-plus pricing forms the foundation where production cost plus a fixed margin ensures cost recovery. For a product with Rs. 2,000 production cost, applying 20 percent margin yields Rs. 2,400 manufacturer price. Value-based pricing is employed by premium brands where price reflects perceived consumer value. Dyson products command Rs. 30,000-45,000 based on innovation and prestige. Dynamic pricing on e-commerce platforms adjusts based on demand patterns and competitor pricing, with prices varying 15-25 percent across platforms.""")

    add_subheading(doc, "3.5 Revenue Analysis")

    add_body(doc, """Total Revenue (TR) equals price multiplied by quantity. For a firm selling 10,000 units at Rs. 2,500 average price, TR equals Rs. 2.5 crore. Revenue distribution shows manufacturers earning 55-60 percent of retail price, distributors earning 12-15 percent, and retailers earning 25-30 percent. The relationship between elasticity and revenue is crucial: when demand is elastic, price reductions increase total revenue; when inelastic, price increases raise total revenue.""")

    # DIAGRAM: Revenue Curves
    add_figure_caption(doc, "Figure 3.2: Revenue Curves under Monopolistic Competition")

    revenue_diagram = """
                    Revenue/Price (Rs.)
                        |
                   3500 |
                        |*
                   3000 | \\
                        |  \\     AR (Average Revenue = Demand)
                   2500 |   \\
                        |    \\
                   2000 |     *-----------------------
                        |      \\
                   1500 |       \\        MR (Marginal Revenue)
                        |        \\
                   1000 |         \\
                        |          *
                    500 |           \\
                        |            \\
                      0 +-------------*------------------
                        0    2000  4000  6000  8000
                                Quantity (Units)

    Key Relationships:
    - AR lies above MR (downward sloping demand)
    - MR = AR when elasticity is infinite
    - MR = 0 when elasticity equals 1
    - Profit maximization: MC = MR
    """
    add_diagram_placeholder(doc, revenue_diagram, 6)

    add_subheading(doc, "3.6 Selling and Marketing Costs")

    add_body(doc, """Selling costs play an important role in influencing demand. Unlike production costs, selling costs aim to shift the demand curve rightward by influencing consumer preferences. Advertising expenditure encompasses digital marketing, social media campaigns, and influencer partnerships, with large brands allocating 8-15 percent of revenue to advertising. Influencer marketing has become particularly significant, with brands allocating 3-8 percent of marketing budgets to influencer collaborations. Packaging investments signal product quality and justify price premiums, costing Rs. 30-150 per unit. After-sales service including warranty fulfilment and customer support accounts for 2-4 percent of revenue.""")

    # TABLE 2: Elasticity Summary
    add_body(doc, "Table 3.1: Elasticity Measures Summary")
    elasticity_headers = ["Elasticity Type", "Estimate", "Nature", "Interpretation"]
    elasticity_rows = [
        ["Price Elasticity (Budget)", "-1.8 to -2.5", "Elastic", "Volume strategy optimal"],
        ["Price Elasticity (Premium)", "-0.4 to -0.7", "Inelastic", "Premium pricing viable"],
        ["Income Elasticity (India)", "+1.5 to +2.2", "Luxury good", "Grows with income"],
        ["Cross Elasticity (Flat Irons)", "+0.4 to +0.6", "Substitute", "Moderate competition"],
        ["Cross Elasticity (Salons)", "+0.6 to +0.8", "Substitute", "Higher substitutability"],
    ]
    add_table(doc, elasticity_headers, elasticity_rows)

    # ========== SECTION 4: SUGGESTIONS ==========
    add_heading(doc, "4. SUGGESTIONS TO IMPROVE PRODUCTION AND DELIVERY EFFICIENCY")

    add_body(doc, """Based on the microeconomic analysis conducted, the following recommendations are proposed to enhance production efficiency and improve supply chain effectiveness:""")

    add_body(doc, """First, manufacturers should increase automation in production processes to reduce dependence on manual labour, improve consistency in product quality, and lower average cost through economies of scale. Investment in automated assembly lines becomes economically viable at production volumes exceeding 150,000 units annually, reducing per-unit labour costs by 70-80 percent.""")

    add_body(doc, """Second, bulk procurement of raw materials and establishment of long-term contracts with suppliers should be adopted to reduce input price fluctuations and achieve lower per-unit material costs. Negotiating annual supply agreements can secure 5-10 percent price reductions compared to spot market purchasing.""")

    add_body(doc, """Third, the supply chain should be streamlined by reducing unnecessary intermediaries, which would lower transportation, storage, and inventory holding costs. Direct-to-retailer distribution models and e-commerce partnerships can eliminate distributor margins of 15-25 percent.""")

    add_body(doc, """Fourth, improved demand forecasting based on past sales trends and seasonal patterns should be implemented to prevent overproduction or underproduction, thereby reducing wastage and excess inventory costs. Advanced analytics can improve forecast accuracy by 15-25 percent.""")

    add_body(doc, """Fifth, adoption of sustainable materials including recycled plastics can reduce material costs by 10-15 percent while enhancing brand image among environmentally conscious consumers. Energy-efficient production processes reduce utility costs and support regulatory compliance.""")

    add_body(doc, """Sixth, investment in quality management systems beyond minimum compliance can reduce warranty claims and returns from 3-5 percent to below 1 percent, substantially reducing after-sales costs and improving customer retention through enhanced brand reputation.""")

    # ========== SECTION 5: LESSONS LEARNED ==========
    add_heading(doc, "5. LESSONS LEARNED")

    add_body(doc, """Through the comprehensive study of hair straightening brushes as a microeconomic product, I have gained valuable insights into the practical application of economic theory to real-world markets. The following points summarize my key learnings from this assignment.""")

    add_body(doc, """First, I learned how production decisions are fundamentally influenced by the cost structure of a firm, particularly the distinction between fixed costs and variable costs. Understanding that fixed costs like machinery and factory rent remain constant regardless of output while variable costs like raw materials increase with production volume helped me appreciate why firms seek economies of scale. The analysis revealed how average cost initially declines as fixed costs are spread over larger output but eventually rises due to diminishing marginal returns, explaining optimal production decisions.""")

    add_body(doc, """Second, the study highlighted the critical importance of price elasticity of demand in determining pricing strategies. I learned that when demand is elastic, as observed in the budget segment, firms must focus on increasing sales volume rather than raising prices because price cuts increase total revenue. Conversely, premium segments with inelastic demand can sustain higher prices without significant volume loss. This differentiated elasticity across market segments explains segment-specific pricing strategies.""")

    add_body(doc, """Third, the analysis demonstrated how efficiency in the supply chain directly affects final market prices and profitability. I understood the economic rationale for intermediaries including wholesalers and retailers who add value through inventory holding, assortment creation, and geographic distribution. The emergence of e-commerce as a more efficient distribution channel illustrated how technological innovation reduces transaction costs and creates consumer welfare gains.""")

    add_body(doc, """Fourth, the assignment provided comprehensive insight into the functioning of monopolistic competition, where product differentiation and selling costs play significant roles in influencing consumer demand. I learned how firms compete not only on price but also on product features, branding, and marketing. The presence of many sellers with differentiated products means no single firm controls market prices, yet differentiation creates limited market power enabling firms to earn short-run supernormal profits.""")

    add_body(doc, """Fifth, I gained appreciation for how government policies including taxation, safety standards, and import regulations impact production costs, market prices, and overall consumer welfare. The 18 percent GST, import duties, and mandatory BIS certification contribute to final product price while serving legitimate policy objectives of revenue generation, domestic manufacturing protection, and consumer safety.""")

    # ========== REFERENCES ==========
    add_heading(doc, "REFERENCES")

    references = [
        "[1] Grand View Research. (2024). Hair Styling Tools Market Size & Share Report.",
        "[2] Market Research Future. (2024). India Hair Care Appliances Market Analysis.",
        "[3] Verified Market Reports. (2024). Hair Straightening Brush Market Report.",
        "[4] IBEF. (2024). Indian Consumer Electronics Industry Analysis.",
        "[5] ClearTax. (2024). GST Rates for Electrical Appliances - HSN Code.",
        "[6] Bureau of Indian Standards. (2024). Certification Requirements for Electrical Products.",
        "[7] Ministry of Commerce. (2024). Import Duty Structure for Consumer Electronics.",
        "[8] Investopedia. (2024). Price Elasticity of Demand - Economic Concepts.",
        "[9] Microeconomics: Theory and Applications, Dominick Salvatore.",
        "[10] Managerial Economics, D.N. Dwivedi.",
    ]

    for ref in references:
        para = doc.add_paragraph()
        run = para.add_run(ref)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        para.paragraph_format.line_spacing = 1.5

    # Save
    output_path = '/mnt/e/AI and Projects/MMS-Prep/Eco Assign/shriya/Hair_Straightening_Brush_Assignment_Final.docx'
    doc.save(output_path)
    print(f"Document saved to: {output_path}")

    word_count = sum(len(p.text.split()) for p in doc.paragraphs)
    print(f"Approximate word count: {word_count}")

if __name__ == "__main__":
    main()
