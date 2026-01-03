#!/usr/bin/env python3
"""
Microeconomics Product Assignment: BUTTON (Garments & Apparel)
Student: Astha Agrawal | Roll: HBM25002
Styled after Nakul's document format
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
import os
from io import BytesIO

# Set output directory
OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# Color constants (matching Nakul's style)
HEADER_BLUE = "1F4E79"  # Dark blue for table headers
HEADER_BLUE_RGB = RGBColor(31, 78, 121)
ALT_ROW_COLOR = "D6DCE4"  # Light gray for alternating rows

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_borders(cell, color="000000", size="4"):
    """Set cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_page_number(doc):
    """Add page numbers to document"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

def create_document():
    """Create the main document with proper formatting"""
    doc = Document()

    # Set page margins (in inches): Top=1, Right=1, Bottom=1, Left=1.5
    for section in doc.sections:
        section.page_width = Inches(8.27)  # A4 width
        section.page_height = Inches(11.69)  # A4 height
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.5)

    return doc

def add_heading_style(doc, text, level=1, bold=True):
    """Add a properly formatted heading"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = bold
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_subheading(doc, text):
    """Add a subheading"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_body_text(doc, text, bold_words=None):
    """Add body text with proper formatting"""
    p = doc.add_paragraph()

    if bold_words:
        remaining = text
        for bold_word in bold_words:
            if bold_word in remaining:
                parts = remaining.split(bold_word, 1)
                if parts[0]:
                    run = p.add_run(parts[0])
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                run = p.add_run(bold_word)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True
                remaining = parts[1] if len(parts) > 1 else ""
        if remaining:
            run = p.add_run(remaining)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    else:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(8)
    return p

def add_bullet_point(doc, text, indent=0.5):
    """Add a bullet point"""
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_elegant_table(doc, headers, data, title=None):
    """Add an elegant table with dark blue headers (Nakul's style)"""
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row - Dark blue background with white text
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        set_cell_shading(cell, HEADER_BLUE)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)  # White text

    # Data rows with alternating colors
    for row_idx, row_data in enumerate(data):
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(cell_data)
            # Alternate row coloring
            if row_idx % 2 == 1:
                set_cell_shading(cell, ALT_ROW_COLOR)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)

    # Add spacing after table
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    return table

def create_cost_curve_graph():
    """Create clean Average Cost and Marginal Cost curves"""
    plt.style.use('seaborn-v0_8-whitegrid')
    fig, ax = plt.subplots(figsize=(7, 4.5))

    # Generate data points
    quantity = np.linspace(50, 500, 100)

    # Average Total Cost (U-shaped)
    atc = 10.5 - 0.025*quantity + 0.00004*quantity**2
    # Marginal Cost
    mc = np.full_like(quantity, 2.0) + 0.00002*quantity
    # Average Variable Cost
    avc = 3.5 - 0.008*quantity + 0.00002*quantity**2

    ax.plot(quantity, atc, 'b-', linewidth=2.5, label='Average Total Cost (ATC)')
    ax.plot(quantity, mc, 'r--', linewidth=2.5, label='Marginal Cost (MC)')
    ax.plot(quantity, avc, 'g-.', linewidth=2, label='Average Variable Cost (AVC)')

    ax.set_xlabel('Quantity (Thousand Gross/Month)', fontsize=10, fontweight='bold')
    ax.set_ylabel('Cost per Piece (Rs.)', fontsize=10, fontweight='bold')
    ax.set_title('Cost Curves in Button Manufacturing', fontsize=12, fontweight='bold', pad=10)
    ax.legend(loc='upper right', fontsize=9, framealpha=0.9)
    ax.set_xlim(50, 500)
    ax.set_ylim(0, 12)
    ax.tick_params(axis='both', labelsize=9)

    # Add optimal scale annotation
    ax.annotate('Optimal Scale', xy=(300, 3.2), xytext=(380, 5.5),
                fontsize=9, ha='center',
                arrowprops=dict(arrowstyle='->', color='darkblue', lw=1.5),
                bbox=dict(boxstyle='round,pad=0.3', facecolor='lightyellow', edgecolor='gray'))

    plt.tight_layout()
    img_buffer = BytesIO()
    plt.savefig(img_buffer, format='png', dpi=180, bbox_inches='tight', facecolor='white')
    plt.close()
    img_buffer.seek(0)
    return img_buffer

def create_demand_curve_graph():
    """Create clean demand and supply curve for buttons"""
    plt.style.use('seaborn-v0_8-whitegrid')
    fig, ax = plt.subplots(figsize=(7, 4.5))

    # Inelastic demand curve (steep)
    price = np.linspace(1.5, 4.5, 100)
    quantity = 500 - 50*price

    ax.plot(quantity, price, 'b-', linewidth=2.5, label='Demand Curve (D)')

    # Add supply curve
    quantity_s = np.linspace(200, 400, 100)
    price_s = 1.0 + 0.008*quantity_s
    ax.plot(quantity_s, price_s, 'r-', linewidth=2.5, label='Supply Curve (S)')

    # Equilibrium point
    ax.plot(325, 2.8, 'ko', markersize=10, zorder=5)
    ax.annotate('Equilibrium (E)', xy=(325, 2.8), xytext=(380, 3.6),
                fontsize=9, ha='center',
                arrowprops=dict(arrowstyle='->', color='black', lw=1.5),
                bbox=dict(boxstyle='round,pad=0.3', facecolor='lightyellow', edgecolor='gray'))

    ax.set_xlabel('Quantity (Million Units)', fontsize=10, fontweight='bold')
    ax.set_ylabel('Price per Button (Rs.)', fontsize=10, fontweight='bold')
    ax.set_title('Demand and Supply Curves for Buttons', fontsize=12, fontweight='bold', pad=10)
    ax.legend(loc='upper right', fontsize=9, framealpha=0.9)
    ax.tick_params(axis='both', labelsize=9)

    plt.tight_layout()
    img_buffer = BytesIO()
    plt.savefig(img_buffer, format='png', dpi=180, bbox_inches='tight', facecolor='white')
    plt.close()
    img_buffer.seek(0)
    return img_buffer

def create_elasticity_graph():
    """Create clean price elasticity visualization"""
    plt.style.use('seaborn-v0_8-whitegrid')
    fig, ax = plt.subplots(figsize=(7, 4.5))

    # Inelastic demand (buttons)
    price = np.linspace(2, 4, 50)
    q_inelastic = 400 - 30*price
    q_elastic = 600 - 120*price

    ax.plot(q_inelastic, price, 'b-', linewidth=2.5, label='Buttons (Inelastic, PED=-0.4)')
    ax.plot(q_elastic, price, 'g--', linewidth=2, label='Velcro (More Elastic)')

    # Show price change effect
    ax.axhline(y=2.5, color='gray', linestyle=':', alpha=0.5, linewidth=1)
    ax.axhline(y=3.5, color='gray', linestyle=':', alpha=0.5, linewidth=1)
    ax.fill_between([295, 325], 2.5, 3.5, alpha=0.15, color='blue')

    ax.annotate('Small Qty Change\n(Inelastic Response)', xy=(310, 3.0), xytext=(380, 3.3),
                fontsize=8, ha='center',
                arrowprops=dict(arrowstyle='->', color='darkblue', lw=1),
                bbox=dict(boxstyle='round,pad=0.2', facecolor='lightblue', edgecolor='gray', alpha=0.8))

    ax.set_xlabel('Quantity Demanded', fontsize=10, fontweight='bold')
    ax.set_ylabel('Price (Rs.)', fontsize=10, fontweight='bold')
    ax.set_title('Price Elasticity: Buttons vs Substitutes', fontsize=12, fontweight='bold', pad=10)
    ax.legend(loc='upper right', fontsize=9, framealpha=0.9)
    ax.tick_params(axis='both', labelsize=9)

    plt.tight_layout()
    img_buffer = BytesIO()
    plt.savefig(img_buffer, format='png', dpi=180, bbox_inches='tight', facecolor='white')
    plt.close()
    img_buffer.seek(0)
    return img_buffer

def create_supply_chain_flowchart():
    """Create clean supply chain flowchart"""
    fig, ax = plt.subplots(figsize=(9, 5.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 7)
    ax.axis('off')

    # Title
    ax.text(5, 6.5, 'BUTTON SUPPLY CHAIN ARCHITECTURE', fontsize=13, fontweight='bold',
            ha='center', va='center', color='#1F4E79')

    # Backend boxes
    backend_boxes = [
        (1.5, 4.8, 'Raw Material\nSuppliers', '#E6F3FF'),
        (1.5, 3.2, 'Button\nManufacturers', '#CCE5FF'),
        (1.5, 1.6, 'Wholesalers', '#B3D9FF'),
    ]

    # Frontend boxes
    frontend_boxes = [
        (8.5, 4.8, 'Garment\nManufacturers', '#E6FFE6'),
        (8.5, 3.2, 'Apparel\nRetailers', '#CCFFCC'),
        (8.5, 1.6, 'End\nConsumers', '#B3FFB3'),
    ]

    # Section labels
    ax.text(1.5, 5.8, 'BACKEND', fontsize=11, fontweight='bold', ha='center', color='#1F4E79')
    ax.text(8.5, 5.8, 'FRONTEND', fontsize=11, fontweight='bold', ha='center', color='#2E7D32')

    # Draw boxes
    for x, y, text, color in backend_boxes + frontend_boxes:
        rect = mpatches.FancyBboxPatch((x-0.8, y-0.5), 1.6, 1.0,
                                        boxstyle="round,pad=0.05,rounding_size=0.1",
                                        facecolor=color, edgecolor='#333333', linewidth=1.5)
        ax.add_patch(rect)
        ax.text(x, y, text, ha='center', va='center', fontsize=9, fontweight='bold')

    # Backend arrows
    for i in range(len(backend_boxes)-1):
        ax.annotate('', xy=(1.5, backend_boxes[i+1][1]+0.55), xytext=(1.5, backend_boxes[i][1]-0.55),
                    arrowprops=dict(arrowstyle='->', color='#1F4E79', lw=2))

    # Frontend arrows
    for i in range(len(frontend_boxes)-1):
        ax.annotate('', xy=(8.5, frontend_boxes[i+1][1]+0.55), xytext=(8.5, frontend_boxes[i][1]-0.55),
                    arrowprops=dict(arrowstyle='->', color='#2E7D32', lw=2))

    # Direct B2B arrow (curved)
    ax.annotate('', xy=(7.6, 4.8), xytext=(2.4, 3.2),
                arrowprops=dict(arrowstyle='->', color='#D32F2F', lw=2,
                               connectionstyle="arc3,rad=-0.15"))
    ax.text(5, 4.6, 'Direct B2B (65-70%)', fontsize=8, ha='center', color='#D32F2F', fontweight='bold')

    # Wholesale channel arrow
    ax.annotate('', xy=(7.6, 4.5), xytext=(2.4, 1.6),
                arrowprops=dict(arrowstyle='->', color='#7B1FA2', lw=1.5, linestyle='--',
                               connectionstyle="arc3,rad=0.2"))
    ax.text(5, 2.2, 'Wholesale (20-25%)', fontsize=8, ha='center', color='#7B1FA2', fontweight='bold')

    plt.tight_layout()
    img_buffer = BytesIO()
    plt.savefig(img_buffer, format='png', dpi=180, bbox_inches='tight', facecolor='white', edgecolor='none')
    plt.close()
    img_buffer.seek(0)
    return img_buffer

def create_revenue_pie_chart():
    """Create clean revenue distribution pie chart"""
    plt.style.use('seaborn-v0_8-whitegrid')
    fig, ax = plt.subplots(figsize=(7, 5))

    sizes = [82, 10, 5, 3]
    labels = ['Bulk Sales to\nGarment Makers\n(82%)',
              'Wholesale\nDistribution\n(10%)',
              'Retail/DIY\n(5%)',
              'E-commerce\n(3%)']
    colors = ['#1F4E79', '#2E7D32', '#F57C00', '#C62828']
    explode = (0.03, 0, 0, 0)

    wedges, texts = ax.pie(sizes, explode=explode, labels=labels, colors=colors,
                           startangle=90, labeldistance=1.15,
                           textprops={'fontsize': 9, 'fontweight': 'bold'})

    ax.set_title('Revenue Distribution by Sales Channel', fontsize=12, fontweight='bold', pad=15)

    plt.tight_layout()
    img_buffer = BytesIO()
    plt.savefig(img_buffer, format='png', dpi=180, bbox_inches='tight', facecolor='white')
    plt.close()
    img_buffer.seek(0)
    return img_buffer

def add_cover_page(doc):
    """Add cover page - Exact Nakul's style from screenshot"""
    # Add space at top (less space - matching screenshot)
    for _ in range(2):
        doc.add_paragraph()

    # Main title - "MICROECONOMICS PRODUCT ASSIGNMENT"
    p = doc.add_paragraph()
    run = p.add_run("MICROECONOMICS PRODUCT ASSIGNMENT")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Space before product name
    for _ in range(4):
        doc.add_paragraph()

    # Product name - Large and bold "BUTTON"
    p = doc.add_paragraph()
    run = p.add_run("BUTTON")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(26)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Space before subtitle
    for _ in range(2):
        doc.add_paragraph()

    # Subtitle - Italic (exactly as in screenshot)
    p = doc.add_paragraph()
    run = p.add_run("A Microeconomic Analysis of Production, Cost, Demand, and Market Dynamics")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Large space before student details
    for _ in range(8):
        doc.add_paragraph()

    # Student details - Bold labels, regular values (Nakul's exact style)
    details = [
        ("Name:", "Astha Agrawal"),
        ("Roll No:", "HBM25002"),
        ("Class:", "MMS Batch 2025-27"),
        ("Semester:", "01st"),
        ("Academic Year:", "2025-2026"),
        ("Institute:", "St. Xavier's College (Autonomous),"),
    ]

    for label, value in details:
        p = doc.add_paragraph()
        run1 = p.add_run(label + " ")
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(12)
        run1.bold = True
        run2 = p.add_run(value)
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Institute continuation (Mumbai on same line style)
    p = doc.add_paragraph()
    run = p.add_run("Mumbai")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _ in range(3):
        doc.add_paragraph()

    # Subject line
    p = doc.add_paragraph()
    run = p.add_run("Subject: Managerial Economics")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run("Faculty: Prof. Agnelo Menezes")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

def add_introduction_section(doc):
    """Add Introduction and Evolution section - As per assignment instructions"""
    add_heading_style(doc, "1. INTRODUCTION AND EVOLUTION", level=1)

    # Definition
    add_subheading(doc, "1.1 Definition of Button as an Economic Product")

    add_body_text(doc,
        "A button, in economic terms, represents a fundamental fastening mechanism employed "
        "in the garment and apparel industry, serving both functional and aesthetic purposes. "
        "From a microeconomic perspective, buttons constitute a derived demand product, meaning "
        "their demand is not autonomous but directly proportional to the demand for finished "
        "apparel. Despite their diminutive scale, typically ranging from 8 to 100 millimeters "
        "in diameter, buttons constitute a substantial market with complex production economics "
        "and intricate supply chain dynamics.")

    add_body_text(doc,
        "The global buttons for clothing market was valued at USD 2,126.29 million "
        "in 2024, with projections to reach USD 2,923.21 million by 2033, reflecting a compound "
        "annual growth rate (CAGR) of 3.6%. This market exists in direct correlation with "
        "apparel industry expansion, which was valued at USD 1.77 trillion in 2024 and is "
        "anticipated to reach USD 2.26 trillion by 2030, expanding at a CAGR of 4.2%. Global "
        "button production reached approximately 5,000 million units in 2024, compared to "
        "4,200 million units in 2019, demonstrating consistent demand growth proportional to "
        "apparel consumption patterns across global markets.")

    # Historical Evolution
    add_subheading(doc, "1.2 Historical Evolution of Buttons")

    add_body_text(doc,
        "The button industry's development reflects remarkable technological progression and "
        "material innovation spanning over a century. India's Mehsi region in East Champaran "
        "district established the world's first organized button manufacturing cluster in 1905, "
        "when sub-inspector Bhulawan Lal initiated handcrafted pearl button production using "
        "oyster shells from the Sikrahna river. By World War II, Mehsi housed 160 factories "
        "employing 10,000 skilled artisans directly, with output reaching approximately 24 lakh "
        "gross annually. The subsequent introduction of synthetic plastic buttons made from "
        "polyester and nylon in the 1950s-1960s fundamentally disrupted this traditional "
        "industry, demonstrating microeconomic principles of technological substitution and "
        "comparative advantage.")

    # Types of buttons
    add_subheading(doc, "1.3 Types of Buttons")

    add_body_text(doc,
        "Contemporary button manufacturing encompasses diverse material categories, each "
        "serving specific market segments and price points:")

    # Types table
    headers = ["Button Type", "Material Composition", "Primary Application", "Market Share"]
    data = [
        ["Polyester Buttons", "Unsaturated polyester resin", "Mass-market shirts, casual wear", "~50%"],
        ["ABS Plastic", "Acrylonitrile-Butadiene-Styrene", "Premium garments, uniforms", "~20%"],
        ["Metal Buttons", "Brass, steel, zinc alloy", "Denim, outerwear, formal wear", "~15%"],
        ["Natural Shell", "Mother-of-pearl, coconut shell", "Luxury tailoring, ethnic wear", "~8%"],
        ["Wooden Buttons", "Hardwoods with resin composite", "Eco-conscious brands", "~5%"],
        ["Snap Buttons", "Metal with fabric covering", "Jackets, children's wear", "~2%"],
    ]
    add_elegant_table(doc, headers, data, "Table 1.1: Classification of Buttons by Material Type")

    # Product Classification
    add_subheading(doc, "1.4 Classification: Evergreen Product")

    add_body_text(doc,
        "Buttons are classified as an Evergreen product within the product lifecycle framework. "
        "Unlike sunrise products representing emerging technologies or sunset products facing "
        "declining demand, buttons demonstrate continuous, structural demand stemming from the "
        "fundamental need for garment fastening across all apparel categories. Global apparel "
        "consumption exhibits consistent growth at 3-4% annually, directly translating to stable "
        "button demand without the cyclical volatility typical of fashion accessories or seasonal "
        "products.",
        bold_words=["Evergreen product"])

    add_body_text(doc,
        "The evergreen nature is evidenced by the fact that despite technological "
        "advances in alternative fasteners like velcro, zippers, and magnetic closures, buttons "
        "remain irreplaceable for formal wear, dress shirts, and traditional garments, ensuring "
        "sustained demand across economic cycles. This classification is further supported by "
        "the observation that button demand has remained resilient even during economic "
        "downturns, with only 2-3% volume contractions during recessions compared to 15-20% "
        "declines experienced by discretionary fashion accessories.")

    # Technological Evolution
    add_subheading(doc, "1.5 Technological Evolution in Button Manufacturing")

    add_body_text(doc,
        "The technological trajectory of button manufacturing represents a fascinating case "
        "study in industrial economics. The transition from artisanal craft production to "
        "capital-intensive, technology-driven manufacturing has fundamentally transformed "
        "industry economics. Manual production in 1905 yielded 50-100 pieces per artisan daily, "
        "whereas modern automated injection molding facilities produce over 20,000 gross "
        "(2.88 million pieces) daily with minimal labor input.")

    add_body_text(doc,
        "Contemporary button manufacturing remains dominated by injection molding technology, "
        "with Italian firms Bonnetti and Giusi dominating global machinery supply. Advanced "
        "machinery incorporates IoT integration, real-time quality monitoring, and modular "
        "design allowing rapid mold changeover, which is critical for short-run customization "
        "demands in the premium segment. The button-making machines market valued at USD 150 "
        "million in 2024 is projected to reach USD 250 million by 2033, growing at 6.5% annually.")

    doc.add_page_break()

def add_backend_section(doc):
    """Add Backend Analysis section - As per assignment instructions"""
    add_heading_style(doc, "2. BACKEND ANALYSIS", level=1)

    # 2.1 Production Modality
    add_subheading(doc, "2.1 Production Modality")

    add_body_text(doc,
        "Button manufacturing employs sophisticated production technologies that have evolved "
        "significantly over the past decades. Injection molding constitutes the predominant "
        "manufacturing methodology for plastic and polyester buttons, accounting for approximately "
        "80-85% of global production. The process unfolds through four sequential stages:")

    add_bullet_point(doc, "Clamping: The mold (typically steel for durability or aluminum for "
                    "lower-volume production) is securely closed under hydraulic pressure.")
    add_bullet_point(doc, "Injection: Melted plastic maintained at material-specific temperatures "
                    "(ABS at 200-240°C, polypropylene at 160-180°C) is injected under high pressure "
                    "into the mold cavity.")
    add_bullet_point(doc, "Cooling: The plastic solidifies within the mold, with cycle times "
                    "varying from 20-60 seconds depending on material properties and button geometry.")
    add_bullet_point(doc, "Ejection: Finished buttons are extracted using ejection pins, typically "
                    "requiring draft angles of 1-2 degrees to facilitate separation.")

    add_body_text(doc,
        "This manufacturing paradigm demonstrates substantial economies of scale. Multi-cavity "
        "molds enable simultaneous production of 4-16 buttons per cycle, dramatically reducing "
        "per-unit production costs as output volume increases. The button-making machines market "
        "was valued at USD 150 million in 2024 and is projected to reach USD 250 million by 2033, "
        "growing at 6.5% annually, reflecting industry-wide automation intensification.")

    # Technology evolution table
    headers = ["Era", "Technology", "Output Capacity", "Labor Requirement"]
    data = [
        ["Pre-1950s", "Hand-crafted (shell/bone)", "50-100 pieces/day", "High (10+ workers)"],
        ["1950s-1980s", "Semi-automatic molding", "5,000 gross/day", "Medium (5-8 workers)"],
        ["1980s-2010s", "Automatic injection molding", "15,000 gross/day", "Low (3-5 workers)"],
        ["2010s-Present", "IoT-enabled automation", "20,000+ gross/day", "Minimal (2-3 operators)"],
    ]
    add_elegant_table(doc, headers, data, "Table 2.1: Evolution of Button Manufacturing Technology")

    # 2.2 Raw Materials
    add_subheading(doc, "2.2 Raw Materials and Factor Market")

    add_body_text(doc,
        "The raw material landscape for button manufacturing reflects geographic cost advantages "
        "and material specialization. Polyester resin, derived from petroleum derivatives, remains "
        "the dominant input material accounting for approximately 50% of button production globally. "
        "India's strategic advantage in button manufacturing reflects low-cost polyester resin "
        "access through domestic petrochemical capacity from companies like Reliance Industries "
        "and GAIL.")

    # Raw materials table
    headers = ["Material", "Source Region", "Cost (Rs./kg)", "Application"]
    data = [
        ["Polyester Resin", "India, China, Middle East", "80-120", "Standard buttons"],
        ["ABS Resin", "Taiwan, South Korea, Germany", "150-200", "Premium buttons"],
        ["Brass/Steel", "India, Southeast Asia", "250-400", "Metal buttons"],
        ["Mother-of-Pearl", "Indonesia, Philippines", "800-1,500", "Luxury segment"],
        ["Recycled PET", "India, Europe", "60-90", "Eco-friendly buttons"],
    ]
    add_elegant_table(doc, headers, data, "Table 2.2: Raw Material Sources and Costs")

    add_body_text(doc,
        "The labor market for button manufacturing exhibits relatively low skill-intensity compared "
        "to apparel assembly. Mold technicians and toolmakers represent high-skill personnel "
        "commanding Rs. 25,000-50,000 monthly in Mumbai and Delhi. Machine operators, classified "
        "as semi-skilled workers, earn Rs. 12,000-18,000 monthly. Quality inspectors receive "
        "Rs. 15,000-25,000 monthly, while logistics and packaging workers earn Rs. 8,000-12,000 "
        "monthly at the unskilled level.")

    add_body_text(doc,
        "Capital markets and machinery finance play crucial roles in industry structure. "
        "Injection molding machine capital requirements range from Rs. 8-15 lakhs for "
        "semi-automatic units to Rs. 25-40 lakhs for fully automatic systems. Mold fabrication "
        "typically costs Rs. 50,000-2,00,000 depending on complexity. Total project capitalization "
        "for a small button manufacturing unit ranges from Rs. 16.03 million, with Internal Rate "
        "of Return (IRR) of 42% under efficient operations.")

    # 2.3 Cost Structure
    add_subheading(doc, "2.3 Cost Structure Analysis")

    add_body_text(doc,
        "The cost structure of button manufacturing demonstrates classical microeconomic principles "
        "of fixed versus variable costs and economies of scale. For a typical 5,000 gross per-day "
        "button manufacturing facility, the capital expenditure breakdown is as follows:")

    # CAPEX table
    headers = ["Component", "Cost (Rs. Lakhs)", "Percentage"]
    data = [
        ["Injection molding machines (3-4 units)", "30-40", "40-45%"],
        ["Molds and tooling", "8-12", "10-15%"],
        ["Finishing and trimming machines", "5-8", "7-10%"],
        ["Quality control equipment", "2-3", "3-5%"],
        ["Infrastructure (building, utilities)", "12-15", "15-18%"],
        ["Working capital (3 months)", "8-10", "10-12%"],
        ["Total CAPEX", "65-88", "100%"],
    ]
    add_elegant_table(doc, headers, data, "Table 2.3: Capital Expenditure (CAPEX) Breakdown")

    # OPEX table
    headers = ["Cost Category", "Monthly (Rs. Lakhs)", "Per Piece (Paisa)", "Percentage"]
    data = [
        ["Raw material (polyester resin)", "2.5-3.2", "1.67-2.13", "40-45%"],
        ["Labor (operators, QC)", "1.2-1.5", "0.80-1.00", "15-18%"],
        ["Utilities (electricity, water)", "0.6-0.8", "0.40-0.53", "7-10%"],
        ["Packaging materials", "0.5-0.7", "0.33-0.47", "6-8%"],
        ["Maintenance and repairs", "0.3-0.4", "0.20-0.27", "4-5%"],
        ["Administrative overheads", "0.4-0.6", "0.27-0.40", "5-7%"],
        ["Total Monthly OPEX", "5.5-7.2", "3.67-4.80", "100%"],
    ]
    add_elegant_table(doc, headers, data, "Table 2.4: Operating Expenditure (OPEX) Structure")

    add_body_text(doc,
        "The fixed costs comprising factory rent, insurance, administrative salaries, and "
        "depreciation amount to approximately Rs. 0.8-1.2 lakhs monthly, representing 15-18% "
        "of total OPEX. Variable costs including raw materials, direct labor, utilities, and "
        "packaging constitute Rs. 4.7-6.0 lakhs monthly, representing 82-85% of OPEX. This cost "
        "structure demonstrates button manufacturing's leverage to volume: at 50% capacity "
        "utilization, average cost reaches Rs. 5.5-6.5 per gross, but at 85% capacity, average "
        "cost declines to Rs. 3.8-4.2 per gross, reflecting strong economies of scale.",
        bold_words=["economies of scale"])

    # Add cost curve graph
    add_body_text(doc, "The following figure illustrates the behavior of Average Total Cost (ATC), "
                 "Marginal Cost (MC), and Average Variable Cost (AVC) curves in button manufacturing:")

    cost_graph = create_cost_curve_graph()
    doc.add_picture(cost_graph, width=Inches(5.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run("Figure 2.1: Cost Curves in Button Manufacturing")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2.4 Supply Chain
    add_subheading(doc, "2.4 Supply Chain Analysis")

    add_body_text(doc,
        "The button supply chain exhibits a three-tier architecture connecting raw material "
        "suppliers to end consumers through multiple intermediaries. The backend supply chain "
        "comprises raw material suppliers (petrochemical companies, metal suppliers, natural "
        "material collectors), button manufacturers, and wholesalers/distributors. The frontend "
        "supply chain connects wholesalers to garment manufacturers, apparel retailers, and "
        "ultimately end consumers.")

    # Add supply chain flowchart
    supply_chain_graph = create_supply_chain_flowchart()
    doc.add_picture(supply_chain_graph, width=Inches(5.2))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run("Figure 2.2: Button Supply Chain Architecture")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_body_text(doc,
        "In Mumbai, wholesale button markets are concentrated in areas such as Mulji Jetha Market "
        "in Kalbadevi, Crawford Market, and Mangaldas Market. These markets serve as critical "
        "aggregation points where garment manufacturers can source diverse button varieties from "
        "multiple suppliers. Typical wholesaler margins range from 8-12% between manufacturer "
        "wholesale price (Rs. 2.50-3.00 per piece) and retail price to garment makers "
        "(Rs. 2.75-3.50 per piece).")

    add_body_text(doc,
        "Large Indian button manufacturers exhibit partial backward integration. Bhansali Buttons, "
        "operating at 20,000 gross per day capacity, maintains in-house resin processing and "
        "quality labs, reducing input costs by 8-15% versus purchasing pre-processed materials.")

    # 2.5 Government Policies
    add_subheading(doc, "2.5 Government Policies")

    add_body_text(doc,
        "The regulatory framework governing button manufacturing in India encompasses taxation, "
        "MSME support, and environmental compliance:")

    add_body_text(doc,
        "Goods and Services Tax (GST): Buttons (HS Code 9606) are subject to 18% GST. "
        "Manufacturers can claim Input Tax Credit (ITC) on inputs including resin, packaging, "
        "and utilities, enabling net tax neutrality on inter-state B2B sales. Small manufacturers "
        "with turnover below Rs. 1.5 crore can opt for the Composition Scheme at 1% fixed GST "
        "without ITC benefit.",
        bold_words=["18% GST"])

    add_body_text(doc,
        "MSME Support Schemes: The Credit-linked Capital Subsidy Scheme (CLCSS) provides "
        "subsidized interest on capital loans for manufacturing equipment. The Prime Minister's "
        "Employment Generation Programme (PMEGP) offers subsidies up to Rs. 25 lakhs for "
        "micro-enterprises.")

    add_body_text(doc,
        "Environmental Regulations: Under the Plastic Waste Management Rules (2016), Extended "
        "Producer Responsibility (EPR) Authorization is mandatory for plastic button manufacturers. "
        "Manufacturers must ensure recycling of 30% of plastic waste generated within 5 years.")

    doc.add_page_break()

def add_frontend_section(doc):
    """Add Frontend Analysis section - As per assignment instructions"""
    add_heading_style(doc, "3. FRONTEND ANALYSIS", level=1)

    # 3.1 Demand Analysis
    add_subheading(doc, "3.1 Demand Analysis")

    add_body_text(doc,
        "Button demand exemplifies derived demand in industrial economics. The demand for buttons "
        "is not autonomous but flows directly from demand for finished apparel. The mathematical "
        "relationship can be expressed as:",
        bold_words=["derived demand"])

    p = doc.add_paragraph()
    run = p.add_run("Demand for Buttons = f(Demand for Apparel, Buttons per Garment, Substitution Rate)")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)

    add_body_text(doc,
        "The elasticity of button demand to apparel demand demonstrates near-unity elastic response: "
        "a 1% increase in apparel production yields approximately 0.95-1.05% increase in button "
        "demand. Global button production reached approximately 5,000 million units in 2024.")

    add_body_text(doc, "The key determinants of button demand include:")

    add_bullet_point(doc, "Apparel Production Volume: The primary driver, with global apparel "
                    "market CAGR of 3.5-4.2% translating directly to button demand growth.")
    add_bullet_point(doc, "Garment Type Composition: Formal wear requires 4-6 buttons per garment, "
                    "while casual t-shirts require 0-2 buttons, affecting aggregate demand.")
    add_bullet_point(doc, "Consumer Preferences: Growing demand for personalized and eco-friendly "
                    "buttons reflects changing consumer preferences.")

    # Add demand curve
    demand_graph = create_demand_curve_graph()
    doc.add_picture(demand_graph, width=Inches(5.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run("Figure 3.1: Demand and Supply Curves for Buttons")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3.2 Elasticity of Demand
    add_subheading(doc, "3.2 Elasticity of Demand")

    add_body_text(doc,
        "Price Elasticity of Demand (PED): Button demand exhibits inelastic pricing characteristics "
        "with an estimated PED of -0.3 to -0.6. This inelasticity is justified by several factors:",
        bold_words=["inelastic"])

    add_bullet_point(doc, "Buttons represent less than 0.5% of total apparel manufacturing cost, "
                    "making price insignificant to garment maker purchasing decisions.")
    add_bullet_point(doc, "No close substitutes are available for basic functional buttons without "
                    "requiring garment redesign and manufacturing process changes.")
    add_bullet_point(doc, "Derived demand inherently exhibits lower price sensitivity than final "
                    "consumer demand.")

    # Elasticity table
    headers = ["Elasticity Type", "Estimated Value", "Interpretation"]
    data = [
        ["Price Elasticity (PED)", "-0.3 to -0.6", "Inelastic - small quantity change for price change"],
        ["Income Elasticity (YED)", "+0.9 to +1.1", "Unit elastic - grows with apparel consumption"],
        ["Cross Elasticity (Velcro)", "+0.45 to +0.65", "Substitute - moderate substitutability"],
        ["Cross Elasticity (Zippers)", "+0.10 to +0.20", "Weak substitute - different applications"],
    ]
    add_elegant_table(doc, headers, data, "Table 3.1: Elasticity Measures for Button Demand")

    # Add elasticity graph
    elasticity_graph = create_elasticity_graph()
    doc.add_picture(elasticity_graph, width=Inches(5.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run("Figure 3.2: Price Elasticity - Buttons vs Substitutes")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_body_text(doc,
        "Income Elasticity of Demand (YED): Button demand exhibits unit-elastic income response "
        "(estimated YED of +0.9 to +1.1). Button demand grows proportionally with real apparel "
        "consumption. Economic downturns correlate with 2-3% button volume reductions.")

    add_body_text(doc,
        "Cross-Price Elasticity of Demand: Velcro demonstrates moderate substitutability "
        "with a cross-elasticity of +0.45 to +0.65. The relatively low overall cross-elasticity "
        "confirms button demand's structural stability.")

    # 3.3 Market Structure
    add_subheading(doc, "3.3 Market Structure")

    add_body_text(doc,
        "The button market exemplifies monopolistic competition, characterized by the following "
        "features:",
        bold_words=["monopolistic competition"])

    add_bullet_point(doc, "Large number of competitors: India has 200+ registered manufacturers "
                    "and 500+ unregistered MSME operations.")
    add_bullet_point(doc, "Product differentiation: Variations in material, sizes (8-100 ligne), "
                    "colors (50+ options), and special effects.")
    add_bullet_point(doc, "Free entry and exit: Low capital requirements (Rs. 60-80 lakhs), "
                    "minimal regulatory barriers, and mature technology.")
    add_bullet_point(doc, "Non-price competition: Design innovation, sustainability certifications, "
                    "quality assurance, and logistics efficiency.")

    add_body_text(doc,
        "Market concentration analysis reveals a fragmented industry structure. The top 5 "
        "manufacturers control approximately 12-15% of the Indian market, the next 20 "
        "control 20-25%, and the remaining 250+ manufacturers control 60-65%.")

    # 3.4 Pricing Mechanism
    add_subheading(doc, "3.4 Pricing Mechanism")

    add_body_text(doc,
        "Button pricing operates through a tiered structure based on order volume, material type, "
        "and value-added features:")

    # Pricing table
    headers = ["Pricing Tier", "Order Volume", "Price/Piece (Rs.)", "Discount"]
    data = [
        ["Minimum Order Quantity", "100-1,000 gross", "3.20-3.50", "Base price"],
        ["Standard Wholesale", "1,000-5,000 gross", "2.80-3.20", "8-12%"],
        ["Bulk Wholesale", "5,000-20,000 gross", "2.40-2.80", "15-20%"],
        ["Large Contract", "20,000+ gross", "2.00-2.40", "25-30%"],
    ]
    add_elegant_table(doc, headers, data, "Table 3.2: Volume-Based Pricing Structure")

    # Channel pricing
    headers = ["Channel", "Price (Rs./piece)", "Margin Structure"]
    data = [
        ["Manufacturer to Wholesaler", "2.50-3.00", "Base cost + margin"],
        ["Wholesaler to Garment Maker", "2.75-3.50", "8-12% markup"],
        ["Retail (DIY, replacement)", "15-50", "400-1,500% markup"],
        ["E-commerce (bulk)", "2.00-2.75", "Competitive to wholesale"],
    ]
    add_elegant_table(doc, headers, data, "Table 3.3: Channel-Based Pricing")

    # 3.5 Revenue Analysis
    add_subheading(doc, "3.5 Revenue Analysis")

    add_body_text(doc,
        "Revenue streams in button manufacturing are distributed across multiple channels. "
        "Bulk sales to garment manufacturers constitute 80-85% of total revenue. For example, "
        "a manufacturer producing 150,000 gross per month at Rs. 2.70 per piece generates "
        "Rs. 58.5 lakhs in monthly revenue.")

    # Add revenue pie chart
    revenue_graph = create_revenue_pie_chart()
    doc.add_picture(revenue_graph, width=Inches(4.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run("Figure 3.3: Revenue Distribution by Sales Channel")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_body_text(doc,
        "The Total Revenue (TR) for button manufacturers is primarily determined by volume "
        "rather than price, given the inelastic demand. Average Revenue (AR) equals price in "
        "this competitive market structure. Profit margins across the value chain vary: "
        "manufacturers achieve 8-12% net margins, wholesalers earn 4-6% net margins, while "
        "retailers capture the highest percentage margins on smaller volumes.")

    # 3.6 Selling and Marketing Costs
    add_subheading(doc, "3.6 Selling and Marketing Costs")

    add_body_text(doc,
        "Button manufacturers invest minimally in B2C marketing (0.2-0.3% of revenue) due to "
        "the derived demand nature of the product. B2B marketing focuses on trade participation, "
        "catalog design, and quality certifications.")

    # Selling costs table
    headers = ["Cost Category", "Rs./1000 pieces", "% of Revenue"]
    data = [
        ["Packaging materials", "8-12", "0.3-0.4%"],
        ["Transportation/Logistics", "20-35", "0.7-1.2%"],
        ["Sales commissions (agents)", "15-25", "0.5-0.8%"],
        ["Marketing/Trade shows", "5-10", "0.2-0.3%"],
        ["Quality certifications", "10-15", "0.3-0.5%"],
        ["Total Selling Costs", "58-97", "2.0-3.2%"],
    ]
    add_elegant_table(doc, headers, data, "Table 3.4: Selling and Marketing Cost Structure")

    doc.add_page_break()

def add_suggestions_section(doc):
    """Add Suggestions section - As per assignment instructions"""
    add_heading_style(doc, "4. SUGGESTIONS TO IMPROVE PRODUCTION AND DELIVERY EFFICIENCY", level=1)

    add_body_text(doc,
        "Based on the microeconomic analysis conducted, the following recommendations are proposed "
        "to enhance production efficiency, reduce costs, and improve supply chain effectiveness:")

    add_subheading(doc, "4.1 Cost Minimization Strategies")

    add_bullet_point(doc, "Adopt multi-cavity mold technology: Transitioning from 8-cavity to "
                    "16-cavity molds can halve per-button production time and reduce average "
                    "costs by 15-20%.")
    add_bullet_point(doc, "Implement hot runner mold systems: These systems reduce material "
                    "waste from 8-10% to 2-3%, with cost premiums amortized over 100 million+ pieces.")
    add_bullet_point(doc, "Negotiate bulk procurement contracts: Large manufacturers achieve "
                    "15-20% discounts on polyester resin purchases versus MSMEs.")

    add_subheading(doc, "4.2 Supply Chain Efficiency Improvements")

    add_bullet_point(doc, "Consolidate shipments using 40-foot containers: This reduces per-unit "
                    "logistics cost by 25-30% versus Less Container Load (LCL) shipments.")
    add_bullet_point(doc, "Establish regional distribution centers: Reducing delivery times to "
                    "garment manufacturers improves service levels and reduces inventory costs.")
    add_bullet_point(doc, "Implement Just-In-Time (JIT) inventory management: Coordinating "
                    "production with demand reduces working capital requirements.")

    add_subheading(doc, "4.3 Sustainable Materials and Recycling")

    add_bullet_point(doc, "Invest in recycled PET button production: Growing 8-12% annually, "
                    "this segment commands only 10-20% cost premium while addressing ESG requirements.")
    add_bullet_point(doc, "Develop bamboo button capacity: With only 5-10% cost premium and "
                    "superior sustainability narrative, bamboo buttons serve eco-conscious brands.")
    add_bullet_point(doc, "Implement Extended Producer Responsibility (EPR) compliance systems.")

    add_subheading(doc, "4.4 Technology Upgradation")

    add_bullet_point(doc, "Deploy AI-enabled quality control: This reduces defect rates from "
                    "2-3% to 0.5-1%, improving customer satisfaction.")
    add_bullet_point(doc, "Integrate IoT-enabled condition monitoring: Reduces unplanned "
                    "downtime by 40% through predictive maintenance.")
    add_bullet_point(doc, "Adopt energy-efficient cold-runner injection molding: Reduces energy "
                    "consumption by 20-25% versus conventional heating methods.")

    doc.add_page_break()

def add_lessons_learned_section(doc):
    """Add Lessons Learned section - As per assignment instructions (first person)"""
    add_heading_style(doc, "5. LESSONS LEARNED", level=1)

    add_body_text(doc,
        "Through this comprehensive microeconomic analysis of the button industry, I have gained "
        "valuable insights into the practical application of economic theory to real-world "
        "manufacturing scenarios. The following reflections capture my key learnings:")

    add_subheading(doc, "5.1 Understanding Derived Demand")

    add_body_text(doc,
        "I learned that buttons represent a classic example of derived demand, where the demand "
        "for an input factor is determined by the demand for the final product it helps create. "
        "This assignment helped me understand that button manufacturers must closely monitor "
        "apparel industry trends rather than direct consumer preferences. The near-unity elasticity "
        "of button demand to apparel demand (0.95-1.05) demonstrates how intermediate goods "
        "markets are intrinsically linked to final goods markets.")

    add_subheading(doc, "5.2 Appreciating Cost Structure Dynamics")

    add_body_text(doc,
        "I gained practical understanding of how fixed and variable costs interact to determine "
        "average total cost behavior. The observation that button manufacturing ATC drops from "
        "Rs. 8.50-10.20 per piece at low volumes to Rs. 2.80-3.50 per piece at high volumes "
        "illustrated economies of scale more vividly than any theoretical exposition. I now "
        "appreciate why manufacturing industries naturally tend toward consolidation.")

    add_subheading(doc, "5.3 Recognizing Market Structure Characteristics")

    add_body_text(doc,
        "I understood how monopolistic competition manifests in practice through the button "
        "industry's structure. The presence of 250+ manufacturers with product differentiation "
        "through materials, sizes, colors, and special features, combined with free entry/exit "
        "and non-price competition, perfectly illustrates Chamberlin's model of monopolistic "
        "competition.")

    add_subheading(doc, "5.4 Applying Price Elasticity Concepts")

    add_body_text(doc,
        "I learned to identify factors that determine price elasticity in industrial markets. "
        "The inelastic demand for buttons (PED of -0.3 to -0.6) results from the product's "
        "small share of total production cost, absence of close substitutes, and derived demand "
        "characteristics. This taught me that B2B markets often exhibit different elasticity "
        "patterns than consumer markets.")

    add_subheading(doc, "5.5 Understanding Supply Chain Economics")

    add_body_text(doc,
        "I appreciated how supply chain structure affects cost distribution and margin allocation "
        "across stakeholders. The observation that wholesalers earn 8-12% margins while retailers "
        "achieve 400-1,500% markups on identical products taught me about value addition through "
        "convenience, risk absorption, and information provision.")

    add_subheading(doc, "5.6 Overall Reflection")

    add_body_text(doc,
        "This comprehensive microeconomic analysis has provided invaluable insights into how "
        "theoretical economic concepts manifest in real-world industrial contexts. The button "
        "industry, despite its seemingly simple product offering, embodies complex microeconomic "
        "dynamics spanning factor markets, production economics, market structure theory, and "
        "price theory.")

    doc.add_page_break()

    # Conclusion
    add_heading_style(doc, "CONCLUSION", level=1)

    add_body_text(doc,
        "The button industry exemplifies elegant microeconomic principles where derived demand, "
        "economies of scale, monopolistic competition, and technological disruption converge "
        "within a seemingly simple product category. While individual buttons command negligible "
        "unit value of Rs. 2-5 per piece, the aggregate market exceeding USD 2 billion reflects "
        "structural demand derived from a USD 1.7+ trillion apparel industry.")

    add_body_text(doc,
        "India's prominence in button manufacturing, serving as a cost-competitive global "
        "supplier while sustaining legacy artisanal production in regions like Mehsi, illustrates "
        "dynamic comparative advantage in labor-intensive manufacturing complemented by technology "
        "adoption. The sector's future depends on navigating sustainability imperatives while "
        "leveraging technological innovation to sustain margins in competitive markets.")

    add_body_text(doc,
        "The button industry's evolution from cottage industry through the polyester revolution "
        "to sustainable innovation demonstrates how microeconomic forces reshape industrial "
        "structure and competitive advantage. This assignment has provided comprehensive "
        "understanding of how microeconomic theory applies to real-world manufacturing contexts.")

    doc.add_page_break()

def add_references_section(doc):
    """Add References section"""
    add_heading_style(doc, "REFERENCES", level=1)

    references = [
        "Market Reports World. (2024). Buttons for Clothing Market Size, Share & Trends Analysis Report.",
        "Grand View Research. (2024). Apparel Market Size, Share & Trends Analysis Report.",
        "East Champaran District Administration. (2024). Mehsi Button Industry - Historical Overview.",
        "Verified Market Reports. (2024). Automatic Button Making Machines Market Report.",
        "Apparel Resources. (2024). India's Bhansali Buttons: Innovation and Quality in Export Markets.",
        "Chetna International. (2024). Quality vs Cost: Buttons in Bulk Procurement.",
        "Market Research Future. (2024). India Apparel Market Size, Share & Industry Analysis.",
        "ClearTax. (2024). GST Rates for Buttons, Press Fasteners - HSN Code 9606.",
        "GST Lawyer. (2024). GST Provisions for the Textile Industry in India.",
        "Corpbiz. (2024). Environmental Compliance for Plastic Button Manufacturers.",
        "Cotton Monk. (2024). Sustainable Buttons: Eco-Friendly Fasteners Guide.",
        "OEC World. (2024). Global Button Trade Statistics and Analysis.",
        "IBEF. (2024). Indian Apparel Industry Analysis and Export Data.",
        "Scribd. (2020). Plastic Button Manufacturing Project Report.",
        "Gabe Clothing. (2024). Clothing Manufacturing Costs Breakdown.",
        "Investopedia. (2024). Derived Demand - Economic Concept Explanation.",
        "6W Research. (2024). India Buttons for Clothing Market Outlook.",
    ]

    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"[{i}] {ref}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

def main():
    """Main function to create the complete document"""
    print("Creating Button Microeconomics Assignment (Nakul Style)...")

    # Create document
    doc = create_document()

    # Add cover page
    print("  Adding cover page...")
    add_cover_page(doc)

    # Add main sections
    print("  Adding Introduction section...")
    add_introduction_section(doc)

    print("  Adding Backend Analysis section...")
    add_backend_section(doc)

    print("  Adding Frontend Analysis section...")
    add_frontend_section(doc)

    print("  Adding Suggestions section...")
    add_suggestions_section(doc)

    print("  Adding Lessons Learned section...")
    add_lessons_learned_section(doc)

    print("  Adding References...")
    add_references_section(doc)

    # Add page numbers
    print("  Adding page numbers...")
    add_page_number(doc)

    # Save document
    output_path = os.path.join(OUTPUT_DIR, "Astha_Agrawal_HBM25002_Button_Assignment.docx")
    doc.save(output_path)
    print(f"\nDocument saved successfully: {output_path}")

    return output_path

if __name__ == "__main__":
    main()
