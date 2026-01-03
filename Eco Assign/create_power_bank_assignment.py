#!/usr/bin/env python3
"""
Generate Power Bank Microeconomics Assignment - 20 PAGE VERSION
~5500 words, 12 tables, 20 pages, strict format compliance
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill_color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_margins(cell, top=50, bottom=50, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, margin_value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(margin_value))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def format_cell(cell, text, bold=False, size=11, center=False, header=False):
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(str(text))
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if header:
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_margins(cell)


def create_table(doc, headers, data, col_widths=None):
    num_cols = len(headers)
    num_rows = len(data) + 1
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    if col_widths:
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(width)

    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, '1F4E79')
        format_cell(cell, header_text, bold=True, size=11, center=True, header=True)

    for row_idx, row_data in enumerate(data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            if row_idx % 2 == 1:
                set_cell_shading(cell, 'F2F2F2')
            format_cell(cell, cell_text, size=11, center=(col_idx > 0))

    return table


def create_document():
    doc = Document()

    # Page setup - As per concept note: 1, 1, 1, 1.5 (Top, Right, Bottom, Left)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1)
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)

    add_page_numbers(doc)

    # ==================== COVER PAGE ====================
    for _ in range(5):
        doc.add_paragraph()

    title = doc.add_paragraph()
    run = title.add_run("MICROECONOMICS PRODUCT ASSIGNMENT")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Times New Roman'
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    product = doc.add_paragraph()
    run = product.add_run("POWER BANK")
    run.bold = True
    run.font.size = Pt(26)
    run.font.name = 'Times New Roman'
    product.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    run = subtitle.add_run("A Microeconomic Analysis of Production, Cost, Demand, and Market Dynamics")
    run.italic = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _ in range(7):
        doc.add_paragraph()

    details = [
        ("Name:", "Nakul Kundra"),
        ("Roll No:", "25066"),
        ("Class:", "MMS Batch 2025-27"),
        ("Semester:", "1st"),
        ("Academic Year:", "2025-2026"),
        ("Institute:", "Sydenham Institute of Management Studies,"),
        ("", "Research and Entrepreneurship Education, Mumbai")
    ]

    for label, value in details:
        para = doc.add_paragraph()
        if label:
            run = para.add_run(label + " ")
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
        run = para.add_run(value)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ==================== 1. INTRODUCTION ====================
    add_heading(doc, "1. INTRODUCTION AND EVOLUTION")

    add_para(doc, "A power bank, also known as a portable charger or battery pack, is a portable device designed to store electrical energy in rechargeable lithium-ion or lithium-polymer batteries. These devices serve as external power sources for charging smartphones, tablets, laptops, wireless earphones, smartwatches, and other portable electronic devices. In today's hyperconnected world, where individuals depend heavily on mobile devices for communication, work, entertainment, and navigation, power banks have evolved from luxury accessories to essential everyday items that ensure uninterrupted connectivity.")

    add_para(doc, "The global power bank market has experienced remarkable growth, reaching a valuation of USD 14.2 billion in 2024. This expansion is driven by several converging factors: the proliferation of mobile devices, increasing battery consumption due to power-hungry applications and 5G connectivity, and the growing culture of mobile-first lifestyles across both developed and emerging economies. With over 6.8 billion smartphone users worldwide and average daily usage exceeding 4.5 hours, the demand for portable charging solutions continues to accelerate at an unprecedented pace.")

    add_para(doc, "The Indian market has emerged as one of the fastest-growing segments globally, valued at USD 963.31 million in 2024 with projections indicating growth to USD 2,565.50 million by 2033, representing a compound annual growth rate of 11.5%. This growth is fueled by India's 750 million smartphone users, rapid digital adoption, increasing disposable incomes in urban areas, and the expanding e-commerce infrastructure that makes power banks accessible even in tier-2 and tier-3 cities.")

    add_para(doc, "The significance of power banks extends beyond mere convenience. For professionals, reliable device power ensures uninterrupted productivity during travel, field work, and meetings. For students, it enables continuous access to educational resources and online learning platforms. For travelers and outdoor enthusiasts, power banks provide essential connectivity in locations without electrical infrastructure. The COVID-19 pandemic further accelerated adoption as remote work, online education, and increased digital entertainment consumption drove higher smartphone usage.")

    add_table_title(doc, "Table 1: Global Power Bank Market Overview (2024)")
    create_table(doc,
        ['Metric', 'Value', 'Growth Rate'],
        [
            ('Global Market Size', 'USD 14.2 Billion', '6.8% CAGR'),
            ('India Market Size', 'USD 963.31 Million', '11.5% CAGR'),
            ('India Market (2033 Projected)', 'USD 2,565.50 Million', '-'),
            ('Global Smartphone Users', '6.8 Billion', '4.2% YoY'),
            ('India Smartphone Users', '750 Million', '8.5% YoY'),
            ('Average Daily Phone Usage', '4.5 Hours', '+12% YoY'),
        ],
        col_widths=[2.5, 1.8, 1.5])
    doc.add_paragraph()

    add_heading(doc, "1.1 Historical Evolution of Power Banks", level=2)

    add_para(doc, "The evolution of power banks closely parallels the development of portable electronics and battery technology over the past two decades. The journey began in the early 2000s when the first commercial power banks emerged, utilizing Nickel-Metal Hydride (NiMH) battery technology. These early devices offered modest energy density of 60-80 Wh/kg and capacities of 1,000-2,000 mAh, sufficient for emergency top-ups but not for complete device charges. Their bulky form factor and limited efficiency restricted adoption to technology enthusiasts and frequent business travelers.")

    add_para(doc, "The lithium-ion revolution from 2008 to 2014 marked a transformative period in power bank development. Lithium-ion cells offered significantly improved energy density of 150-200 Wh/kg, enabling higher capacities in smaller, more portable form factors. This era saw power banks achieve 500-1000 charge cycles, making them economically viable for regular daily use. Manufacturing costs decreased substantially as production scaled globally, particularly in China, where integrated supply chains emerged for battery cells, electronic components, and housing materials.")

    add_para(doc, "The fast charging era from 2014 to 2020 introduced technologies like Qualcomm Quick Charge 2.0, 3.0, and subsequently 4.0, enabling charging speeds of 18W to 65W. This addressed the significant consumer frustration with slow charging times and made power banks practical for emergency rapid charging situations. Concurrently, USB-C adoption began standardizing connectors and enabling bidirectional power flow, allowing a single port to both charge the power bank and deliver power to connected devices.")

    add_para(doc, "Contemporary power banks from 2020 to the present incorporate advanced features including wireless charging capabilities at 5-15W using Qi standard, USB-C Power Delivery supporting up to 100W output for laptop charging, and capacities exceeding 30,000 mAh. Gallium Nitride (GaN) technology has enabled more compact, efficient charging circuitry with less heat generation. Smart features like LED displays showing precise battery percentage, pass-through charging allowing simultaneous charging and discharging, and multiple output ports catering to diverse device ecosystems have become standard in premium offerings.")

    add_table_title(doc, "Table 2: Power Bank Technology Evolution")
    create_table(doc,
        ['Era', 'Period', 'Technology', 'Energy Density', 'Key Features'],
        [
            ('Early', '2001-2008', 'NiMH', '60-80 Wh/kg', '1,000-2,000 mAh'),
            ('Li-ion', '2008-2014', 'Lithium-ion', '150-200 Wh/kg', '500-1000 cycles'),
            ('Fast Charge', '2014-2020', 'QC 2.0/3.0', '200-250 Wh/kg', '18W-65W charging'),
            ('Current', '2020-Present', 'GaN/USB-PD', '250-300 Wh/kg', '100W PD, Wireless'),
        ],
        col_widths=[0.9, 1.1, 1.1, 1.2, 1.5])
    doc.add_paragraph()

    add_heading(doc, "1.2 Types and Versions of Power Banks", level=2)

    add_para(doc, "The power bank market exhibits significant product differentiation to address diverse consumer needs and use cases across different market segments. Standard power banks with capacities ranging from 5,000 to 20,000 mAh represent the largest market segment at approximately 60% market share. These devices offer reliable charging for smartphones and tablets at accessible price points of Rs. 500-2,000, making them suitable for everyday use by price-conscious consumers who need basic portable charging functionality.")

    add_para(doc, "Fast charging power banks have emerged as the fastest-growing segment with 18% year-over-year growth. Supporting protocols like Qualcomm Quick Charge 3.0/4.0, USB Power Delivery 3.0, and proprietary technologies like OnePlus Warp Charge and Xiaomi HyperCharge, these devices deliver 18W-65W output, reducing charging time by 50-70% compared to standard chargers. Premium pricing of Rs. 1,500-5,000 reflects the advanced circuitry, higher-quality cells, and sophisticated thermal management required for safe fast charging operations.")

    add_para(doc, "Wireless power banks represent an emerging category experiencing 25% annual growth, driven by the proliferation of Qi-compatible smartphones. These devices incorporate wireless charging pads on their surfaces, enabling cable-free charging for compatible devices. While currently limited to 5-15W output speeds, the convenience of eliminating cables appeals to users seeking simplified charging experiences. Solar power banks with integrated photovoltaic panels cater to outdoor enthusiasts and emergency preparedness, while high-capacity models (20,000-50,000 mAh) target power users and those charging multiple devices, with aviation regulations limiting lithium batteries to 100Wh for carry-on luggage.")

    add_table_title(doc, "Table 3: Power Bank Categories and Market Segments")
    create_table(doc,
        ['Category', 'Capacity', 'Price (Rs.)', 'Market Share', 'Growth'],
        [
            ('Standard', '5,000-20,000 mAh', '500-2,000', '60%', '8%'),
            ('Fast Charging', '10,000-20,000 mAh', '1,500-5,000', '22%', '18%'),
            ('Wireless', '10,000-15,000 mAh', '1,500-4,000', '8%', '25%'),
            ('Solar', '10,000-25,000 mAh', '1,500-6,000', '3%', '12%'),
            ('High-Capacity', '20,000-50,000 mAh', '2,500-10,000', '7%', '15%'),
        ],
        col_widths=[1.2, 1.4, 1.2, 1.0, 0.8])
    doc.add_paragraph()

    add_heading(doc, "1.3 Product Classification: Sunrise, Sunset, or Evergreen?", level=2)

    add_para(doc, "Analyzing the power bank product category through the lens of product lifecycle classification reveals characteristics of an EVERGREEN product with distinct SUNRISE attributes. This classification is supported by multiple quantitative indicators demonstrating sustained market growth, continuous technological innovation, and expanding use cases that suggest long-term market viability.")

    add_para(doc, "From a demand perspective, the Indian power bank market's 11.5% CAGR and global market's 6.8% CAGR indicate robust, sustained growth rather than the explosive growth typical of pure sunrise products or the decline characteristic of sunset products. The smartphone dependency factor provides structural support for evergreen classification - with average daily usage of 4.5 hours and manufacturers prioritizing device thinness over battery size, a persistent gap exists between power consumption and device battery capacity that power banks effectively address.")

    add_para(doc, "Technological innovation continues driving product evolution, a characteristic typically associated with sunrise products. The 25% annual growth in wireless charging power banks, emergence of GaN-based high-power devices, and integration of smart features like battery health monitoring indicate ongoing category development. Unlike sunset products that see minimal innovation and declining R&D investment, power banks continue attracting significant product differentiation efforts from both established brands and new market entrants.")

    doc.add_page_break()

    # ==================== 2. BACKEND ANALYSIS ====================
    add_heading(doc, "2. BACKEND ANALYSIS: PRODUCTION AND COST ECONOMICS")

    add_para(doc, "The backend analysis of power bank manufacturing reveals a complex production ecosystem characterized by global supply chains, significant economies of scale, and evolving regulatory requirements. Understanding these production dynamics is essential for analyzing market structure, pricing strategies, and competitive positioning within the industry.")

    add_heading(doc, "2.1 Production Modality and Technology", level=2)

    add_para(doc, "Power bank manufacturing involves a multi-stage assembly process that combines battery cell production, electronic circuit assembly, enclosure fabrication, and comprehensive quality testing. The production process can be categorized into three distinct modalities based on scale, automation level, and target market positioning, each with distinct economic characteristics and competitive implications.")

    add_para(doc, "Small-scale manufacturing operations require initial investments of Rs. 20-40 lakhs and typically produce 5,000-10,000 units monthly. These facilities rely primarily on manual assembly with basic testing equipment and quality control through sampling-based inspection. Labour intensity remains high at 15-20 workers per 1,000 units produced. Such operations typically serve local markets, unbranded segments, or specialize in customized corporate gifting products where volumes are limited but margins can be higher.")

    add_para(doc, "Medium-scale manufacturing facilities invest Rs. 75 lakhs to 2 crores and achieve monthly capacities of 50,000-100,000 units. These operations employ semi-automated Surface Mount Technology (SMT) lines for PCB assembly, reducing labour requirements to 8-12 workers per 1,000 units while improving consistency. Quality systems include BIS certification compliance and structured testing protocols. These manufacturers typically produce for regional brands or serve as OEM/ODM suppliers to established companies.")

    add_para(doc, "Large-scale manufacturing facilities represent significant capital commitments of Rs. 5-15 crores with monthly capacities exceeding 500,000 units. Fully automated SMT lines achieve placement rates of 25,000-50,000 components per hour with defect rates below 0.5%. Labour efficiency improves dramatically to 3-5 workers per 1,000 units. These facilities incorporate advanced testing including cell matching for uniform performance, thermal cycling tests, and accelerated life testing. Major brands like Xiaomi, Anker, and Samsung operate or contract with such facilities to ensure consistent quality at scale.")

    add_table_title(doc, "Table 4: Manufacturing Scale and Efficiency Metrics")
    create_table(doc,
        ['Scale', 'Investment (Rs.)', 'Capacity/Month', 'Workers/1000 Units', 'Automation'],
        [
            ('Small-Scale', '20-40 Lakhs', '5,000-10,000', '15-20', 'Manual'),
            ('Medium-Scale', '75L-2 Crores', '50,000-100,000', '8-12', 'Semi-Auto'),
            ('Large-Scale', '5-15 Crores', '500,000+', '3-5', 'Fully Auto'),
        ],
        col_widths=[1.1, 1.3, 1.2, 1.3, 1.0])
    doc.add_paragraph()

    add_heading(doc, "2.2 Raw Materials and Factor Market", level=2)

    add_para(doc, "The power bank supply chain is globally distributed, with significant concentration of critical components in East Asian manufacturing hubs. Understanding the raw material ecosystem is essential for analyzing cost structures, supply chain risks, and competitive dynamics within the industry. The factor market for power bank manufacturing exhibits characteristics of derived demand, with input requirements directly linked to final product demand.")

    add_para(doc, "Battery cells represent the single largest cost component at 35-40% of total manufacturing cost. The global battery cell market is dominated by Chinese manufacturers, with CATL commanding 34% global share, followed by BYD at 16%. Korean manufacturers Samsung SDI and LG Energy Solution together hold approximately 25% market share, offering premium cells with higher energy density and longer cycle life. Japanese manufacturer Panasonic maintains approximately 10% share, primarily serving premium applications. Indian power bank manufacturers predominantly source cells from Chinese suppliers due to cost advantages of 15-25% compared to Korean alternatives, though this creates supply chain concentration risk.")

    add_para(doc, "Printed Circuit Boards (PCBs) and Power Management Integrated Circuits (PMICs) constitute 18-22% of manufacturing cost. These components require sophisticated semiconductor fabrication capabilities concentrated in Taiwan, South Korea, and China. The PMIC manages critical functions including voltage conversion, charging protocol negotiation, and protection against overcharging, over-discharging, and short circuits. Premium power banks incorporate dedicated chips for fast charging protocols, adding USD 0.50-1.50 to component costs but enabling significant product differentiation.")

    add_para(doc, "Enclosure materials including ABS plastic, polycarbonate, and aluminum represent 12-15% of costs. ABS plastic remains the most common material due to its balance of durability, weight, and cost. Aluminum enclosures command premium pricing but offer superior heat dissipation crucial for fast-charging applications. Connectors, cables, and accessories contribute an additional 11-14% to total costs, with USB-C Power Delivery certification requiring specific e-marker chips in cables.")

    add_table_title(doc, "Table 5: Raw Material Cost Structure")
    create_table(doc,
        ['Component', 'Source', 'Cost (USD)', 'Cost (Rs.)', '% Total'],
        [
            ('Li-ion/LiPo Cells', 'China 70%, Korea 25%', '2.50-4.00', '210-335', '35-40%'),
            ('PCB & PMIC', 'China, Taiwan', '1.50-2.50', '125-210', '18-22%'),
            ('ABS/Al Housing', 'India, China', '1.00-2.00', '85-170', '12-15%'),
            ('USB Ports/Connectors', 'China', '0.50-1.00', '40-85', '6-8%'),
            ('Cables & Accessories', 'China, India', '0.40-0.80', '35-65', '5-6%'),
            ('Packaging', 'Local', '0.50-1.00', '40-85', '5-7%'),
        ],
        col_widths=[1.3, 1.4, 1.0, 0.9, 0.7])
    doc.add_paragraph()

    add_heading(doc, "2.3 Cost Structure Analysis", level=2)

    add_para(doc, "Analyzing the complete cost structure of power bank manufacturing reveals the interplay between variable and fixed costs, the significance of economies of scale, and the cost reduction potential through operational improvements. A detailed cost breakdown for a standard 10,000 mAh fast-charging power bank illustrates these dynamics and provides insights into pricing flexibility and margin structures.")

    add_para(doc, "Variable costs constitute approximately 85-90% of total manufacturing cost, dominated by material inputs. Battery cells at Rs. 210-335 represent the largest single cost element, followed by PCB and power management components at Rs. 125-210. This high variable cost proportion creates strong incentives for volume manufacturing and supply chain optimization. Material cost volatility, particularly in lithium pricing which has ranged from USD 15,000 to USD 80,000 per metric ton, directly impacts profitability and pricing decisions.")

    add_para(doc, "Fixed costs including facility depreciation, equipment maintenance, and administrative overhead represent 10-15% of total costs at typical production volumes. However, the fixed cost proportion varies significantly with scale. Small manufacturers operating at 10,000 units monthly face fixed costs of approximately Rs. 200 per unit, while large manufacturers producing 200,000+ units reduce fixed costs to Rs. 35 per unit - an 82.5% reduction in per-unit fixed cost allocation that creates substantial competitive advantages.")

    add_table_title(doc, "Table 6: Complete Cost Breakdown per 10,000 mAh Power Bank")
    create_table(doc,
        ['Cost Component', 'USD', 'INR', '%', 'Type'],
        [
            ('Battery Cells', '2.50-4.00', '210-335', '35-40%', 'Variable'),
            ('PCB & Power Management', '1.50-2.50', '125-210', '18-22%', 'Variable'),
            ('Housing (ABS/Aluminum)', '1.00-2.00', '85-170', '12-15%', 'Variable'),
            ('USB Ports & Connectors', '0.50-1.00', '40-85', '6-8%', 'Variable'),
            ('Labour & Assembly', '0.30-0.60', '25-50', '4-5%', 'Semi-Var'),
            ('Quality Testing', '0.20-0.40', '17-35', '2-3%', 'Variable'),
            ('Overheads/Depreciation', '0.50-0.80', '40-65', '6-8%', 'Fixed'),
            ('TOTAL', '7.20-12.00', '600-1,000', '100%', '-'),
        ],
        col_widths=[1.6, 0.9, 0.9, 0.7, 0.8])
    doc.add_paragraph()

    add_para(doc, "Cost curve analysis across different production volumes demonstrates significant economies of scale in power bank manufacturing. At 10,000 units monthly production, Average Total Cost (ATC) stands at Rs. 950, comprising Rs. 200 AFC and Rs. 750 AVC. Increasing production to 50,000 units reduces ATC to Rs. 750, a 21% reduction. At 100,000 units, ATC falls to Rs. 690, and at 200,000 units reaches Rs. 655 - a cumulative 31% reduction from the baseline. The Minimum Efficient Scale (MES) is approximately 100,000 units/month where ATC approaches MC and further scale benefits become marginal.")

    add_table_title(doc, "Table 7: Cost Curve Analysis at Different Production Levels")
    create_table(doc,
        ['Production (Units/Month)', 'AFC (Rs.)', 'AVC (Rs.)', 'ATC (Rs.)', 'MC (Rs.)'],
        [
            ('10,000', '200', '750', '950', '720'),
            ('50,000', '80', '670', '750', '650'),
            ('100,000', '50', '640', '690', '620'),
            ('200,000', '35', '620', '655', '610'),
        ],
        col_widths=[1.6, 0.9, 0.9, 0.9, 0.9])
    doc.add_paragraph()

    add_heading(doc, "2.4 Supply Chain Analysis", level=2)

    add_para(doc, "The power bank value chain encompasses multiple stages from raw material sourcing through final retail, with each stage adding value and extracting margins. Understanding the margin structure across the value chain illuminates competitive dynamics and identifies opportunities for value capture optimization by different market participants.")

    add_para(doc, "Manufacturing represents the primary value creation stage, with manufacturers capturing 25-35% gross margins on production. This margin compensates for capital investment in equipment and facilities, technology development and R&D, quality assurance systems, and production management overhead. Manufacturers with integrated cell production capabilities or long-term supply agreements achieve superior margins compared to those dependent on spot market procurement.")

    add_para(doc, "Distribution channels add significant costs but provide essential market access. National distributors typically operate on 8-12% margins, handling logistics, inventory financing, and retailer relationships across geographic regions. Regional distributors add another 5-8%, providing last-mile distribution to smaller retailers. Retail margins of 18-25% reflect the costs of customer acquisition, display space, sales staff, and return handling. E-commerce platforms operate on lower direct margins of 8-15% but charge commissions and advertising fees. Direct-to-Consumer (D2C) channels offer 35-45% margins but require significant investment in customer acquisition and brand building.")

    add_table_title(doc, "Table 8: Value Chain Margin Structure")
    create_table(doc,
        ['Stage', 'Margin (%)', 'Value Add (Rs.)', 'Cumulative Price'],
        [
            ('Raw Materials', '-', '600-850', '600-850'),
            ('Manufacturing', '25-35%', '180-280', '780-1,130'),
            ('National Distributor', '8-12%', '70-120', '850-1,250'),
            ('Retailer', '18-25%', '180-300', '1,080-1,640'),
            ('GST (18%)', '18%', '195-295', '1,275-1,935'),
            ('Final MRP', '-', '-', '1,299-1,999'),
        ],
        col_widths=[1.5, 1.0, 1.3, 1.3])
    doc.add_paragraph()

    add_heading(doc, "2.5 Wholesale Markets in Mumbai", level=2)

    add_para(doc, "Mumbai serves as a major distribution hub for power banks and electronic accessories in Western India. Several concentrated wholesale markets enable bulk purchasing at competitive rates, serving retailers, corporate buyers, and resellers across Maharashtra and neighboring states. Understanding these wholesale channels is essential for analyzing the complete distribution infrastructure.")

    add_para(doc, "Lamington Road in Grant Road area represents Mumbai's oldest and most prominent electronics wholesale market. Spanning approximately 500 meters, this market hosts over 200 shops specializing in electronic components, mobile accessories, and power banks. Wholesale buyers can procure power banks at 15-25% below MRP with minimum order quantities of 10-50 units. Major wholesalers include Vijay Sales Wholesale, Kohinoor Electronics, and Prime Electronics, offering brands like Mi, Ambrane, Portronics, and unbranded options. Operating hours are typically 10 AM to 8 PM, Monday through Saturday.")

    add_para(doc, "Manish Market near Crawford Market is another significant wholesale destination, particularly for budget and unbranded power banks. This market offers the lowest prices in Mumbai, with discounts of 25-40% below retail. However, quality verification is essential as BIS certification compliance varies among sellers. Heera Panna in Haji Ali has evolved as a premium electronics wholesale market, focusing on branded and authorized products with 10-18% below MRP and better warranty support. Online platforms like IndiaMART and TradeIndia have emerged as virtual wholesale markets, enabling price discovery and bulk ordering without physical market visits.")

    add_heading(doc, "2.6 Government Policies and Regulations", level=2)

    add_para(doc, "The regulatory environment for power banks in India has evolved significantly, creating compliance requirements that impact manufacturing costs, market structure, and competitive dynamics. Multiple regulatory frameworks address taxation, product safety, environmental responsibility, and domestic manufacturing incentives.")

    add_para(doc, "The Goods and Services Tax applies at 18% to power banks classified under HSN 8507 for lithium-ion accumulators, adding Rs. 180-300 to consumer prices. Bureau of Indian Standards certification under IS 17018:2018 became mandatory in 2019, with certification costs of Rs. 50,000-1.5 lakhs per model creating entry barriers for small manufacturers. Battery Waste Management Rules 2022 impose Extended Producer Responsibility requiring 70% collection by 2024-25 with compliance costs of 2-4% of revenue. The PLI Scheme for Advanced Chemistry Cell offers 5-20% incentives for domestic manufacturing. Import duty stands at 15% on finished goods and 5-10% on components, influencing make-vs-buy decisions.")

    doc.add_page_break()

    # ==================== 3. FRONTEND ANALYSIS ====================
    add_heading(doc, "3. FRONTEND ANALYSIS: MARKET DEMAND AND REVENUE DYNAMICS")

    add_para(doc, "The frontend analysis examines demand determinants, elasticity characteristics, market structure, pricing mechanisms, and revenue dynamics in the power bank market. Understanding these consumer-facing dynamics is essential for strategic positioning, pricing optimization, and competitive strategy development in this rapidly evolving market.")

    add_heading(doc, "3.1 Nature of Demand", level=2)

    add_para(doc, "Power bank demand is characterized by derived demand, meaning it derives from the demand for primary products that power banks complement, primarily smartphones and tablets. As smartphone adoption increases, particularly in emerging markets like India, power bank demand correspondingly expands. This derived demand relationship implies that power bank market growth is structurally linked to the broader mobile device ecosystem and its evolution.")

    add_para(doc, "Multiple determinants influence power bank demand with varying degrees of impact. Product price exhibits an inverse relationship with quantity demanded, consistent with the law of demand. However, the price elasticity varies significantly across market segments, with budget consumers highly sensitive to price changes while premium segment consumers prioritize features, brand reputation, and aesthetic design over price considerations.")

    add_para(doc, "Consumer income positively correlates with power bank demand, particularly for premium products. Rising disposable incomes in urban India, with monthly household incomes exceeding Rs. 50,000 among target consumers, enable spending on accessory categories. Smartphone prices demonstrate complementary good relationships with power banks - when smartphone prices decrease and adoption increases, power bank demand rises correspondingly. Conversely, competing products like car chargers and wireless charging pads exhibit substitute good relationships. Seasonal demand patterns significantly impact sales, with festive seasons from October through January generating 40-60% volume spikes.")

    add_table_title(doc, "Table 9: Demand Determinants and Quantitative Impact")
    create_table(doc,
        ['Determinant', 'Current Status', 'Impact on Demand', 'Elasticity'],
        [
            ('Product Price', 'Rs. 500-5,000 range', 'Inverse relationship', 'PED: -0.8 to -1.2'),
            ('Consumer Income', 'Rs. 50K+ urban income', 'Positive correlation', 'YED: +0.95 to +1.15'),
            ('Smartphone Prices', 'Avg. Rs. 12,000', 'Complementary effect', 'XED: -0.3 to -0.5'),
            ('Competitor Prices', 'Rs. 800-2,500 avg.', 'Substitution effect', 'XED: +0.2 to +0.4'),
            ('Festive Seasons', 'Oct-Jan peak', '40-60% volume spike', 'Index: 1.5'),
        ],
        col_widths=[1.3, 1.4, 1.4, 1.3])
    doc.add_paragraph()

    add_heading(doc, "3.2 Elasticity of Demand Analysis", level=2)

    add_para(doc, "Price elasticity of demand varies substantially across market segments, creating opportunities for differentiated pricing strategies. The budget segment, comprising products priced Rs. 500-1,000, exhibits elastic demand with PED of -1.3 to -1.5. Consumers in this segment are highly price-sensitive, readily switching between brands based on promotional pricing and discounts. A 10% price reduction generates 13-15% quantity increase, resulting in 3-5% revenue increase despite lower unit prices - making promotional pricing an effective strategy.")

    add_para(doc, "The mid-range segment at Rs. 1,000-2,000 demonstrates approximately unit elastic demand with PED of -0.9 to -1.1. Price changes generate proportional quantity responses, maintaining relatively stable revenue regardless of pricing direction. This segment balances price sensitivity with feature preferences, responding to both promotional pricing and value-added features like fast charging support.")

    add_para(doc, "Premium segments priced Rs. 2,000-3,500 exhibit inelastic demand with PED of -0.6 to -0.8. Brand reputation, advanced features, and aesthetic design drive purchase decisions more than price considerations. A 10% price reduction generates only 6-8% quantity increase, resulting in 2-4% revenue decline - suggesting premium brands should avoid price competition. Ultra-premium products above Rs. 3,500 demonstrate highly inelastic demand with PED of -0.4 to -0.6, where status signaling and brand alignment drive purchase decisions.")

    add_para(doc, "Income elasticity analysis reveals power banks function as normal goods across all segments. Basic power banks demonstrate income elasticity of +0.95 to +1.05, indicating proportional demand response with slight necessity characteristics. Premium power banks exhibit income elasticity of +1.05 to +1.15, indicating mild luxury good characteristics where demand increases faster than income. Cross-price elasticity with smartphones (-0.3 to -0.5) confirms complementary goods relationship, while positive XED with car chargers (+0.2 to +0.4) indicates substitute goods relationship.")

    add_table_title(doc, "Table 10: Price Elasticity of Demand by Market Segment")
    create_table(doc,
        ['Segment', 'Price Range', 'PED Value', 'Type', '10% Price Cut Impact'],
        [
            ('Budget', 'Rs. 500-1,000', '-1.3 to -1.5', 'Elastic', '+3% to +5% revenue'),
            ('Mid-Range', 'Rs. 1,000-2,000', '-0.9 to -1.1', 'Unit Elastic', '0% to +1% revenue'),
            ('Premium', 'Rs. 2,000-3,500', '-0.6 to -0.8', 'Inelastic', '-2% to -4% revenue'),
            ('Ultra-Premium', 'Rs. 3,500+', '-0.4 to -0.6', 'Inelastic', '-4% to -6% revenue'),
        ],
        col_widths=[1.1, 1.2, 1.0, 1.0, 1.6])
    doc.add_paragraph()

    add_heading(doc, "3.3 Market Structure Analysis", level=2)

    add_para(doc, "The Indian power bank market operates under MONOPOLISTIC COMPETITION, characterized by numerous sellers offering differentiated products, relatively low entry barriers, significant non-price competition, and free entry and exit. This market structure creates specific competitive dynamics and strategic imperatives for market participants seeking sustainable competitive advantages.")

    add_para(doc, "Market concentration analysis using the Herfindahl-Hirschman Index yields values of 850-1,100, indicating moderate competition without dominant players. No single brand commands dominant market share, with Xiaomi/Mi leading at 18-22%, followed by Anker (12-15%), Samsung (8-10%), and Ambrane (7-9%). The top five brands collectively control approximately 50-55% of the market, leaving substantial space for smaller players and new entrants. This distribution reflects relatively low barriers to entry and the ability of differentiated products to carve sustainable market positions.")

    add_para(doc, "Product differentiation manifests across multiple dimensions including capacity, charging speed, design aesthetics, brand image, and ecosystem integration. Xiaomi differentiates through value proposition and ecosystem integration with smartphones. Anker emphasizes quality and technology leadership. Samsung leverages brand reputation and smartphone ecosystem integration. Entry barriers remain relatively low with minimum viable investments of Rs. 20-50 lakhs, though effective competition requires BIS certification, distribution relationships, and brand building investments. Non-price competition consumes 8-15% of revenue through marketing, influencer partnerships, and packaging investments.")

    add_table_title(doc, "Table 11: Indian Power Bank Market Share Analysis (2024)")
    create_table(doc,
        ['Brand', 'Market Share', 'Revenue (Cr.)', 'Segment', 'YoY Growth'],
        [
            ('Xiaomi/Mi', '18-22%', '1,450-1,780', 'Value + Mid-range', '+12%'),
            ('Anker', '12-15%', '970-1,210', 'Premium', '+8%'),
            ('Samsung', '8-10%', '645-810', 'Premium Ecosystem', '+5%'),
            ('Ambrane', '7-9%', '565-730', 'Budget + Value', '+15%'),
            ('Realme/OnePlus', '6-8%', '485-645', 'Mid-range', '+22%'),
            ('Others/Unbranded', '35-40%', '2,830-3,230', 'Budget', '+6%'),
        ],
        col_widths=[1.2, 1.0, 1.1, 1.3, 0.9])
    doc.add_paragraph()

    add_heading(doc, "3.4 Pricing Mechanism", level=2)

    add_para(doc, "Pricing strategy in the power bank market reflects cost-plus fundamentals modified by competitive positioning, segment targeting, and channel requirements. The price build-up from manufacturing cost to consumer price reveals value distribution across the supply chain and pricing flexibility available to manufacturers at different stages.")

    add_para(doc, "Cost-plus pricing establishes the floor price, with manufacturers targeting 25-35% gross margins over variable production costs. Fixed cost allocation adds Rs. 80-120 per unit at typical production volumes. Channel margins are relatively fixed by industry practice: distributors 10%, retailers 20%. Premium brands price 30-50% above category averages justified by perceived quality and brand reputation, while value brands price 10-20% below to drive volume. Psychological pricing techniques using price points ending in 99 (Rs. 1,299, Rs. 1,999) are standard. Promotional pricing during festive seasons can temporarily reduce prices by 20-40%.")

    add_heading(doc, "3.5 Revenue and Selling Costs", level=2)

    add_para(doc, "Revenue analysis across distribution channels reveals evolving patterns informing strategic channel investments. E-commerce channels command 45-50% revenue share with 18% YoY growth, offering lower distribution costs but intense price competition. Modern retail through Croma and Reliance Digital represents 20-25% with 8% growth. D2C channels show fastest growth at 35% YoY, currently 8-12% of revenue, offering 35-45% margins by eliminating intermediaries.")

    add_para(doc, "Marketing and selling costs consume 15-25% of revenue across multiple channels: digital advertising (3-5% of revenue, ROAS 3.5-5.0x), influencer marketing (2-4%, engagement 4-8%), e-commerce platform ads (2-3%, ACoS 8-15%), e-commerce commission (5-15% of sale), and retail trade marketing (1-2%). Customer Acquisition Cost ranges Rs. 150-300 depending on channel and brand positioning.")

    doc.add_page_break()

    # ==================== 4. SUGGESTIONS ====================
    add_heading(doc, "4. SUGGESTIONS FOR IMPROVING PRODUCTION AND DELIVERY EFFICIENCY")

    add_para(doc, "Based on the comprehensive analysis of production economics, cost structures, and market dynamics, several strategic recommendations emerge for improving production efficiency, reducing costs, and enhancing market competitiveness. These recommendations are prioritized by expected return on investment and implementation feasibility for different scales of manufacturers.")

    add_para(doc, "D2C Channel Development offers the fastest payback period of 8-12 months with investment of Rs. 20-80 lakhs. Building D2C capabilities through e-commerce websites, payment integration, and customer service infrastructure captures 20-30% higher margins by eliminating intermediaries. D2C success requires parallel investment in digital marketing and brand building to drive traffic cost-effectively, but represents the most accessible path to margin improvement for brands with established recognition.")

    add_para(doc, "Just-in-Time Inventory Systems offer attractive short-term returns with investment of Rs. 50-150 lakhs in demand forecasting systems, supplier integration platforms, and warehouse optimization. Reducing inventory holding from 60-90 days to 20-30 days frees 20-30% of working capital currently tied in stock. The 12-18 month payback makes this an attractive priority for mid-sized manufacturers seeking immediate efficiency gains without major capital commitments.")

    add_para(doc, "Quality Control System investments of Rs. 15-25 lakhs improve defect rates from typical 3% to best-in-class 0.5%. Reduced warranty claims, return handling costs, and brand reputation damage deliver 12-18 month payback. Investment areas include incoming inspection equipment, in-process monitoring, and automated final testing. Quality improvements also enable premium positioning and higher price realization.")

    add_para(doc, "Production Scale-up to achieve minimum efficient scale of 100,000+ units monthly delivers 15-20% unit cost reductions with 18-24 month payback. Investment of Rs. 2-5 crores in additional capacity and automation equipment improves operational leverage. Automation Investment of Rs. 1-3 crores reduces labour costs by 40-60% with 24-36 month payback. Domestic Cell Manufacturing (Rs. 50-200 Cr) offers 15-25% cost reduction but requires 3-5 year commitment, suitable for large manufacturers or industry consortiums.")

    add_table_title(doc, "Table 12: Strategic Recommendations with Expected Impact")
    create_table(doc,
        ['Strategy', 'Investment', 'Expected Benefit', 'Payback'],
        [
            ('D2C Channel Development', 'Rs. 20-80 L', '20-30% higher margins', '8-12 months'),
            ('JIT Inventory System', 'Rs. 50-150 L', '20-30% working capital saving', '12-18 months'),
            ('Quality Control Systems', 'Rs. 15-25 L', 'Defect: 3% to 0.5%', '12-18 months'),
            ('Production Scale-up (100K+)', 'Rs. 2-5 Cr', '15-20% lower unit cost', '18-24 months'),
            ('Automation Investment', 'Rs. 1-3 Cr', '40-60% labour cost reduction', '24-36 months'),
        ],
        col_widths=[1.7, 1.2, 1.8, 1.2])
    doc.add_paragraph()

    doc.add_page_break()

    # ==================== 5. LESSONS LEARNED ====================
    add_heading(doc, "5. LESSONS LEARNED")

    add_para(doc, "This comprehensive microeconomic analysis of the power bank industry reveals several key learnings applicable to understanding consumer electronics markets, manufacturing economics, and competitive strategy. These lessons synthesize the detailed analysis into actionable insights for students of microeconomics and industry practitioners.")

    add_para(doc, "1. Cost Structure Dominance of Variable Costs: With 85-90% of total costs being variable and dominated by raw material inputs (particularly battery cells at 35-40%), volume-based strategies become essential for profitability. The high variable cost proportion means that production scale directly impacts competitiveness, explaining the industry's consolidation trend toward larger manufacturers with superior purchasing power and operational efficiency.")

    add_para(doc, "2. Significant Economies of Scale: The 31% reduction in Average Total Cost between 10,000 and 200,000 units monthly production illustrates why minimum efficient scale matters. Manufacturers operating below MES of approximately 100,000 units face structural cost disadvantages of 15-40% compared to larger competitors, constraining pricing flexibility and profit margins while limiting resources available for marketing and innovation.")

    add_para(doc, "3. Segment-Specific Elasticity Enables Differentiated Pricing: The variation from elastic demand in budget segments (PED -1.3 to -1.5) to inelastic demand in premium segments (PED -0.4 to -0.6) creates opportunities for differentiated pricing approaches. Budget brands should compete aggressively on price to capture volume, while premium brands should emphasize differentiation and resist price competition that would erode margins without proportional volume gains.")

    add_para(doc, "4. Normal Good Classification Supports Market Growth: Income elasticity of +0.95 to +1.15 confirms power banks function as normal goods with upward potential as Indian incomes rise. The mild luxury characteristics of premium products (YED >1) suggest premium segment growth will outpace overall market growth as consumer purchasing power increases, informing product portfolio strategy toward premium offerings.")

    add_para(doc, "5. Monopolistic Competition Rewards Differentiation: The combination of low barriers to entry, significant product differentiation, and active non-price competition creates an environment where brand building and innovation matter as much as cost efficiency. Marketing investments of 15-25% of revenue reflect the competitive necessity of differentiation rather than optional discretionary spending.")

    add_para(doc, "6. Value Chain Analysis Highlights D2C Opportunity: Manufacturers capture only 50-58% of consumer price, with the remainder distributed across distribution, retail, and taxation. This margin compression highlights the strategic value of direct-to-consumer channels that can capture 35-45% margins by eliminating intermediaries. The D2C opportunity represents the most accessible path to margin improvement for established brands with customer acquisition capabilities.")

    doc.add_page_break()

    # References
    add_heading(doc, "REFERENCES")

    refs = [
        "1. Grand View Research (2024). Power Bank Market Size, Share & Trends Analysis Report, 2024-2030.",
        "2. Mordor Intelligence (2024). India Power Bank Market - Growth, Trends, and Forecasts (2024-2033).",
        "3. Bureau of Indian Standards. IS 17018:2018 - Specification for Portable Secondary Lithium Cells and Batteries.",
        "4. Central Board of Indirect Taxes and Customs. GST Rate Schedule for Electronic Products - HSN 8507.",
        "5. Ministry of Environment, Forest and Climate Change. Battery Waste Management Rules, 2022.",
        "6. Ministry of Heavy Industries. Production Linked Incentive Scheme for Advanced Chemistry Cell Battery Storage.",
        "7. Statista (2024). Smartphone Users in India 2024-2028: Market Report.",
        "8. India Brand Equity Foundation (2024). Indian Electronics Industry Report.",
        "9. Counterpoint Research (2024). India Smartphone Accessories Market Report Q3 2024.",
        "10. BloombergNEF (2024). Lithium-Ion Battery Price Survey 2024."
    ]

    for ref in refs:
        para = doc.add_paragraph()
        run = para.add_run(ref)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        para.paragraph_format.space_after = Pt(6)

    return doc


def add_heading(doc, text, level=1):
    """Add a heading - 14pt bold as per concept note"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    para.space_before = Pt(14) if level == 1 else Pt(10)
    para.space_after = Pt(8)


def add_para(doc, text):
    """Add a justified paragraph with selective bold for key metrics - Font 12pt, Line spacing 1.5"""
    import re
    para = doc.add_paragraph()

    # Pattern to match key metrics that should be bolded (non-capturing for split)
    bold_pattern = r'((?:Rs\.|USD|INR)\s*[\d,\.]+(?:\s*(?:billion|million|crores?|lakhs?|L))?(?:\s*to\s*(?:Rs\.|USD|INR)?\s*[\d,\.]+(?:\s*(?:billion|million|crores?|lakhs?|L))?)?|\d+(?:,\d{3})*(?:\.\d+)?\s*(?:billion|million|crores?|lakhs?|Wh/kg|mAh|CAGR|YoY)|\d+(?:\.\d+)?%(?:\s*(?:CAGR|YoY))?|\d+-\d+%|PED:?\s*[\-\d\.]+(?:\s*to\s*[\-\d\.]+)?|YED:?\s*[\+\-\d\.]+(?:\s*to\s*[\+\-\d\.]+)?|XED:?\s*[\+\-\d\.]+(?:\s*to\s*[\+\-\d\.]+)?|\d+(?:\.\d+)?W)'

    # Use finditer to get positions and process text
    last_end = 0
    for match in re.finditer(bold_pattern, text, flags=re.IGNORECASE):
        # Add text before match (normal)
        if match.start() > last_end:
            run = para.add_run(text[last_end:match.start()])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        # Add matched text (bold)
        run = para.add_run(match.group())
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True

        last_end = match.end()

    # Add remaining text after last match
    if last_end < len(text):
        run = para.add_run(text[last_end:])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = 1.5
    para.space_after = Pt(8)


def add_table_title(doc, title):
    """Add a table title - 14pt bold for uniformity"""
    para = doc.add_paragraph()
    run = para.add_run(title)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.space_before = Pt(10)
    para.space_after = Pt(4)


def add_page_numbers(doc):
    """Add page numbers"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


if __name__ == "__main__":
    print("Creating 20-PAGE Power Bank Assignment...")
    print("=" * 60)
    doc = create_document()
    output_path = "/mnt/e/AI and Projects/MMS-Prep/Eco Assign/Power_Bank_Microeconomics_Assignment.docx"
    doc.save(output_path)
    print(f"Saved to: {output_path}")
    print("=" * 60)
    print("Features:")
    print("  - ~5500 words, targeting 20 pages")
    print("  - 12 essential tables")
    print("  - Margins: 1, 1, 1, 1.5 (Top, Right, Bottom, Left)")
    print("  - Body: 12pt Times New Roman, 1.5 line spacing")
    print("  - Headings & Table Titles: 14pt Bold")
    print("  - Full justification with pagination")
